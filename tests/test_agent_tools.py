from __future__ import annotations

import logging
import base64
import tempfile
import unittest
from unittest.mock import patch
from pathlib import Path

import pandas as pd
from docx import Document
from openpyxl import load_workbook

from analysis.dry_curves import dry_statistics
from analysis.event_response import analyze_event_response
from analysis.patterns import analyze_patterns
from analysis.rainfall import rainfall_events
from analysis.rdii import analyze_rdii
from analysis.reporting import build_report
from analysis.risk import assess_risk
from agent.deps import AgentDeps, AgentSettings, Paths, SessionState, ensure_directories
from agent.tools.module_tools import (
    _save_pattern_curve_pngs,
    _save_rdii_curve_pngs,
    analyze_patterns_impl,
    analyze_rdii_impl,
    analyze_rainfall_impl,
    check_data_impl,
    data_filter_impl,
)
from agent.tools.python_tool import run_python_impl


def make_deps(root: Path) -> AgentDeps:
    paths = Paths.from_root(root)
    ensure_directories(paths)
    return AgentDeps(
        paths=paths,
        settings=AgentSettings(model="test", base_url=None, api_key=None),
        logger=logging.getLogger("test"),
        session=SessionState(),
        project_notes="",
    )


def write_flow(deps: AgentDeps) -> None:
    pd.DataFrame(
        {
            "timestamp": pd.date_range("2026-01-01", periods=5, freq="min"),
            "flow": [1, 2, 3, 4, 5],
            "level": [0.1, 0.2, 0.3, 0.4, 0.5],
            "velocity": [0.2, 0.2, 0.3, 0.3, 0.4],
        }
    ).to_csv(deps.paths.flow_dir / "100_W1.csv", index=False)
    deps.paths.rainfall_file.write_text("timestamp,rain\n2026-01-01,0\n", encoding="utf-8")


def write_filter_flow(deps: AgentDeps) -> None:
    rows = 4 * 1440
    pd.DataFrame(
        {
            "timestamp": pd.date_range("2026-01-01", periods=rows, freq="min"),
            "flow": [1.0] * rows,
            "level": [0.2] * rows,
            "velocity": [0.3] * rows,
        }
    ).to_csv(deps.paths.flow_dir / "100_W1.csv", index=False)
    pd.DataFrame(
        {
            "timestamp": pd.date_range("2026-01-01", periods=4, freq="D"),
            "rain": [0.0, 0.0, 3.0, 0.0],
        }
    ).to_csv(deps.paths.rainfall_file, index=False)


def write_rdii_sample(deps: AgentDeps) -> None:
    timestamps = pd.date_range("2026-01-01", periods=4 * 1440, freq="min")
    pd.DataFrame(
        {
            "timestamp": timestamps,
            "flow": [1.0] * (2 * 1440) + [3.0] * 1440 + [1.0] * 1440,
            "level": [0.2] * (4 * 1440),
            "velocity": [0.3] * (4 * 1440),
        }
    ).to_csv(deps.paths.flow_dir / "100_W1.csv", index=False)
    pd.DataFrame(
        {
            "timestamp": pd.date_range("2026-01-01", periods=4 * 1440, freq="min"),
            "rain": [0.0] * (2 * 1440) + [0.0, 3.0] + [0.0] * (2 * 1440 - 2),
        }
    ).to_csv(deps.paths.rainfall_file, index=False)
    pd.DataFrame({"点位编号": ["W1", "W2"], "管径": [1.0, 1.2]}).to_excel(
        deps.paths.site_info_file,
        index=False,
    )
    analyze_rainfall_impl(deps, output="events")


class AgentToolTests(unittest.TestCase):
    def test_check_data_reads_flow_summary(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            deps = make_deps(Path(tmp))
            write_flow(deps)
            for point_id in ("W10", "W2"):
                pd.DataFrame(
                    {
                        "timestamp": pd.date_range("2026-01-01", periods=5, freq="min"),
                        "flow": [1, 2, 3, 4, 5],
                        "level": [0.1, 0.2, 0.3, 0.4, 0.5],
                        "velocity": [0.2, 0.2, 0.3, 0.3, 0.4],
                    }
                ).to_csv(deps.paths.flow_dir / f"100_{point_id}.csv", index=False)
            result = check_data_impl(deps)
            self.assertEqual(result["status"], "ok")
            self.assertIn("数据收集率统计完成", result["summary"])
            df = pd.read_excel(deps.paths.combined_xlsx, sheet_name="数据收集率统计")
            self.assertEqual(
                list(df.columns),
                ["点位编号", "监测数据条数", "监测天数", "理论数据条数", "数据收集率(%)"],
            )
            self.assertEqual(df["点位编号"].tolist(), ["W1", "W2", "W10"])
            self.assertNotIn("数据体检", load_workbook(deps.paths.combined_xlsx).sheetnames)

    def test_data_filter_writes_filter_result(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            deps = make_deps(Path(tmp))
            write_filter_flow(deps)
            result = data_filter_impl(deps)
            self.assertEqual(result["status"], "ok")
            self.assertEqual(result["data"]["selected"], {"W1": ["2026-01-02"]})
            workbook = load_workbook(deps.paths.filter_result)
            sheet = workbook["筛选结果"]
            self.assertEqual(sheet.cell(row=1, column=1).value, "点位编号")
            self.assertEqual(sheet.cell(row=2, column=1).value, "当天雨量")
            self.assertEqual(sheet.cell(row=1, column=sheet.max_column).value, "筛选说明")
            self.assertTrue(str(sheet.cell(row=3, column=3).fill.start_color.index).upper().endswith("92D050"))

    def test_patterns_use_filter_result_instead_of_internal_zero_filter(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            deps = make_deps(Path(tmp))
            rows = []
            for day_idx, day in enumerate(pd.date_range("2026-01-01", periods=4, freq="D")):
                for minute, ts in enumerate(pd.date_range(day, periods=1440, freq="min")):
                    active = minute < (60 + day_idx)
                    rows.append(
                        {
                            "timestamp": ts,
                            "flow": 8.0 if active else 0.0,
                            "level": 0.2,
                            "velocity": 0.3,
                        }
                    )
            pd.DataFrame(rows).to_csv(deps.paths.flow_dir / "100_W8.csv", index=False)
            pd.DataFrame(
                {
                    "timestamp": pd.date_range("2026-01-01", periods=4, freq="D"),
                    "rain": [0.0, 0.0, 0.0, 0.0],
                }
            ).to_csv(deps.paths.rainfall_file, index=False)

            data_filter_impl(deps)
            result = analyze_patterns_impl(deps)

            self.assertEqual(result["status"], "ok", result)
            point_ids = {row["point_id"] for row in result["data"]["table"]}
            self.assertIn("W8", point_ids)

    def test_dry_statistics_uses_daily_scale_for_flow(self) -> None:
        flow = pd.DataFrame(
            {
                "timestamp": pd.to_datetime(
                    [
                        "2026-01-01 00:00",
                        "2026-01-01 00:01",
                        "2026-01-02 00:00",
                    ]
                ),
                "device_id": ["100", "100", "100"],
                "point_id": ["W1", "W1", "W1"],
                "flow_lps": [10.0, 20.0, 30.0],
                "level_m": [0.1, 0.3, 0.5],
                "velocity_mps": [0.2, 0.4, 0.6],
            }
        )

        stats = dry_statistics(flow).set_index("point_id")

        self.assertAlmostEqual(stats.loc["W1", "daily_flow_m3d"], 22.5 * 86.4)
        self.assertAlmostEqual(stats.loc["W1", "max_flow_lps"], 30.0)
        self.assertAlmostEqual(stats.loc["W1", "min_flow_lps"], 15.0)

    def test_analyze_patterns_outputs_pipeline_columns(self) -> None:
        timestamps = pd.date_range("2026-01-02", periods=1440, freq="min")
        minutes = timestamps.hour * 60 + timestamps.minute
        flow_values = [
            1.0
            + (2.5 if 7 * 60 <= minute <= 9 * 60 else 0.0)
            + (3.0 if 19 * 60 <= minute <= 21 * 60 else 0.0)
            for minute in minutes
        ]
        flow = pd.DataFrame(
            {
                "timestamp": timestamps,
                "device_id": ["100"] * 1440,
                "point_id": ["W1"] * 1440,
                "flow_lps": flow_values,
                "level_m": [0.2] * 1440,
                "velocity_mps": [0.3] * 1440,
            }
        )

        table = analyze_patterns(flow, smooth_window_minutes=1)["patterns"]

        self.assertEqual(int(table.loc[0, "category"]), 1)
        self.assertIn("category_name", table.columns)
        self.assertIn("peak_count", table.columns)
        self.assertIn("diagnosis_reason", table.columns)
        self.assertIn("description", table.columns)

    def test_pattern_curve_pngs_use_pipeline_overlay_style(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_dir = Path(tmp) / "charts"
            minutes = list(range(1440))
            curves = {
                "W1": pd.DataFrame(
                    {
                        "minute_of_day": minutes,
                        "flow_lps": [50 + minute % 120 for minute in minutes],
                        "level_m": [0.2 + (minute % 60) / 1000 for minute in minutes],
                    }
                )
            }
            rows = []
            for day, offset in [("2026-01-01", 0), ("2026-01-02", 10)]:
                for minute, ts in enumerate(pd.date_range(day, periods=1440, freq="min")):
                    rows.append(
                        {
                            "timestamp": ts,
                            "point_id": "W1",
                            "flow_lps": 40 + offset + minute % 180,
                            "level_m": 0.1 + offset / 100 + (minute % 90) / 1000,
                        }
                    )
            dry_flow = pd.DataFrame(rows)

            saved = _save_pattern_curve_pngs(curves, dry_flow, output_dir)

            flow_path = output_dir / "W1_流量特征曲线.png"
            self.assertIn(str(flow_path), saved["W1"])
            self.assertTrue(flow_path.exists())
            import matplotlib.image as mpimg

            image = mpimg.imread(flow_path)
            self.assertGreater(image.shape[1], image.shape[0] * 1.7)

    def test_rainfall_analysis_writes_charts(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            deps = make_deps(Path(tmp))
            pd.DataFrame(
                {
                    "timestamp": pd.date_range("2026-01-01", periods=4, freq="D"),
                    "rain": [0.0, 2.0, 0.0, 8.0],
                }
            ).to_csv(deps.paths.rainfall_file, index=False)

            result = analyze_rainfall_impl(deps)

            self.assertEqual(result["status"], "ok")
            workbook = load_workbook(deps.paths.combined_xlsx)
            sheet = workbook["降雨概况"]
            self.assertEqual(len(sheet._charts), 2)
            self.assertIn("降雨场次分析", workbook.sheetnames)
            self.assertNotIn("日降雨量统计", workbook.sheetnames)
            self.assertNotIn("场次降雨统计", workbook.sheetnames)
            self.assertTrue((deps.paths.outputs / "降雨分析图" / "日降雨量时间序列图.png").exists())
            self.assertTrue((deps.paths.outputs / "降雨分析图" / "降雨日占比饼图.png").exists())

    def test_rainfall_events_match_pipeline_hourly_windows(self) -> None:
        rain = pd.DataFrame(
            {
                "timestamp": pd.date_range("2026-01-01 00:00", periods=24, freq="h"),
                "rain_mm": [1.0, 2.0, 0.5, 0.0, 3.0, 1.0] + [0.0] * 18,
            }
        )

        events = rainfall_events(rain)
        row = events.iloc[0]

        self.assertEqual(
            list(events.columns),
            [
                "event_id",
                "start_time",
                "end_time",
                "total_rain_mm",
                "duration_h",
                "peak_intensity_mmh",
                "max_3h_rain_mm",
                "max_6h_rain_mm",
                "max_12h_rain_mm",
                "max_24h_rain_mm",
                "avg_intensity_mmh",
                "rain_level",
            ],
        )
        self.assertAlmostEqual(row["total_rain_mm"], 7.5)
        self.assertAlmostEqual(row["duration_h"], 5.0)
        self.assertAlmostEqual(row["peak_intensity_mmh"], 3.0)
        self.assertAlmostEqual(row["max_3h_rain_mm"], 4.0)
        self.assertAlmostEqual(row["max_6h_rain_mm"], 7.5)
        self.assertTrue(pd.isna(row["max_12h_rain_mm"]))
        self.assertTrue(pd.isna(row["max_24h_rain_mm"]))
        self.assertAlmostEqual(row["avg_intensity_mmh"], 1.5)

    def test_rainfall_window_keeps_overlapping_full_event_and_original_id(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            deps = make_deps(Path(tmp))
            pd.DataFrame(
                {
                    "timestamp": pd.to_datetime(
                        [
                            "2026-02-05 19:00",
                            "2026-02-15 15:00",
                            "2026-02-23 09:00",
                            "2026-02-25 19:00",
                            "2026-02-25 22:00",
                            "2026-02-26 02:00",
                            "2026-02-26 08:00",
                        ]
                    ),
                    "rain": [1.9, 3.2, 1.1, 2.1, 2.5, 3.0, 3.0],
                }
            ).to_csv(deps.paths.rainfall_file, index=False)

            result = analyze_rainfall_impl(
                deps,
                time_range=["2026-02-25", "2026-02-26"],
                output="events",
            )

            events = result["data"]["events"]
            self.assertTrue(result["data"]["has_rainfall_coverage"])
            self.assertEqual(len(events), 1)
            self.assertEqual(events[0]["event_id"], 4)
            self.assertEqual(events[0]["start_time"], "2026-02-25 19:00")
            self.assertEqual(events[0]["end_time"], "2026-02-26 08:00")
            self.assertAlmostEqual(events[0]["total_rain_mm"], 10.6)

    def test_event_response_outputs_pipeline_wide_stats(self) -> None:
        flow = pd.DataFrame(
            {
                "timestamp": pd.to_datetime(["2026-01-02 00:30", "2026-01-02 06:00"]),
                "point_id": ["W1", "W1"],
                "flow_lps": [1.0, 3.0],
                "level_m": [0.2, 0.8],
                "velocity_mps": [0.1, 0.2],
            }
        )
        events = pd.DataFrame(
            {
                "event_id": [1],
                "start_time": pd.to_datetime(["2026-01-02 00:00"]),
                "end_time": pd.to_datetime(["2026-01-02 01:00"]),
            }
        )

        table = analyze_event_response(flow, events, [1]).set_index("point_id")

        self.assertAlmostEqual(table.loc["W1", "场次1_最大液位(m)"], 0.8)
        self.assertAlmostEqual(table.loc["W1", "场次1_平均流量(m³/d)"], 172.8)
        self.assertAlmostEqual(table.loc["W1", "场次1_峰值流量(L/s)"], 3.0)

    def test_event_response_skips_events_without_monitoring_coverage(self) -> None:
        flow = pd.DataFrame(
            {
                "timestamp": pd.to_datetime(["2026-01-02 00:30", "2026-01-02 06:00"]),
                "point_id": ["W1", "W1"],
                "flow_lps": [1.0, 3.0],
                "level_m": [0.2, 0.8],
                "velocity_mps": [0.1, 0.2],
            }
        )
        events = pd.DataFrame(
            {
                "event_id": [1, 2],
                "start_time": pd.to_datetime(["2025-12-01 00:00", "2026-01-02 00:00"]),
                "end_time": pd.to_datetime(["2025-12-01 01:00", "2026-01-02 01:00"]),
            }
        )

        table = analyze_event_response(flow, events, [1, 2])

        self.assertNotIn("场次1_最大液位(m)", table.columns)
        self.assertIn("场次2_最大液位(m)", table.columns)
        self.assertIn("场次2_平均流量(m³/d)", table.columns)

    def test_event_response_is_empty_when_no_selected_event_has_coverage(self) -> None:
        flow = pd.DataFrame(
            {
                "timestamp": pd.to_datetime(["2026-03-15 03:30"]),
                "point_id": ["W1"],
                "flow_lps": [1.0],
                "level_m": [0.2],
                "velocity_mps": [0.1],
            }
        )
        events = pd.DataFrame(
            {
                "event_id": [4],
                "start_time": pd.to_datetime(["2026-02-25 19:00"]),
                "end_time": pd.to_datetime(["2026-02-26 08:00"]),
            }
        )

        table = analyze_event_response(flow, events, [4])

        self.assertTrue(table.empty)
        self.assertEqual(list(table.columns), [])

    def test_assess_risk_uses_site_ratios_and_rain_delay(self) -> None:
        dry_stats = pd.DataFrame(
            {
                "point_id": ["W1"],
                "daily_flow_m3d": [100.0],
                "max_level_m": [0.8],
                "avg_velocity_mps": [0.25],
            }
        )
        sites = pd.DataFrame(
            {
                "点位编号": ["W1"],
                "管径(m)": [1.0],
                "井深(m)": [1.0],
                "管道类型": ["分流"],
            }
        )
        flow = pd.DataFrame(
            {
                "timestamp": pd.to_datetime(["2026-01-02 00:30", "2026-01-02 06:00"]),
                "point_id": ["W1", "W1"],
                "level_m": [0.5, 1.1],
            }
        )
        events = pd.DataFrame(
            {
                "event_id": [1],
                "start_time": pd.to_datetime(["2026-01-02 00:00"]),
                "end_time": pd.to_datetime(["2026-01-02 01:00"]),
                "rain_level": ["中雨"],
            }
        )

        result = assess_risk(dry_stats, scope="all", sites=sites, flow=flow, events=events, event_ids=[1])
        dry = result["dry_risk"].set_index("point_id")
        rainy = result["rainy_risk"].set_index("point_id")

        self.assertAlmostEqual(dry.loc["W1", "max_fullness"], 0.8)
        self.assertAlmostEqual(dry.loc["W1", "overflow_value"], 0.8)
        self.assertEqual(dry.loc["W1", "running_risk"], "低风险")
        self.assertEqual(dry.loc["W1", "overflow_risk"], "中溢流风险")
        self.assertEqual(dry.loc["W1", "silting_risk"], "高淤积风险")
        self.assertAlmostEqual(rainy.loc["W1", "max_level_m"], 1.1)
        self.assertEqual(rainy.loc["W1", "overflow_risk"], "已发生溢流")

    def test_analyze_rdii_uses_aligned_dry_curve_total(self) -> None:
        flow = pd.DataFrame(
            {
                "timestamp": pd.date_range("2026-01-02 00:00", periods=3, freq="min"),
                "point_id": ["W1", "W1", "W1"],
                "flow_lps": [3.0, 3.0, 3.0],
                "level_m": [0.1, 0.1, 0.1],
            }
        )
        dry_curves = {
            "W1": pd.DataFrame(
                {
                    "minute_of_day": list(range(1440)),
                    "flow_lps": [1.0] * 1440,
                }
            )
        }
        events = pd.DataFrame(
            {
                "event_id": [1],
                "start_time": pd.to_datetime(["2026-01-02 00:00"]),
                "end_time": pd.to_datetime(["2026-01-02 00:02"]),
            }
        )

        result = analyze_rdii(flow, dry_curves, events, [1])
        table = result["rdii_total"].set_index("point_id")

        self.assertAlmostEqual(table.loc["W1", "场次1"], 0.36)
        self.assertEqual(result["rdii_curve_data"][1]["W1"]["rdii_lps"].tolist(), [2.0, 2.0, 2.0])

    def test_analyze_rdii_skips_events_without_monitoring_coverage(self) -> None:
        flow = pd.DataFrame(
            {
                "timestamp": pd.date_range("2026-01-02 00:00", periods=3, freq="min"),
                "point_id": ["W1", "W1", "W1"],
                "flow_lps": [3.0, 3.0, 3.0],
                "level_m": [0.1, 0.1, 0.1],
            }
        )
        dry_curves = {
            "W1": pd.DataFrame(
                {
                    "minute_of_day": list(range(1440)),
                    "flow_lps": [1.0] * 1440,
                }
            )
        }
        events = pd.DataFrame(
            {
                "event_id": [1, 2],
                "start_time": pd.to_datetime(["2025-12-01 00:00", "2026-01-02 00:00"]),
                "end_time": pd.to_datetime(["2025-12-01 00:02", "2026-01-02 00:02"]),
            }
        )

        result = analyze_rdii(flow, dry_curves, events, [1, 2])
        table = result["rdii_total"].set_index("point_id")

        self.assertNotIn("场次1", table.columns)
        self.assertIn("场次2", table.columns)
        self.assertNotIn(1, result["rdii_curve_data"])
        self.assertIn(2, result["rdii_curve_data"])

    def test_analyze_rdii_is_empty_when_no_selected_event_has_coverage(self) -> None:
        flow = pd.DataFrame(
            {
                "timestamp": pd.date_range("2026-03-15 03:00", periods=3, freq="min"),
                "point_id": ["W1", "W1", "W1"],
                "flow_lps": [3.0, 3.0, 3.0],
                "level_m": [0.1, 0.1, 0.1],
            }
        )
        dry_curves = {
            "W1": pd.DataFrame(
                {"minute_of_day": list(range(1440)), "flow_lps": [1.0] * 1440}
            )
        }
        events = pd.DataFrame(
            {
                "event_id": [4],
                "start_time": pd.to_datetime(["2026-02-25 19:00"]),
                "end_time": pd.to_datetime(["2026-02-26 08:00"]),
            }
        )

        result = analyze_rdii(flow, dry_curves, events, [4])

        self.assertTrue(result["rdii_total"].empty)
        self.assertEqual(result["rdii_curve_data"], {})

    def test_save_rdii_curve_pngs_outputs_pipeline_paths(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_dir = Path(tmp) / "charts"
            rain_times = pd.date_range("2026-01-02 00:00", periods=720, freq="min").tolist()
            rain = pd.DataFrame(
                {
                    "timestamp": rain_times,
                    "rain_mm": [0.2 if idx % 60 == 0 else 0.0 for idx in range(720)],
                }
            )
            events = pd.DataFrame(
                {
                    "event_id": [1],
                    "start_time": [pd.Timestamp("2026-01-02 00:00")],
                    "end_time": [pd.Timestamp("2026-01-02 00:05")],
                }
            )
            rdii_curve_data = {
                1: {
                    "W1": pd.DataFrame(
                        {
                            "rain_flow_lps": [3.0, 4.0, 5.0],
                            "dry_flow_lps": [1.0, 1.0, 1.0],
                            "rdii_lps": [2.0, 3.0, 4.0],
                        },
                        index=pd.date_range("2026-01-02 00:00", periods=3, freq="min"),
                    )
                }
            }

            saved = _save_rdii_curve_pngs(rdii_curve_data, rain, events, output_dir, selected_events=[1])

            expected = output_dir / "rdii_curve" / "event1_1_2" / "W1_event1.png"
            self.assertEqual(saved[1]["W1"], str(expected))
            self.assertTrue(expected.exists())

    def test_save_rdii_curve_pngs_does_not_use_pandas_bar_plot(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_dir = Path(tmp) / "charts"
            rain = pd.DataFrame(
                {
                    "timestamp": pd.date_range("2026-01-02 00:00", periods=720, freq="min").tolist(),
                    "rain_mm": [0.1] * 720,
                }
            )
            events = pd.DataFrame(
                {
                    "event_id": [1],
                    "start_time": [pd.Timestamp("2026-01-02 00:00")],
                    "end_time": [pd.Timestamp("2026-01-02 00:05")],
                }
            )
            rdii_curve_data = {
                1: {
                    "W1": pd.DataFrame(
                        {
                            "rain_flow_lps": [3.0, 4.0, 5.0],
                            "dry_flow_lps": [1.0, 1.0, 1.0],
                            "rdii_lps": [2.0, 3.0, 4.0],
                        },
                        index=pd.date_range("2026-01-02 00:00", periods=3, freq="min"),
                    )
                }
            }

            original_plot = pd.Series.plot

            def fail_bar_plot(self, *args, **kwargs):
                if kwargs.get("kind") == "bar":
                    raise ValueError("Must supply freq for datetime value")
                return original_plot(self, *args, **kwargs)

            with patch.object(pd.Series, "plot", fail_bar_plot):
                saved = _save_rdii_curve_pngs(rdii_curve_data, rain, events, output_dir, selected_events=[1])

            self.assertTrue((output_dir / "rdii_curve" / "event1_1_2" / "W1_event1.png").exists())
            self.assertIn("W1", saved[1])

    def test_analyze_rdii_impl_writes_curve_pngs(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            deps = make_deps(Path(tmp))
            write_rdii_sample(deps)

            result = analyze_rdii_impl(deps, event_ids=[1])

            expected = deps.paths.outputs / "rdii_curve" / "event1_1_3" / "W1_event1.png"
            self.assertEqual(result["status"], "ok", result)
            self.assertTrue(expected.exists())
            self.assertIn(str(expected), result["data"]["chart_paths"][1].values())

    def test_partial_rdii_without_export_writes_no_csv_or_png(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            deps = make_deps(Path(tmp))
            write_rdii_sample(deps)

            result = analyze_rdii_impl(deps, event_ids=[1], points=["W1"], export=False)

            self.assertEqual(result["status"], "ok", result)
            self.assertEqual(result["data"]["result_destinations"][0]["kind"], "not_persisted")
            self.assertFalse(list(deps.paths.outputs.glob("*.csv")))
            self.assertFalse(list(deps.paths.outputs.rglob("*.png")))
            workbook = load_workbook(deps.paths.combined_xlsx)
            self.assertNotIn("RDII总量统计", workbook.sheetnames)

    def test_partial_rdii_with_export_writes_named_csv_and_png(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            deps = make_deps(Path(tmp))
            write_rdii_sample(deps)

            result = analyze_rdii_impl(deps, event_ids=[1], points=["W1"], export=True)

            self.assertEqual(result["status"], "ok", result)
            self.assertTrue((deps.paths.outputs / "W1_RDII总量统计.csv").exists())
            self.assertTrue((deps.paths.outputs / "W1_RDII曲线.png").exists())
            self.assertFalse((deps.paths.outputs / "rdii_curve" / "event1_1_3" / "W1_event1.png").exists())
            workbook = load_workbook(deps.paths.combined_xlsx)
            self.assertNotIn("RDII总量统计", workbook.sheetnames)

    def test_run_python_success_and_workspace_artifact(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            deps = make_deps(Path(tmp))
            write_flow(deps)
            result = run_python_impl(deps, "(WORKSPACE_DIR / 'out.txt').write_text(str(len(load_flow())), encoding='utf-8')")
            self.assertEqual(result["status"], "ok", result)
            self.assertEqual((deps.paths.workspace / "out.txt").read_text(encoding="utf-8"), "5")

    def test_build_report_uses_template_modules_and_workbook_tables(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            outputs = root / "outputs"
            templates = root / "templates"
            outputs.mkdir()
            templates.mkdir()
            combined = outputs / "综合分析结果.xlsx"
            with pd.ExcelWriter(combined, engine="openpyxl") as writer:
                pd.DataFrame({"日期": ["2026-01-01", "2026-01-02"], "日降雨量(mm)": [0.0, 8.0]}).to_excel(
                    writer, sheet_name="降雨概况", index=False
                )
                pd.DataFrame({"场次编号": [1], "总降雨量(mm)": [8.0]}).to_excel(writer, sheet_name="降雨场次分析", index=False)
                pd.DataFrame({"点位编号": ["W1"], "运行风险": ["低风险"]}).to_excel(writer, sheet_name="旱天风险", index=False)

            template = templates / "template.docx"
            doc = Document()
            doc.add_heading("监测数据分析报告", level=0)
            doc.add_heading("降雨分析", level=1)
            doc.add_paragraph("这里是模板中的降雨章节。")
            doc.save(template)

            output = outputs / "分析报告.docx"
            result = build_report(
                output,
                "排水监测数据分析报告",
                template_file=template,
                combined_xlsx=combined,
                outputs_dir=outputs,
                sections=["降雨分析", "污水系统运行风险分析"],
            )

            rendered = Document(output)
            paragraph_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
            self.assertNotIn("自动分析摘要", paragraph_text)
            self.assertIn("降雨分析共统计", paragraph_text)
            self.assertIn("污水系统运行风险分析", paragraph_text)
            self.assertEqual(result["templated_sections"], ["降雨分析"])
            self.assertEqual(result["generated_sections"], ["污水系统运行风险分析"])
            self.assertGreaterEqual(result["stats"]["inserted_tables"], 3)

    def test_build_report_fills_existing_collection_rate_template_table(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            outputs = root / "outputs"
            templates = root / "templates"
            outputs.mkdir()
            templates.mkdir()
            combined = outputs / "综合分析结果.xlsx"
            with pd.ExcelWriter(combined, engine="openpyxl") as writer:
                pd.DataFrame(
                    {
                        "点位编号": ["W1"],
                        "记录数": [1440],
                        "监测天数": [1],
                        "理论数据条数": [1440],
                        "开始时间": ["2026-01-01 00:00"],
                        "结束时间": ["2026-01-01 23:59"],
                        "收集率": [1.0],
                    }
                ).to_excel(writer, sheet_name="数据体检", index=False)

            template = templates / "template.docx"
            doc = Document()
            doc.add_heading("监测概况", level=1)
            doc.add_paragraph("本轮共布设44个流量监测点位，时间段选择2024/9/18日-2024/11/26日。")
            table = doc.add_table(rows=2, cols=5)
            for idx, header in enumerate(["点位编号", "监测数据条数", "监测天数", "理论数据条数", "数据收集率"]):
                table.rows[0].cells[idx].text = header
                table.rows[1].cells[idx].text = "模板旧值"
            doc.save(template)

            output = outputs / "分析报告.docx"
            result = build_report(
                output,
                "排水监测数据分析报告",
                template_file=template,
                combined_xlsx=combined,
                sections=["监测概况"],
            )

            rendered = Document(output)
            filled = rendered.tables[0].rows[1].cells
            self.assertEqual(filled[0].text, "W1")
            self.assertEqual(filled[1].text, "1440")
            self.assertEqual(filled[3].text, "1440")
            self.assertEqual(filled[4].text, "100.0%")
            self.assertEqual(result["stats"]["filled_template_tables"], 1)
            self.assertEqual(result["stats"]["text_replaced"], 1)
            self.assertEqual(len(rendered.tables), 1)
            text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
            self.assertNotIn("2024/9/18", text)
            self.assertIn("本轮共布设1个流量监测点位", text)

    def test_build_report_does_not_warn_site_info_missing_when_filled_from_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            outputs = root / "outputs"
            templates = root / "templates"
            data = root / "data"
            outputs.mkdir()
            templates.mkdir()
            data.mkdir()
            combined = outputs / "综合分析结果.xlsx"
            site_info = data / "点位信息.xlsx"
            with pd.ExcelWriter(combined, engine="openpyxl") as writer:
                pd.DataFrame(
                    {
                        "点位编号": ["W1"],
                        "记录数": [1440],
                        "监测天数": [1],
                        "理论数据条数": [1440],
                        "收集率": [1.0],
                    }
                ).to_excel(writer, sheet_name="数据体检", index=False)
            pd.DataFrame(
                {
                    "安装监测点位": ["W1"],
                    "形状": ["圆管"],
                    "管径(m)": [1.5],
                    "井深(m)": [5.4],
                    "设备安装时间": ["2026-01-22"],
                    "类型": ["污水"],
                }
            ).to_excel(site_info, index=False)

            template = templates / "template.docx"
            doc = Document()
            doc.add_heading("监测概况", level=1)
            table = doc.add_table(rows=2, cols=6)
            for idx, header in enumerate(["监测点位", "设备类型", "形状", "管径(m)", "井深(m)", "设备安装时间"]):
                table.rows[0].cells[idx].text = header
            doc.save(template)

            output = outputs / "分析报告.docx"
            result = build_report(
                output,
                "排水监测数据分析报告",
                template_file=template,
                combined_xlsx=combined,
                site_info_file=site_info,
                sections=["监测概况"],
            )

            rendered = Document(output)
            self.assertEqual(rendered.tables[0].rows[1].cells[0].text, "W1")
            self.assertNotIn("监测点位安装信息汇总", result["missing_sheets"])

    def test_build_report_rebuilds_pattern_section_from_template_boundary(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            outputs = root / "outputs"
            templates = root / "templates"
            image_dir = outputs / "特征曲线图"
            outputs.mkdir()
            templates.mkdir()
            image_dir.mkdir()
            combined = outputs / "综合分析结果.xlsx"
            with pd.ExcelWriter(combined, engine="openpyxl") as writer:
                pd.DataFrame(
                    {
                        "点位编号": ["W1", "W2"],
                        "分类": [1, 3],
                        "分类名称": ["第1类 符合生活用水规律", "第3类 曲线平坦/异常"],
                        "排污规律描述": ["W1点位符合生活用水规律。", "W2点位曲线无明显波峰或波谷。"],
                    }
                ).to_excel(writer, sheet_name="排污规律分析", index=False)

            png_bytes = base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
            )
            (image_dir / "W1_流量特征曲线.png").write_bytes(png_bytes)

            template = templates / "template.docx"
            doc = Document()
            doc.add_heading("旱天排污规律统计分析", level=1)
            doc.add_paragraph("模板示例点位内容，需要删除。")
            doc.add_heading("污水系统运行风险分析", level=1)
            doc.save(template)

            output = outputs / "分析报告.docx"
            result = build_report(
                output,
                "排水监测数据分析报告",
                template_file=template,
                combined_xlsx=combined,
                outputs_dir=outputs,
                sections=["旱天排污规律统计分析"],
            )

            rendered = Document(output)
            text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
            self.assertNotIn("模板示例点位内容", text)
            self.assertIn("本轮监测的2个点位", text)
            self.assertIn("W1点位符合生活用水规律", text)
            self.assertIn("W2点位曲线无明显波峰或波谷", text)
            self.assertGreaterEqual(result["stats"]["text_replaced"], 5)
            self.assertEqual(result["stats"]["inserted_images"], 1)

    def test_build_report_writes_grouped_risk_section(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            outputs = root / "outputs"
            templates = root / "templates"
            outputs.mkdir()
            templates.mkdir()
            combined = outputs / "综合分析结果.xlsx"
            with pd.ExcelWriter(combined, engine="openpyxl") as writer:
                pd.DataFrame(
                    {
                        "点位编号": ["W1", "W2", "W3"],
                        "最大充满度": [0.5, 0.85, 2.2],
                        "溢流风险值": [0.4, 0.75, 1.2],
                        "旱天流速(m/s)": [0.8, 0.45, 0.2],
                    }
                ).to_excel(writer, sheet_name="旱天风险", index=False)

            template = templates / "template.docx"
            doc = Document()
            doc.add_heading("污水系统运行风险分析", level=1)
            doc.add_paragraph("旧风险章节示例内容。")
            doc.save(template)

            output = outputs / "分析报告.docx"
            result = build_report(
                output,
                "排水监测数据分析报告",
                template_file=template,
                combined_xlsx=combined,
                sections=["污水系统运行风险分析"],
            )

            rendered = Document(output)
            text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
            self.assertIn("旱天最大充满度情况如下", text)
            self.assertIn("溢流风险值情况如下", text)
            self.assertIn("淤积风险情况如下", text)
            self.assertIn("本章小结", text)
            self.assertIn("W3", text)
            self.assertGreaterEqual(result["stats"]["text_replaced"], 10)

    def test_build_report_inserts_rainfall_images_before_captions(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            outputs = root / "outputs"
            templates = root / "templates"
            chart_dir = outputs / "降雨分析图"
            outputs.mkdir()
            templates.mkdir()
            chart_dir.mkdir()
            combined = outputs / "综合分析结果.xlsx"
            with pd.ExcelWriter(combined, engine="openpyxl") as writer:
                pd.DataFrame({"日期": ["2026-01-01", "2026-01-02"], "日降雨量(mm)": [0.0, 8.0]}).to_excel(
                    writer, sheet_name="降雨概况", index=False
                )
                pd.DataFrame({"场次编号": [1], "总降雨量(mm)": [8.0]}).to_excel(writer, sheet_name="降雨场次分析", index=False)

            png_bytes = base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
            )
            (chart_dir / "日降雨量时间序列图.png").write_bytes(png_bytes)
            (chart_dir / "降雨日占比饼图.png").write_bytes(png_bytes)

            template = templates / "template.docx"
            doc = Document()
            doc.add_heading("降雨分析", level=1)
            doc.add_paragraph("图 15 日降雨量时间序列图")
            doc.add_paragraph("图 16 降雨日与非降雨日占比图")
            doc.save(template)

            output = outputs / "分析报告.docx"
            result = build_report(
                output,
                "排水监测数据分析报告",
                template_file=template,
                combined_xlsx=combined,
                outputs_dir=outputs,
                sections=["降雨分析"],
            )

            rendered = Document(output)
            text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
            self.assertIn("图 15 日降雨量时间序列图", text)
            self.assertIn("图 16 降雨日与非降雨日占比图", text)
            self.assertEqual(result["stats"]["inserted_images"], 2)
            self.assertEqual(len(rendered.inline_shapes), 2)


if __name__ == "__main__":
    unittest.main()
