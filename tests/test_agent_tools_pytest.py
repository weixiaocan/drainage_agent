from __future__ import annotations

import base64
import logging
import time
import zipfile
from pathlib import Path
from typing import get_args

import pandas as pd
import pytest
from docx import Document

from analysis import io
from analysis.reporting import build_report
from agent.deps import AgentDeps, AgentSettings, Paths, SessionState, ensure_directories
from agent.tools.inspect_tools import list_results_impl
from agent.tools.manifest import record_result
from agent.tools.memory_tool import record_note_impl
from agent.tools.module_tools import (
    analyze_event_response_impl,
    analyze_patterns_impl,
    analyze_rainfall_impl,
    analyze_rdii_impl,
    assess_risk_impl,
    check_data_impl,
    data_filter_impl,
    generate_report_impl,
    is_full_network,
    _report_actual_time_range,
    _time_result_prefix,
)
from analysis.pipeline_report_assembler.assembler import _scope_period_text
from analysis.pipeline_report_assembler.template_scanner import scan_template
from agent.tools.python_tool import run_python_impl
from agent.types import ToolStatus, ok


def make_deps(root: Path) -> AgentDeps:
    paths = Paths.from_root(root)
    ensure_directories(paths)
    session = SessionState()
    session.auto_confirm_filter_result = True
    return AgentDeps(
        paths=paths,
        settings=AgentSettings(model="test", base_url=None, api_key=None),
        logger=logging.getLogger("test.agent_tools"),
        session=session,
        project_notes="",
    )


def write_sample_data(deps: AgentDeps) -> None:
    deps.paths.flow_dir.mkdir(parents=True, exist_ok=True)
    rows = 60
    pd.DataFrame(
        {
            "数据时间": pd.date_range("2026-01-01", periods=rows, freq="min"),
            "设备编号": ["100"] * rows,
            "流量(L/s)(均值)": [1.0 + i / 100 for i in range(rows)],
            "液位(m)(均值)": [0.2 + i / 1000 for i in range(rows)],
            "流速(m/s)(均值)": [0.3 + i / 1000 for i in range(rows)],
        }
    ).to_csv(deps.paths.flow_dir / "100_W1.csv", index=False, encoding="utf-8-sig")
    pd.DataFrame(
        {
            "timestamp": pd.date_range("2026-01-01", periods=6, freq="h"),
            "rain": [1, 2, 0, 0, 0, 0],
        }
    ).to_csv(deps.paths.rainfall_file, index=False)
    pd.DataFrame({"点位编号": ["W1"], "管径": [1.0]}).to_excel(deps.paths.site_info_file, index=False)


def write_filter_sample_data(deps: AgentDeps) -> None:
    deps.paths.flow_dir.mkdir(parents=True, exist_ok=True)
    rows = 4 * 1440
    pd.DataFrame(
        {
            "数据时间": pd.date_range("2026-01-01", periods=rows, freq="min"),
            "设备编号": ["100"] * rows,
            "流量(L/s)(均值)": [1.0] * rows,
            "液位(m)(均值)": [0.2] * rows,
            "流速(m/s)(均值)": [0.3] * rows,
        }
    ).to_csv(deps.paths.flow_dir / "100_W1.csv", index=False, encoding="utf-8-sig")
    pd.DataFrame(
        {
            "timestamp": pd.date_range("2026-01-01", periods=4, freq="D"),
            "rain": [0.0, 0.0, 3.0, 0.0],
        }
    ).to_csv(deps.paths.rainfall_file, index=False)
    pd.DataFrame({"点位编号": ["W1"], "管径": [1.0]}).to_excel(deps.paths.site_info_file, index=False)


def write_two_point_data(deps: AgentDeps) -> None:
    write_sample_data(deps)
    rows = 60
    pd.DataFrame(
        {
            "数据时间": pd.date_range("2026-01-01", periods=rows, freq="min"),
            "设备编号": ["200"] * rows,
            "流量(L/s)(均值)": [2.0 + i / 100 for i in range(rows)],
            "液位(m)(均值)": [0.4 + i / 1000 for i in range(rows)],
            "流速(m/s)(均值)": [0.5 + i / 1000 for i in range(rows)],
        }
    ).to_csv(deps.paths.flow_dir / "200_W2.csv", index=False, encoding="utf-8-sig")
    pd.DataFrame({"点位编号": ["W1", "W2"], "管径": [1.0, 1.2]}).to_excel(
        deps.paths.site_info_file,
        index=False,
    )


def sample_two_point_pattern_flow() -> pd.DataFrame:
    frames = []
    for point_id, offset in (("W1", 0.0), ("W2", 1.0)):
        for day in ("2026-01-01", "2026-01-02"):
            timestamps = pd.date_range(day, periods=120, freq="min")
            frames.append(
                pd.DataFrame(
                    {
                        "timestamp": timestamps,
                        "device_id": point_id,
                        "point_id": point_id,
                        "flow_lps": [offset + 1.0 + i / 100 for i in range(120)],
                        "level_m": [offset / 10 + 0.2 + i / 1000 for i in range(120)],
                        "velocity_mps": [0.3] * 120,
                    }
                )
            )
    return pd.concat(frames, ignore_index=True)


def test_tool_status_values_are_v2_only() -> None:
    assert set(get_args(ToolStatus)) == {"ok", "needs_input", "needs_confirmation", "error"}


def test_check_data_success(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_sample_data(deps)

    check = check_data_impl(deps)

    assert check["status"] == "ok"
    assert not deps.paths.combined_xlsx.exists()


def test_full_network_standalone_analysis_does_not_write_combined_sheet(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_two_point_data(deps)

    result = check_data_impl(deps, points=["W1", "W2"])

    assert result["status"] == "ok"
    assert result["data"]["result_destinations"] == [
        {"kind": "not_persisted", "path": None, "sheet": None}
    ]
    assert not deps.paths.combined_xlsx.exists()
    assert "综合分析结果.xlsx" not in result["summary"]


def test_partial_analysis_without_export_does_not_persist_table(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_two_point_data(deps)

    result = check_data_impl(deps, points=["W1"], export=False)

    assert result["status"] == "ok"
    assert result["data"]["result_destinations"] == [
        {"kind": "not_persisted", "path": None, "sheet": None}
    ]
    assert not deps.paths.combined_xlsx.exists()
    assert not list(deps.paths.outputs.glob("*.csv"))


def test_partial_analysis_with_export_writes_csv_not_combined(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_two_point_data(deps)

    result = check_data_impl(deps, points=["W1"], export=True)

    destination = result["data"]["result_destinations"][0]
    assert result["status"] == "ok"
    assert destination == {
        "kind": "csv",
        "path": "outputs/W1_全时段_数据收集率统计.csv",
        "sheet": None,
    }
    assert (deps.paths.outputs / "W1_全时段_数据收集率统计.csv").exists()
    assert not deps.paths.combined_xlsx.exists()
    assert "已导出 CSV" in result["summary"]
    assert "综合分析结果.xlsx" not in result["summary"]
    assert result["artifacts"] == ["outputs/W1_全时段_数据收集率统计.csv"]


def test_partial_analysis_does_not_replace_existing_full_network_sheet(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_two_point_data(deps)
    pd.DataFrame([{"point_id": "W1", "record_count": 1}]).to_excel(
        deps.paths.combined_xlsx, sheet_name="数据收集率统计", index=False
    )
    before = pd.read_excel(deps.paths.combined_xlsx, sheet_name="数据收集率统计")

    result = check_data_impl(deps, points=["W1"], export=False)
    after = pd.read_excel(deps.paths.combined_xlsx, sheet_name="数据收集率统计")

    assert result["data"]["result_destinations"][0]["kind"] == "not_persisted"
    pd.testing.assert_frame_equal(after, before)


def test_full_network_partial_time_does_not_write_combined_and_exports_range_csv(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_two_point_data(deps)
    start = "2026-01-01 00:10:00"
    end = "2026-01-01 00:19:00"

    without_export = check_data_impl(deps, start=start, end=end)
    with_export = check_data_impl(deps, start=start, end=end, export=True)

    assert without_export["data"]["result_destinations"][0]["kind"] == "not_persisted"
    assert with_export["data"]["result_destinations"][0] == {
        "kind": "csv",
            "path": "outputs/全网_2026-01-01_00-10-00_2026-01-01_00-19-00_数据收集率统计.csv",
        "sheet": None,
    }
    assert (deps.paths.outputs / "全网_2026-01-01_00-10-00_2026-01-01_00-19-00_数据收集率统计.csv").exists()
    assert not deps.paths.combined_xlsx.exists()


def test_partial_points_partial_time_does_not_write_combined(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_two_point_data(deps)

    result = check_data_impl(
        deps,
        points=["W1"],
        start="2026-01-01 00:10:00",
        end="2026-01-01 00:19:00",
        export=True,
    )

    assert result["data"]["result_destinations"][0]["kind"] == "csv"
    assert (deps.paths.outputs / "W1_2026-01-01_00-10-00_2026-01-01_00-19-00_数据收集率统计.csv").exists()
    assert not deps.paths.combined_xlsx.exists()


def test_full_network_partial_time_does_not_replace_full_time_sheet(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_two_point_data(deps)
    pd.DataFrame([{"point_id": "W1", "record_count": 1}]).to_excel(
        deps.paths.combined_xlsx, sheet_name="数据收集率统计", index=False
    )
    before = pd.read_excel(deps.paths.combined_xlsx, sheet_name="数据收集率统计")

    result = check_data_impl(
        deps,
        start="2026-01-01 00:10:00",
        end="2026-01-01 00:19:00",
    )
    after = pd.read_excel(deps.paths.combined_xlsx, sheet_name="数据收集率统计")

    assert result["data"]["result_destinations"][0]["kind"] == "not_persisted"
    pd.testing.assert_frame_equal(after, before)


def test_full_network_patterns_write_no_combined_and_full_network_pngs(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    deps = make_deps(tmp_path)
    write_two_point_data(deps)
    flow = sample_two_point_pattern_flow()
    monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_args, **_kwargs: flow)

    result = analyze_patterns_impl(deps)

    assert result["status"] == "ok"
    assert result["data"]["result_destinations"][0]["kind"] == "not_persisted"
    assert not deps.paths.combined_xlsx.exists()
    assert (deps.paths.outputs / "特征曲线图" / "全网_全时段" / "W1_流量特征曲线.png").exists()
    assert (deps.paths.outputs / "特征曲线图" / "全网_全时段" / "W2_流量特征曲线.png").exists()


def test_partial_patterns_without_export_write_no_table_or_png(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    deps = make_deps(tmp_path)
    write_two_point_data(deps)
    flow = sample_two_point_pattern_flow()
    monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_args, **_kwargs: flow[flow["point_id"] == "W1"])

    result = analyze_patterns_impl(deps, points=["W1"], export=False)

    assert result["status"] == "ok"
    assert result["data"]["result_destinations"][0]["kind"] == "not_persisted"
    assert not deps.paths.combined_xlsx.exists()
    assert not list(deps.paths.outputs.glob("*.csv"))
    assert not list(deps.paths.outputs.rglob("*.png"))


def test_partial_patterns_with_export_write_named_csv_and_png_only(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    deps = make_deps(tmp_path)
    write_two_point_data(deps)
    flow = sample_two_point_pattern_flow()
    monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_args, **_kwargs: flow[flow["point_id"] == "W1"])

    result = analyze_patterns_impl(deps, points=["W1"], export=True)

    assert result["status"] == "ok"
    assert (deps.paths.outputs / "W1_全时段_排污规律分析.csv").exists()
    point_images = [Path(path) for path in result["data"]["curve_images"]["W1"]]
    assert len(point_images) == 2
    assert {path.name for path in point_images} == {"W1_流量特征曲线.png", "W1_液位特征曲线.png"}
    assert all(path.exists() for path in point_images)
    assert len({path.read_bytes() for path in point_images}) == 2
    assert not deps.paths.combined_xlsx.exists()
    assert not (deps.paths.outputs / "特征曲线图" / "全网_全时段" / "W1_流量特征曲线.png").exists()


def test_partial_patterns_do_not_overwrite_full_network_sheet_or_fixed_png(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    deps = make_deps(tmp_path)
    write_two_point_data(deps)
    flow = sample_two_point_pattern_flow()
    monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_args, **_kwargs: flow)
    analyze_patterns_impl(deps)
    pd.DataFrame([{"point_id": "W2", "category": 1}]).to_excel(
        deps.paths.combined_xlsx, sheet_name="排污规律分析", index=False
    )
    sheet_before = pd.read_excel(deps.paths.combined_xlsx, sheet_name="排污规律分析")
    fixed_png = deps.paths.outputs / "特征曲线图" / "全网_全时段" / "W1_流量特征曲线.png"
    png_before = fixed_png.read_bytes()

    monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_args, **_kwargs: flow[flow["point_id"] == "W1"])
    analyze_patterns_impl(deps, points=["W1"], export=True)
    sheet_after = pd.read_excel(deps.paths.combined_xlsx, sheet_name="排污规律分析")

    pd.testing.assert_frame_equal(sheet_after, sheet_before)
    assert fixed_png.read_bytes() == png_before
    assert (deps.paths.outputs / "特征曲线图" / "W1_全时段" / "W1_流量特征曲线.png").exists()
    assert (deps.paths.outputs / "特征曲线图" / "W1_全时段" / "W1_液位特征曲线.png").exists()


def test_rainfall_time_window_pngs_are_range_named(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_sample_data(deps)

    full = analyze_rainfall_impl(deps)
    window = analyze_rainfall_impl(deps, time_range=["2026-01-01 00:00:00", "2026-01-01 01:00:00"])

    full_daily = Path(full["data"]["chart_paths"]["daily_bar"])
    window_daily = Path(window["data"]["chart_paths"]["daily_bar"])
    assert full_daily.exists()
    assert window_daily.exists()
    assert full_daily != window_daily
    assert "全网_全时段" in full_daily.name
    assert "全网_2026-01-01_00-00-00_2026-01-01_01-00-00" in window_daily.name


def test_full_network_pattern_window_pngs_use_separate_range_dir(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    deps = make_deps(tmp_path)
    write_two_point_data(deps)
    flow = sample_two_point_pattern_flow()
    monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_args, **_kwargs: flow)

    full = analyze_patterns_impl(deps)
    window = analyze_patterns_impl(deps, start="2026-01-01 00:00:00", end="2026-01-01 00:59:00")

    full_png = Path(full["data"]["curve_images"]["W1"][0])
    window_png = Path(window["data"]["curve_images"]["W1"][0])
    assert full_png.exists()
    assert window_png.exists()
    assert full_png != window_png
    assert full_png.parent.name == "全网_全时段"
    assert window_png.parent.name == "全网_2026-01-01_00-00-00_2026-01-01_00-59-00"


def test_partial_pattern_export_window_png_includes_point_and_time_range(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    deps = make_deps(tmp_path)
    write_two_point_data(deps)
    flow = sample_two_point_pattern_flow()
    monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_args, **_kwargs: flow[flow["point_id"] == "W1"])

    result = analyze_patterns_impl(
        deps,
        points=["W1"],
        export=True,
        start="2026-01-01 00:00:00",
        end="2026-01-01 00:59:00",
    )

    assert result["status"] == "ok"
    point_images = [Path(path) for path in result["data"]["curve_images"]["W1"]]
    assert len(point_images) == 2
    assert all(path.exists() for path in point_images)
    assert all(path.parent.name == "W1_2026-01-01_00-00-00_2026-01-01_00-59-00" for path in point_images)
    assert not (deps.paths.outputs / "W1_2026-01-01_00-00-00_2026-01-01_00-59-00_排污规律曲线.png").exists()


def test_patterns_time_window_slices_before_analysis(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    deps = make_deps(tmp_path)
    flow = sample_two_point_pattern_flow()
    captured: dict[str, pd.DataFrame] = {}
    monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_args, **_kwargs: flow)

    def capture_analysis(window_flow: pd.DataFrame, **_kwargs: object) -> dict[str, object]:
        captured["flow"] = window_flow.copy()
        return {"patterns": pd.DataFrame([{"point_id": "W1"}]), "curves": {}, "descriptions": {}}

    monkeypatch.setattr("agent.tools.module_tools.analyze_patterns", capture_analysis)

    result = analyze_patterns_impl(
        deps,
        points=["W1"],
        start="2026-01-02 00:10",
        end="2026-01-02 00:19",
    )

    assert result["status"] == "ok"
    assert len(captured["flow"]) == 10
    assert captured["flow"]["timestamp"].min() == pd.Timestamp("2026-01-02 00:10")
    assert captured["flow"]["timestamp"].max() == pd.Timestamp("2026-01-02 00:19")


def test_patterns_none_window_preserves_full_input(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    deps = make_deps(tmp_path)
    flow = sample_two_point_pattern_flow()
    captured: list[pd.DataFrame] = []
    monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_args, **_kwargs: flow)

    def capture_analysis(input_flow: pd.DataFrame, **_kwargs: object) -> dict[str, object]:
        captured.append(input_flow.copy())
        return {"patterns": pd.DataFrame([{"point_id": "W1"}]), "curves": {}, "descriptions": {}}

    monkeypatch.setattr("agent.tools.module_tools.analyze_patterns", capture_analysis)

    before = analyze_patterns_impl(deps, points=["W1"])
    after = analyze_patterns_impl(deps, points=["W1"], start=None, end=None)

    assert before["status"] == after["status"] == "ok"
    assert before["data"] == after["data"]
    assert "window_coverage" not in after["data"]
    assert len(captured) == 1
    pd.testing.assert_frame_equal(captured[0], flow)


def test_patterns_time_window_rejects_no_coverage(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    deps = make_deps(tmp_path)
    flow = sample_two_point_pattern_flow()
    monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_args, **_kwargs: flow)

    result = analyze_patterns_impl(
        deps,
        points=["W1"],
        start="2026-02-01",
        end="2026-02-02",
    )

    assert result["status"] == "needs_input"
    assert result["missing"] == "data_coverage"
    assert "无数据覆盖" in result["summary"]


def test_patterns_time_window_uses_partial_point_coverage_and_reports_range(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    deps = make_deps(tmp_path)
    flow = sample_two_point_pattern_flow()
    flow.loc[flow["point_id"] == "W2", "timestamp"] += pd.Timedelta(days=10)
    captured: dict[str, pd.DataFrame] = {}
    monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_args, **_kwargs: flow)

    def capture_analysis(window_flow: pd.DataFrame, **_kwargs: object) -> dict[str, object]:
        captured["flow"] = window_flow.copy()
        return {"patterns": pd.DataFrame([{"point_id": "W1"}]), "curves": {}, "descriptions": {}}

    monkeypatch.setattr("agent.tools.module_tools.analyze_patterns", capture_analysis)

    result = analyze_patterns_impl(
        deps,
        points=["W1", "W2"],
        start="2026-01-02",
        end="2026-01-02 00:09",
    )

    assert result["status"] == "ok"
    assert captured["flow"]["point_id"].unique().tolist() == ["W1"]
    assert len(captured["flow"]) == 10
    assert result["data"]["excluded_points"][0]["point_id"] == "W2"
    assert result["data"]["window_coverage"]["actual_start"] == "2026-01-02 00:00:00"
    assert result["data"]["window_coverage"]["actual_end"] == "2026-01-02 00:09:00"
    assert "实际分析范围" in result["summary"]


def test_patterns_cross_month_date_window_keeps_both_endpoints(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    deps = make_deps(tmp_path)
    flow = pd.DataFrame(
        {
            "timestamp": pd.to_datetime(
                ["2026-02-27 23:59", "2026-02-28 00:00", "2026-03-01 23:59", "2026-03-02 00:00"]
            ),
            "device_id": ["W1"] * 4,
            "point_id": ["W1"] * 4,
            "flow_lps": [1.0, 2.0, 3.0, 4.0],
            "level_m": [0.1] * 4,
            "velocity_mps": [0.2] * 4,
        }
    )
    captured: dict[str, pd.DataFrame] = {}
    monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_args, **_kwargs: flow)

    def capture_analysis(window_flow: pd.DataFrame, **_kwargs: object) -> dict[str, object]:
        captured["flow"] = window_flow.copy()
        return {"patterns": pd.DataFrame([{"point_id": "W1"}]), "curves": {}, "descriptions": {}}

    monkeypatch.setattr("agent.tools.module_tools.analyze_patterns", capture_analysis)

    result = analyze_patterns_impl(deps, points=["W1"], start="2026-02-28", end="2026-03-01")

    assert result["status"] == "ok"
    assert captured["flow"]["timestamp"].tolist() == [
        pd.Timestamp("2026-02-28 00:00"),
        pd.Timestamp("2026-03-01 23:59"),
    ]


def test_dry_risk_time_window_uses_only_window_rows(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    deps = make_deps(tmp_path)
    flow = sample_two_point_pattern_flow()
    captured: dict[str, pd.DataFrame] = {}
    monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_args, **_kwargs: flow)

    def capture_stats(window_flow: pd.DataFrame, _sites: pd.DataFrame) -> pd.DataFrame:
        captured["flow"] = window_flow.copy()
        return pd.DataFrame([{"point_id": "W1"}])

    monkeypatch.setattr("agent.tools.module_tools.dry_statistics", capture_stats)
    monkeypatch.setattr(
        "agent.tools.module_tools.assess_risk",
        lambda *_args, **_kwargs: {"dry_risk": pd.DataFrame(), "rainy_risk": pd.DataFrame()},
    )

    result = assess_risk_impl(
        deps,
        scope="dry",
        points=["W1"],
        start="2026-01-01 00:30",
        end="2026-01-01 00:39",
    )

    assert result["status"] == "ok"
    assert len(captured["flow"]) == 10
    assert captured["flow"]["timestamp"].min() == pd.Timestamp("2026-01-01 00:30")
    assert captured["flow"]["timestamp"].max() == pd.Timestamp("2026-01-01 00:39")
    assert not deps.paths.combined_xlsx.exists()


def test_data_filter_writes_pipeline_style_filter_result(tmp_path: Path) -> None:
    from openpyxl import load_workbook

    deps = make_deps(tmp_path)
    deps.session.auto_confirm_filter_result = True
    write_filter_sample_data(deps)
    raw_flow = io.load_flow(root=deps.paths.root)

    result = data_filter_impl(deps)
    filtered_flow = io.load_filtered_flow(root=deps.paths.root)

    assert result["status"] == "ok"
    assert len(raw_flow) == 4 * 1440
    assert deps.paths.filter_result.exists()
    assert result["data"]["selected"] == {"W1": ["2026-01-02"]}
    assert len(filtered_flow) == 1440
    assert set(filtered_flow["timestamp"].dt.strftime("%Y-%m-%d")) == {"2026-01-02"}

    wb = load_workbook(deps.paths.filter_result)
    ws = wb["筛选结果"]
    headers = [cell.value for cell in ws[1]]
    assert headers[:5] == ["点位编号", "2026-01-01", "2026-01-02", "2026-01-03", "2026-01-04"]
    assert headers[-1] == "筛选说明"
    assert ws.cell(row=2, column=1).value == "当天雨量"
    assert ws.cell(row=3, column=1).value == "W1"
    assert str(ws.cell(row=3, column=3).fill.start_color.index).upper().endswith("92D050")


def test_rainfall_and_event_needs_input(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_sample_data(deps)

    rain = analyze_rainfall_impl(deps)
    response = analyze_event_response_impl(deps)

    assert rain["status"] == "ok"
    assert not deps.paths.combined_xlsx.exists()
    assert "综合分析结果.xlsx" not in rain["summary"]
    assert response["status"] == "needs_input"
    assert response["missing"] == "event_ids"
    assert response["options"]


def test_event_response_rdii_and_risk_with_event_ids(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_sample_data(deps)
    analyze_rainfall_impl(deps)

    response = analyze_event_response_impl(deps, event_ids=[1])
    rdii = analyze_rdii_impl(deps, event_ids=[1])
    risk = assess_risk_impl(deps, scope="all", event_ids=[1])

    assert response["status"] == "ok"
    assert rdii["status"] == "ok"
    assert risk["status"] == "ok"


def test_event_response_impl_marks_no_monitoring_coverage(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    deps = make_deps(tmp_path)
    write_sample_data(deps)
    analyze_rainfall_impl(deps)
    monkeypatch.setattr("agent.tools.module_tools.analyze_event_response", lambda *_args, **_kwargs: pd.DataFrame())

    response = analyze_event_response_impl(deps, event_ids=[1], points=["W9"])

    assert response["status"] == "needs_input"
    assert response["missing"] == "data_coverage"
    assert "该时段/该点位无数据，无法分析" in response["summary"]
    assert deps.session.unavailable_event_ids == [1]


def test_event_response_guard_uses_real_timestamps_for_partial_coverage(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_sample_data(deps)
    pd.DataFrame(
        {
            "数据时间": pd.date_range("2026-02-01", periods=60, freq="min"),
            "设备编号": ["200"] * 60,
            "流量(L/s)(均值)": [2.0] * 60,
            "液位(m)(均值)": [0.4] * 60,
            "流速(m/s)(均值)": [0.5] * 60,
        }
    ).to_csv(deps.paths.flow_dir / "200_W2.csv", index=False, encoding="utf-8-sig")
    analyze_rainfall_impl(deps)

    result = analyze_event_response_impl(deps, event_ids=[1], points=["W1", "W2"])

    assert result["status"] == "ok"
    assert result["data"]["covered_points"] == ["W1"]
    assert result["data"]["excluded_points"] == [
        {"point_id": "W2", "reason": "该时段/该点位无数据，无法分析"}
    ]
    assert [row["point_id"] for row in result["data"]["table"]] == ["W1"]


@pytest.mark.parametrize("tool_name", ["event_response", "rdii", "risk"])
def test_event_tools_guard_before_analysis_when_no_point_has_coverage(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    tool_name: str,
) -> None:
    deps = make_deps(tmp_path)
    write_sample_data(deps)
    analyze_rainfall_impl(deps)
    monkeypatch.setattr(
        "agent.tools.module_tools._event_data_coverage",
        lambda *_args, **_kwargs: (
            pd.DataFrame(),
            pd.DataFrame(),
            [],
            [{"point_id": "W1", "reason": "该时段/该点位无数据，无法分析"}],
        ),
    )

    if tool_name == "event_response":
        result = analyze_event_response_impl(deps, event_ids=[1], points=["W1"])
    elif tool_name == "rdii":
        result = analyze_rdii_impl(deps, event_ids=[1], points=["W1"])
    else:
        result = assess_risk_impl(deps, scope="rainy", event_ids=[1])

    assert result["status"] == "needs_input"
    assert result["missing"] == "data_coverage"
    assert "该时段/该点位无数据，无法分析" in result["summary"]


@pytest.mark.parametrize("tool_name", ["event_response", "rdii", "risk"])
def test_event_tools_exclude_uncovered_points_and_continue(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    tool_name: str,
) -> None:
    deps = make_deps(tmp_path)
    write_sample_data(deps)
    analyze_rainfall_impl(deps)
    covered_flow = io.load_flow(points=["W1"], root=deps.paths.root)
    events = pd.DataFrame(
        {
            "event_id": [1],
            "start_time": ["2026-01-01 00:00"],
            "end_time": ["2026-01-01 01:00"],
            "rain_level": ["小雨"],
        }
    )
    excluded = [{"point_id": "W2", "reason": "该时段/该点位无数据，无法分析"}]
    monkeypatch.setattr(
        "agent.tools.module_tools._event_data_coverage",
        lambda *_args, **_kwargs: (covered_flow, events, ["W1"], excluded),
    )

    if tool_name == "event_response":
        result = analyze_event_response_impl(deps, event_ids=[1], points=["W1", "W2"])
    elif tool_name == "rdii":
        monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_args, **_kwargs: covered_flow)
        monkeypatch.setattr("agent.tools.module_tools.build_dry_curves", lambda *_args, **_kwargs: {})
        monkeypatch.setattr(
            "agent.tools.module_tools.analyze_rdii",
            lambda *_args, **_kwargs: {
                "rdii_total": pd.DataFrame([{"event_id": 1, "point_id": "W1", "rdii_total_m3": 1.0}]),
                "rdii_curve_data": {},
            },
        )
        result = analyze_rdii_impl(deps, event_ids=[1], points=["W1", "W2"])
    else:
        monkeypatch.setattr(
            "agent.tools.module_tools._dry_inputs",
            lambda *_args, **_kwargs: (pd.DataFrame(), pd.DataFrame(), {}),
        )
        result = assess_risk_impl(deps, scope="rainy", event_ids=[1])

    assert result["status"] == "ok"
    assert result["data"]["covered_points"] == ["W1"]
    assert result["data"]["excluded_points"] == excluded
    assert "剔除无覆盖点位 ['W2']" in result["summary"]


def test_report_refuses_event_without_monitoring_coverage(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_sample_data(deps)
    deps.session.selected_event_ids = [4]
    deps.session.unavailable_event_ids = [4]

    report = generate_report_impl(deps, sections=["RDII", "雨天风险"], event_ids=[4])

    assert report["status"] == "error"
    assert "无时间重叠" in report["summary"]
    assert not (deps.paths.outputs / "分析报告.docx").exists()


def test_patterns_and_report_success(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_sample_data(deps)

    patterns = analyze_patterns_impl(deps)
    report = generate_report_impl(deps, sections=["数据体检", "排污规律"])

    assert patterns["status"] == "ok"
    assert report["status"] == "ok"
    assert (deps.paths.outputs / "全网_全时段_分析报告.docx").exists()


def test_list_results_exposes_fresh_manifest_metadata(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_sample_data(deps)
    artifact = deps.paths.outputs / "sample.xlsx"
    artifact.write_text("sample", encoding="utf-8")
    record_result(deps, "data_filter", [artifact.relative_to(deps.paths.root).as_posix()], params={})

    result = list_results_impl(deps)
    item = result["data"]["manifest"]["data_filter"]

    assert result["status"] == "ok"
    assert item["fresh"] is True
    assert item["params"] == {}


def _install_report_stubs(monkeypatch: pytest.MonkeyPatch, captured: dict, counts: dict) -> None:
    def selected(points: list[str] | None) -> list[str]:
        return points or ["W1", "W2"]

    def fake_check(_deps, points=None, export=False, start=None, end=None):
        counts["check"] = counts.get("check", 0) + 1
        captured["check_scope"] = (points, start, end)
        return ok("checked", table=[{"point_id": point, "record_count": 10} for point in selected(points)])

    def fake_rain(_deps, time_range=None, output="all", rainfall_gap_hours=12):
        counts["rain"] = counts.get("rain", 0) + 1
        captured["rain_scope"] = time_range
        return ok(
            "rain",
            daily=[{"date": "2026-03-10", "daily_rain_mm": 5.0, "is_rainy": True}],
            events=[{
                "event_id": 1,
                "start_time": "2026-03-10 01:00",
                "end_time": "2026-03-10 03:00",
                "total_rain_mm": 5.0,
            }],
        )

    def fake_patterns(_deps, points=None, output="all", export=False, start=None, end=None, report_charts=False):
        counts["patterns"] = counts.get("patterns", 0) + 1
        captured["pattern_scope"] = (points, start, end)
        return ok("patterns", table=[{"point_id": point, "category": 1} for point in selected(points)])

    def fake_risk(_deps, scope="all", event_ids=None, points=None, export=False, start=None, end=None):
        counts["risk"] = counts.get("risk", 0) + 1
        captured["risk_scope"] = (scope, event_ids, points, start, end)
        rows = selected(points)
        return ok(
            "risk",
            dry_analysis=[{"point_id": point} for point in rows],
            dry_risk=[{"point_id": point, "overflow_value": 0.2} for point in rows],
            rainy_risk=[{"point_id": point, "overflow_value": 0.3} for point in rows],
        )

    def fake_build(output_file: Path, *_args, **kwargs):
        counts["build"] = counts.get("build", 0) + 1
        captured["build"] = kwargs
        output_file.write_bytes(b"report")
        return {
            "output_file": str(output_file),
            "templated_sections": kwargs.get("sections") or [],
            "generated_sections": [],
            "stats": {},
        }

    monkeypatch.setattr("agent.tools.module_tools.check_data_impl", fake_check)
    monkeypatch.setattr("agent.tools.module_tools.analyze_rainfall_impl", fake_rain)
    monkeypatch.setattr("agent.tools.module_tools.analyze_patterns_impl", fake_patterns)
    monkeypatch.setattr("agent.tools.module_tools.assess_risk_impl", fake_risk)
    monkeypatch.setattr("agent.tools.module_tools.build_report", fake_build)


def test_default_full_report_has_all_scope_and_nonempty_rainy_risk(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    deps = make_deps(tmp_path)
    captured: dict = {}
    counts: dict = {}
    _install_report_stubs(monkeypatch, captured, counts)

    result = generate_report_impl(deps, event_ids=[1])

    assert result["status"] == "ok"
    build = captured["build"]
    assert build["point_ids"] is None
    assert build["start"] is None and build["end"] is None
    assert build["sections"] == ["监测概况", "降雨分析", "旱天排污规律统计分析", "污水系统运行风险分析"]
    assert not build["analysis_tables"]["rainy_overflow_risk"].empty
    assert set(build["analysis_tables"]["rainy_overflow_risk"]["point_id"]) == {"W1", "W2"}


def test_single_point_report_does_not_fall_back_to_full_network(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    deps = make_deps(tmp_path)
    captured: dict = {}
    counts: dict = {}
    _install_report_stubs(monkeypatch, captured, counts)

    result = generate_report_impl(deps, points=["W1"], event_ids=[1])

    assert result["status"] == "ok"
    build = captured["build"]
    assert build["point_ids"] == ["W1"]
    for name in ("data_collection", "pattern_analysis", "dry_risk", "rainy_overflow_risk"):
        assert set(build["analysis_tables"][name]["point_id"]) == {"W1"}
    assert (deps.paths.outputs / "W1_全时段_分析报告.docx").exists()


def test_time_window_report_passes_one_scope_to_all_analyses(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    deps = make_deps(tmp_path)
    captured: dict = {}
    counts: dict = {}
    _install_report_stubs(monkeypatch, captured, counts)
    monkeypatch.setattr(
        "agent.tools.module_tools._resolved_report_time_range",
        lambda *_args: ["2026-03-07", "2026-03-10"],
    )
    monkeypatch.setattr(
        "agent.tools.module_tools._report_actual_time_range",
        lambda *_args, **_kwargs: ("2026-03-08 00:00:00", "2026-03-09 23:59:00"),
    )

    result = generate_report_impl(
        deps,
        start="2026-03-07",
        end="2026-03-10",
        event_ids=[1],
    )

    assert result["status"] == "ok"
    assert captured["check_scope"] == (None, "2026-03-07", "2026-03-10")
    assert captured["rain_scope"] == ["2026-03-07", "2026-03-10"]
    assert captured["pattern_scope"] == (None, "2026-03-07", "2026-03-10")
    assert captured["risk_scope"] == ("all", [1], None, "2026-03-07", "2026-03-10")
    assert captured["build"]["start"] == "2026-03-08 00:00:00"
    assert captured["build"]["end"] == "2026-03-09 23:59:00"
    combined = deps.paths.outputs / "全网_2026-03-07_2026-03-10_综合分析结果.xlsx"
    assert combined.exists()
    assert not deps.paths.combined_xlsx.exists()
    assert result["data"]["result_destinations"][0]["path"] == "outputs/全网_2026-03-07_2026-03-10_综合分析结果.xlsx"


def test_report_actual_time_range_uses_raw_flow_even_for_dry_reports(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    deps = make_deps(tmp_path)
    raw_flow = pd.DataFrame(
        {
            "timestamp": pd.to_datetime(["2026-03-10 00:00:00", "2026-03-15 23:59:00"]),
            "point_id": ["W1", "W1"],
            "flow_lps": [1.0, 2.0],
            "level_m": [0.1, 0.2],
            "velocity_mps": [0.3, 0.4],
        }
    )
    filtered_dry_flow = pd.DataFrame(
        {
            "timestamp": pd.to_datetime(["2026-03-11 00:00:00", "2026-03-14 23:59:00"]),
            "point_id": ["W1", "W1"],
            "flow_lps": [1.0, 2.0],
            "level_m": [0.1, 0.2],
            "velocity_mps": [0.3, 0.4],
        }
    )
    monkeypatch.setattr("agent.tools.module_tools.io.load_flow", lambda *_args, **_kwargs: raw_flow)
    monkeypatch.setattr(
        "agent.tools.module_tools._load_filtered_dry_flow",
        lambda *_args, **_kwargs: filtered_dry_flow,
    )

    start, end = _report_actual_time_range(deps, points=["W1"], start=None, end=None)

    assert start == "2026-03-10"
    assert end == "2026-03-15"
    assert (start, end) != ("2026-03-11", "2026-03-14")


def test_generate_report_uses_raw_flow_period_and_matching_combined_name(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    deps = make_deps(tmp_path)
    captured: dict = {}
    counts: dict = {}
    _install_report_stubs(monkeypatch, captured, counts)
    raw_flow = pd.DataFrame(
        {
            "timestamp": pd.to_datetime(["2026-03-10 00:00:00", "2026-03-15 23:59:00"]),
            "point_id": ["W1", "W1"],
            "flow_lps": [1.0, 2.0],
            "level_m": [0.1, 0.2],
            "velocity_mps": [0.3, 0.4],
        }
    )
    monkeypatch.setattr("agent.tools.module_tools.io.load_flow", lambda *_args, **_kwargs: raw_flow)

    result = generate_report_impl(deps, sections=["数据概况", "排污规律", "旱天风险"])

    assert result["status"] == "ok"
    assert captured["build"]["start"] == "2026-03-10"
    assert captured["build"]["end"] == "2026-03-15"
    assert (captured["build"]["start"], captured["build"]["end"]) != (
        "2026-03-11",
        "2026-03-14",
    )
    report_path = Path(result["data"]["output_file"])
    combined_path = deps.paths.root / result["data"]["result_destinations"][0]["path"]
    expected_combined_stem = report_path.stem.removesuffix("_分析报告") + "_综合分析结果"
    assert combined_path.name == expected_combined_stem + ".xlsx"
    assert combined_path.exists()


def test_check_data_time_window_uses_only_window_rows(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_two_point_data(deps)

    result = check_data_impl(
        deps,
        points=["W1"],
        start="2026-01-01 00:10:00",
        end="2026-01-01 00:19:00",
    )

    assert result["status"] == "ok"
    assert result["data"]["table"][0]["record_count"] == 10
    assert result["data"]["table"][0]["theoretical_count"] == 10
    assert result["data"]["table"][0]["collection_rate"] == 1.0
    assert result["data"]["window_coverage"]["actual_start"].startswith("2026-01-01 00:10:00")
    assert not deps.paths.combined_xlsx.exists()


def test_report_with_selected_sections_only_computes_selected_data(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    deps = make_deps(tmp_path)
    captured: dict = {}
    counts: dict = {}
    _install_report_stubs(monkeypatch, captured, counts)

    result = generate_report_impl(deps, sections=["监测概况"])

    assert result["status"] == "ok"
    assert counts == {"check": 1, "build": 1}
    assert captured["build"]["sections"] == ["监测概况"]
    assert set(captured["build"]["analysis_tables"]) == {"data_collection"}
    combined = deps.paths.outputs / "全网_全时段_综合分析结果.xlsx"
    assert pd.ExcelFile(combined).sheet_names == ["数据收集率统计"]
    assert not deps.paths.combined_xlsx.exists()


def test_custom_dry_only_section_names_still_generate_docx(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    deps = make_deps(tmp_path)
    captured: dict = {}
    counts: dict = {}
    _install_report_stubs(monkeypatch, captured, counts)

    sections = ["数据概况", "旱天排污规律", "旱天运行风险评估", "结论与建议"]
    result = generate_report_impl(deps, sections=sections)

    assert result["status"] == "ok"
    assert counts == {"check": 1, "patterns": 1, "risk": 1, "build": 1}
    assert captured["risk_scope"][0] == "dry"
    output = Path(result["data"]["output_file"])
    assert output.suffix == ".docx" and output.exists()
    assert not list(deps.paths.outputs.glob("*.md"))


def test_dry_only_report_keeps_monitoring_and_writes_matching_combined_sheets(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    deps = make_deps(tmp_path)
    captured: dict = {}
    counts: dict = {}
    _install_report_stubs(monkeypatch, captured, counts)

    result = generate_report_impl(deps, sections=["旱天分析", "排污规律", "旱天风险"])

    assert result["status"] == "ok"
    assert counts == {"check": 1, "patterns": 1, "risk": 1, "build": 1}
    assert captured["risk_scope"][0] == "dry"
    assert captured["build"]["sections"][0] == "监测概况"
    assert set(captured["build"]["analysis_tables"]) == {
        "data_collection",
        "pattern_analysis",
        "dry_analysis",
        "dry_risk",
    }
    combined = deps.paths.outputs / "全网_全时段_综合分析结果.xlsx"
    assert pd.ExcelFile(combined).sheet_names == [
        "数据收集率统计",
        "排污规律分析",
        "旱天分析",
        "旱天风险",
    ]
    assert not deps.paths.combined_xlsx.exists()
    assert result["data"]["report_combined_sheets"] == [
        "数据收集率统计",
        "排污规律分析",
        "旱天分析",
        "旱天风险",
    ]


def test_repeated_same_scope_report_reuses_analysis(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    deps = make_deps(tmp_path)
    pd.DataFrame({"点位编号": ["W1", "W2"], "管径": [1.0, 1.0]}).to_excel(
        deps.paths.site_info_file, index=False
    )
    flow = sample_two_point_pattern_flow()
    monkeypatch.setattr("agent.tools.module_tools._load_filtered_dry_flow", lambda *_a, **_k: flow)
    calls = {"analysis": 0, "build": 0}

    def fake_analysis(_flow, **_kwargs):
        calls["analysis"] += 1
        return {"patterns": pd.DataFrame([{"point_id": "W1", "category": 1}]), "curves": {}}

    def fake_build(output_file: Path, *_args, **kwargs):
        calls["build"] += 1
        output_file.write_bytes(b"report")
        return {"output_file": str(output_file), "templated_sections": kwargs["sections"], "generated_sections": [], "stats": {}}

    monkeypatch.setattr("agent.tools.module_tools.analyze_patterns", fake_analysis)
    monkeypatch.setattr("agent.tools.module_tools.build_report", fake_build)

    first = generate_report_impl(deps, points=["W1"], sections=["排污规律"])
    second = generate_report_impl(deps, points=["W1"], sections=["排污规律"])

    assert first["status"] == second["status"] == "ok"
    assert calls == {"analysis": 1, "build": 2}


def _fixed_template_tables() -> dict[str, pd.DataFrame]:
    return {
        "data_collection": pd.DataFrame([{
            "point_id": "W1", "record_count": 1440, "monitoring_days": 1,
            "theoretical_count": 1440, "collection_rate": 1.0,
        }]),
        "rainfall_daily": pd.DataFrame([{"date": "2026-03-10", "daily_rain_mm": 5.0}]),
        "rainfall_events": pd.DataFrame([{
            "event_id": 1, "start_time": "2026-03-10 01:00", "end_time": "2026-03-10 03:00",
            "total_rain_mm": 5.0, "duration_h": 2.0, "avg_intensity_mmh": 2.5, "rain_level": "小雨",
        }]),
        "pattern_analysis": pd.DataFrame([{
            "point_id": "W1", "category": 1, "category_name": "第1类 符合生活用水规律",
            "description": "W1点位符合生活用水规律。",
        }]),
        "dry_risk": pd.DataFrame([{
            "index": 1, "point_id": "W1", "diameter_m": 1.0, "well_depth_m": 3.0,
            "dry_velocity_mps": 0.5, "max_level_m": 0.8, "max_fullness": 0.8,
            "overflow_value": 0.27, "silting_risk": "中风险", "running_risk": "低风险",
            "overflow_risk": "低风险",
        }]),
        "rainy_overflow_risk": pd.DataFrame([{
            "event_id": 1, "point_id": "W1", "max_level_m": 1.2, "well_depth_m": 3.0,
            "overflow_value": 0.4, "overflow_risk": "低风险",
        }]),
    }


def _write_test_png(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
    ))


def _word_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    values.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    return "\n".join(values)


def test_fixed_template_report_uses_in_memory_tables_and_fills_rainy_section(tmp_path: Path) -> None:
    template = next((Path(__file__).parents[1] / "templates").glob("*.docx"))
    site_info = tmp_path / "点位信息.xlsx"
    pd.DataFrame([{
        "点位编号": "W1", "设备类型": "流量计", "形状": "圆管", "管径(m)": 1.0,
        "井深(m)": 3.0, "设备安装时间": "2026-03-01",
    }]).to_excel(site_info, index=False)
    output = tmp_path / "W1_分析报告.docx"
    flow_png = tmp_path / "charts" / "W1_流量特征曲线.png"
    level_png = tmp_path / "charts" / "W1_液位特征曲线.png"
    _write_test_png(flow_png)
    _write_test_png(level_png)

    result = build_report(
        output,
        "排水监测数据分析报告",
        template_file=template,
        analysis_tables=_fixed_template_tables(),
        site_info_file=site_info,
        sections=["监测概况", "降雨分析", "旱天排污规律统计分析", "污水系统运行风险分析"],
        point_ids=["W1"],
        start="2026-03-07",
        end="2026-03-10",
        pattern_chart_paths={"W1": [str(flow_png), str(level_png)]},
    )

    rendered = Document(output)
    text = _word_text(rendered)
    assert result["template_used"] is True
    assert result["stats"]["points_processed"] == 1
    assert not any("pipeline_report_assembler failed" in warning for warning in result.get("warnings", []))
    assert "雨天运行风险分析" in text
    assert "雨天溢流风险数据暂不完整" not in text
    assert "W1" in text
    assert len(rendered._element.xpath(".//w:drawing")) >= 4
    assert sum(1 for rel in rendered.part.rels.values() if "image" in rel.reltype) >= 4
    assert not any(value in text for value in ("W21", "W32", "44个流量监测点位", "2024/9/18", "2024/11/26"))
    assert not any(point in text for point in ("W20", "W28", "W34"))
    assert output.suffix == ".docx"
    assert zipfile.is_zipfile(output)
    assert not list(tmp_path.glob("*.md"))


def test_fixed_template_selected_sections_remove_unselected_chapters(tmp_path: Path) -> None:
    template = next((Path(__file__).parents[1] / "templates").glob("*.docx"))
    site_info = tmp_path / "点位信息.xlsx"
    pd.DataFrame([{
        "点位编号": "W1", "设备类型": "流量计", "形状": "圆管", "管径(m)": 1.0,
        "井深(m)": 3.0, "设备安装时间": "2026-03-01",
    }]).to_excel(site_info, index=False)
    output = tmp_path / "概况报告.docx"

    build_report(
        output,
        "排水监测数据分析报告",
        template_file=template,
        analysis_tables={"data_collection": _fixed_template_tables()["data_collection"]},
        site_info_file=site_info,
        sections=["监测概况"],
        point_ids=["W1"],
    )

    text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert "监测概况" in text
    assert "降雨分析" not in text
    assert "旱天排污规律统计分析" not in text
    assert "污水系统运行风险" not in text


def test_fixed_template_dry_report_keeps_overview_and_removes_only_rain_sections(tmp_path: Path) -> None:
    template = next((Path(__file__).parents[1] / "templates").glob("*.docx"))
    site_info = tmp_path / "点位信息.xlsx"
    pd.DataFrame([{
        "点位编号": "W1", "设备类型": "流量计", "形状": "圆管", "管径(m)": 1.0,
        "井深(m)": 3.0, "设备安装时间": "2026-03-01",
    }]).to_excel(site_info, index=False)
    output = tmp_path / "旱天报告.docx"
    flow_png = tmp_path / "charts" / "W1_流量特征曲线.png"
    level_png = tmp_path / "charts" / "W1_液位特征曲线.png"
    _write_test_png(flow_png)
    _write_test_png(level_png)

    build_report(
        output,
        "排水监测数据分析报告",
        template_file=template,
        analysis_tables={
            "data_collection": _fixed_template_tables()["data_collection"],
            "pattern_analysis": _fixed_template_tables()["pattern_analysis"],
            "dry_risk": _fixed_template_tables()["dry_risk"],
        },
        site_info_file=site_info,
        sections=["监测概况", "旱天排污规律统计分析", "旱天风险"],
        point_ids=["W1"],
        has_rainfall_data=False,
        pattern_chart_paths={"W1": [str(flow_png), str(level_png)]},
    )

    text = _word_text(Document(output))
    assert "监测概况" in text
    assert "监测设备安装" in text
    assert "监测数据质量" in text
    assert "旱天排污规律统计分析" in text
    assert "旱天运行风险分析" in text
    assert "降雨分析" not in text
    assert "降雨日分析" not in text
    assert "降雨场次分析" not in text
    assert "雨天运行风险分析" not in text


def test_fixed_template_dry_risk_only_removes_rain_and_exactly_sizes_table(tmp_path: Path) -> None:
    template = next((Path(__file__).parents[1] / "templates").glob("*.docx"))
    site_info = tmp_path / "点位信息.xlsx"
    pd.DataFrame([{
        "点位编号": "W1", "设备类型": "流量计", "形状": "圆管", "管径(m)": 1.0,
        "井深(m)": 3.0, "设备安装时间": "2026-03-01",
    }]).to_excel(site_info, index=False)
    output = tmp_path / "旱天风险报告.docx"

    build_report(
        output,
        "排水监测数据分析报告",
        template_file=template,
        analysis_tables={"dry_risk": _fixed_template_tables()["dry_risk"]},
        site_info_file=site_info,
        sections=["旱天风险"],
        point_ids=["W1"],
        has_rainfall_data=False,
    )

    rendered = Document(output)
    text = _word_text(rendered)
    template_map = scan_template(rendered)
    assert "旱天运行风险分析" in text
    assert "降雨分析" not in text
    assert "雨天运行风险分析" not in text
    assert template_map.get("rainy_overflow_risk") is None
    dry_table = template_map.get("dry_risk")
    assert dry_table is not None
    assert len(dry_table.rows) == 3  # two header rows + one actual data row
    assert all(any(cell.text.strip() for cell in row.cells) for row in dry_table.rows[2:])


def test_report_rejects_markdown_output(tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="必须是 .docx"):
        build_report(tmp_path / "分析报告.md", "排水监测数据分析报告")


def test_window_rainfall_events_are_locally_renumbered(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    pd.DataFrame({
        "timestamp": pd.to_datetime([
            "2026-01-01 00:00", "2026-01-20 00:00",
            "2026-02-05 00:00", "2026-02-20 00:00",
        ]),
        "rain": [2.0, 2.0, 3.0, 4.0],
    }).to_csv(deps.paths.rainfall_file, index=False)

    result = analyze_rainfall_impl(
        deps, time_range=["2026-02-01", "2026-02-28"], output="events"
    )

    events = result["data"]["events"]
    assert [event["source_event_id"] for event in events] == [3, 4]
    assert [event["event_id"] for event in events] == [1, 2]


def test_window_local_event_id_is_translated_for_downstream_analysis(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    deps = make_deps(tmp_path)
    deps.session.window_event_id_map = {1: 6}
    captured: dict[str, list[int]] = {}

    def fake_coverage(_deps, event_ids, points):
        captured["coverage"] = event_ids
        return pd.DataFrame(), pd.DataFrame(), ["W1"], []

    def fake_response(_flow, _events, event_ids):
        captured["analysis"] = event_ids
        return pd.DataFrame([{"event_id": 6, "point_id": "W1", "peak_flow_lps": 1.0}])

    monkeypatch.setattr("agent.tools.module_tools._event_data_coverage", fake_coverage)
    monkeypatch.setattr("agent.tools.module_tools.analyze_event_response", fake_response)

    result = analyze_event_response_impl(deps, event_ids=[1], points=["W1"])

    assert result["status"] == "ok"
    assert captured == {"coverage": [6], "analysis": [6]}
    assert result["data"]["table"][0]["event_id"] == 1
    assert "场次 [1]" in result["summary"]


def test_window_report_keeps_source_event_internal_and_renders_local_id(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    deps = make_deps(tmp_path)
    captured: dict = {}
    counts: dict = {}
    _install_report_stubs(monkeypatch, captured, counts)

    def fake_rain(_deps, time_range=None, output="all", rainfall_gap_hours=12):
        return ok(
            "rain",
            daily=[{"date": "2026-03-10", "daily_rain_mm": 5.0, "is_rainy": True}],
            events=[{
                "event_id": 1,
                "source_event_id": 6,
                "start_time": "2026-03-10 01:00",
                "end_time": "2026-03-10 03:00",
                "total_rain_mm": 5.0,
            }],
        )

    def fake_risk(_deps, scope="all", event_ids=None, points=None, export=False, start=None, end=None):
        captured["risk_event_ids"] = event_ids
        return ok(
            "risk",
            dry_analysis=[],
            dry_risk=[],
            rainy_risk=[{"event_id": 6, "point_id": "W1", "overflow_value": 0.3}],
        )

    monkeypatch.setattr("agent.tools.module_tools.analyze_rainfall_impl", fake_rain)
    monkeypatch.setattr("agent.tools.module_tools.assess_risk_impl", fake_risk)
    monkeypatch.setattr(
        "agent.tools.module_tools._resolved_report_time_range",
        lambda *_args: ["2026-03-10", "2026-03-12"],
    )

    result = generate_report_impl(
        deps,
        start="2026-03-10",
        end="2026-03-12",
        sections=["降雨分析", "雨天风险"],
        event_ids=[6],
    )

    assert result["status"] == "ok"
    assert captured["risk_event_ids"] == [6]
    rainy_table = captured["build"]["analysis_tables"]["rainy_overflow_risk"]
    assert rainy_table["event_id"].tolist() == [1]
    assert "窗口内降雨场次编号 [1]" in result["summary"]
    assert "窗口内降雨场次编号 [6]" not in result["summary"]


def test_open_time_ranges_use_natural_language_and_filename_tokens() -> None:
    assert _scope_period_text("2026-03-10", "") == "2026/03/10日之后"
    assert _scope_period_text("", "2026-03-10") == "2026/03/10日之前"
    assert "不限日" not in _scope_period_text("2026-03-10", "")
    assert _time_result_prefix("2026-03-10", None) == "2026-03-10_之后"
    assert _time_result_prefix(None, "2026-03-10") == "2026-03-10_之前"


def test_full_network_aliases_and_complete_list_use_short_full_network_filename(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    deps = make_deps(tmp_path)
    all_points = [f"W{index}" for index in range(1, 20)]
    pd.DataFrame({"点位编号": all_points, "管径": [1.0] * len(all_points)}).to_excel(
        deps.paths.site_info_file, index=False
    )
    for alias in (["全网"], ["全部点"], ["全部点位"], ["19个点"], all_points):
        assert is_full_network(alias, deps)

    captured: dict = {}
    counts: dict = {}
    _install_report_stubs(monkeypatch, captured, counts)
    full_result = generate_report_impl(deps, points=all_points, sections=["监测概况"])
    multi_result = generate_report_impl(deps, points=["W1", "W2", "W3"], sections=["监测概况"])

    assert full_result["status"] == multi_result["status"] == "ok"
    assert (deps.paths.outputs / "全网_全时段_分析报告.docx").exists()
    multi_files = list(deps.paths.outputs.glob("3点_*_分析报告.docx"))
    assert len(multi_files) == 1
    assert len(multi_files[0].name) <= 80
    assert "W2" not in multi_files[0].name and "W3" not in multi_files[0].name


def test_record_note_success(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    result = record_note_impl(deps, "sample note")
    assert result["status"] == "ok"
    assert "sample note" in deps.paths.notes.read_text(encoding="utf-8")


def test_run_python_success_with_analysis_io(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    write_sample_data(deps)
    result = run_python_impl(deps, "df = load_flow()\nprint(len(df))\n(WORKSPACE_DIR / 'out.txt').write_text(str(len(df)), encoding='utf-8')")
    assert result["status"] == "ok"
    assert (deps.paths.workspace / "out.txt").read_text(encoding="utf-8") == "60"


def test_run_python_rejects_markdown_report_fallback(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    result = run_python_impl(
        deps,
        "(OUTPUTS_DIR / '旱天分析报告.md').write_text('报告', encoding='utf-8')",
    )

    assert result["status"] == "error"
    assert "generate_report" in result["summary"]
    assert not (deps.paths.outputs / "旱天分析报告.md").exists()


def test_run_python_returns_error_without_crashing(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)
    result = run_python_impl(deps, "raise RuntimeError('kernel boom')")
    assert result["status"] == "error"
    assert "kernel boom" in result["data"]["stderr"]


def test_run_python_forces_utf8_stdout(tmp_path: Path) -> None:
    deps = make_deps(tmp_path)

    result = run_python_impl(deps, "print('RDII 单位：m³')")

    assert result["status"] == "ok"
    assert "RDII 单位：m³" in result["data"]["stdout"]


def test_run_python_timeout_is_killed(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    deps = make_deps(tmp_path)
    monkeypatch.setattr("agent.tools.python_tool.TIMEOUT_SECONDS", 1)

    started = time.monotonic()
    result = run_python_impl(deps, "import time\ntime.sleep(30)")
    elapsed = time.monotonic() - started

    assert result["status"] == "error"
    assert elapsed < 10
    assert result["data"]["script"]
