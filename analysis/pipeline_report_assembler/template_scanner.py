"""Semantic scanner for the report Word template."""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import Dict, List, Optional

from docx import Document
from docx.table import Table


@dataclass
class TemplateMap:
    """Tables in the Word template keyed by report semantic role."""

    tables: Dict[str, Table] = field(default_factory=dict)
    curve_tables: List[Table] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)

    def get(self, role: str) -> Optional[Table]:
        return self.tables.get(role)


def scan_template(doc: Document) -> TemplateMap:
    """Identify template table roles by headers and structure."""
    mapping = TemplateMap()

    for idx, table in enumerate(doc.tables):
        text = _table_text(table)
        header0 = _row_text(table, 0)
        header1 = _row_text(table, 1)

        role = None
        if _has_all(header0, "监测点位", "设备类型", "管径", "井深"):
            role = "site_info"
        elif _has_all(header0, "点位编号", "监测数据条数", "理论数据条数", "数据收集率"):
            role = "collection_rate"
        elif _has_all(header0, "日期", "日降雨量"):
            role = "rainfall_daily"
        elif _has_all(header0, "开始时间", "结束时间") and ("总降雨量" in header0 or "总降雨量/mm" in header0):
            role = "rainfall_events"
        elif _is_curve_table(table, text):
            mapping.curve_tables.append(table)
        elif _has_all(header1, "序号", "监测点位", "流速", "最大充满度", "淤积风险"):
            role = "dry_risk"
        elif _has_all(header0, "点位名称", "最大液位", "溢流风险值", "溢流风险"):
            role = "rainy_overflow_risk"

        if role:
            if role in mapping.tables:
                mapping.warnings.append(f"模板中发现重复表格角色 {role}，保留第一个，忽略 table[{idx}]")
            else:
                mapping.tables[role] = table

    required = ["site_info", "collection_rate", "dry_risk"]
    for role in required:
        if role not in mapping.tables:
            mapping.warnings.append(f"模板缺少关键表格角色: {role}")

    return mapping


def _table_text(table: Table) -> str:
    return "\n".join(
        " | ".join(cell.text.strip() for cell in row.cells)
        for row in table.rows
    )


def _row_text(table: Table, row_idx: int) -> str:
    if row_idx >= len(table.rows):
        return ""
    return " | ".join(cell.text.strip() for cell in table.rows[row_idx].cells)


def _has_all(text: str, *needles: str) -> bool:
    return all(needle in text for needle in needles)


def _is_curve_table(table: Table, text: str) -> bool:
    if len(table.rows) != 1:
        return False
    if len(table.rows[0].cells) != 2:
        return False
    return "流量特征曲线" in text and "液位特征曲线" in text
