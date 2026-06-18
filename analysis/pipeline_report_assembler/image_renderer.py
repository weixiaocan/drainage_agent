"""Curve image rendering for report templates."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.table import Table

from .style_writer import add_picture_to_cell, set_cell_text


def render_curve_images(
    doc: Document,
    curve_tables: List[Table],
    image_dir: Path,
    point_ids: Iterable[str],
) -> tuple[int, list[str]]:
    """Render flow/level curve images into semantic curve tables."""
    points = list(point_ids)
    warnings: list[str] = []
    if not points:
        return 0, warnings
    if not curve_tables:
        warnings.append("模板未找到特征曲线图表格")
        return 0, warnings

    _resize_curve_tables(doc, curve_tables, len(points))
    tables = _current_curve_tables(doc)
    inserted = 0

    for table, point_id in zip(tables, points):
        row = table.rows[0]
        flow_path = image_dir / f"{point_id}_流量特征曲线.png"
        level_path = image_dir / f"{point_id}_液位特征曲线.png"
        combined_path = image_dir / f"{point_id}_特征曲线.png"

        inserted += _insert_or_warn(row.cells[0], flow_path, combined_path, f"{point_id} 流量特征曲线", warnings)
        inserted += _insert_or_warn(row.cells[1], level_path, combined_path, f"{point_id} 液位特征曲线", warnings)

    return inserted, warnings


def _resize_curve_tables(doc: Document, curve_tables: List[Table], target_count: int) -> None:
    current = len(curve_tables)
    if current == target_count:
        return
    if current > target_count:
        for table in reversed(curve_tables[target_count:]):
            table._element.getparent().remove(table._element)
        return

    template = curve_tables[-1]
    parent = template._element.getparent()
    insert_after = template._element
    for _ in range(target_count - current):
        new_element = deepcopy(template._element)
        parent.insert(parent.index(insert_after) + 1, new_element)
        insert_after = new_element


def _current_curve_tables(doc: Document) -> List[Table]:
    result = []
    for table in doc.tables:
        if len(table.rows) == 1 and len(table.rows[0].cells) == 2:
            text = table.rows[0].cells[0].text + table.rows[0].cells[1].text
            if "流量特征曲线" in text or "液位特征曲线" in text or not text.strip():
                result.append(table)
    return result


def _insert_or_warn(cell, primary: Path, fallback: Path, label: str, warnings: list[str]) -> int:
    path = primary if primary.exists() else fallback
    if not path.exists():
        set_cell_text(cell, f"缺少{label}图片")
        warnings.append(f"缺少图片: {primary.name}")
        return 0
    try:
        add_picture_to_cell(cell, str(path), width_inches=2.8)
        return 1
    except Exception as exc:
        set_cell_text(cell, f"{label}图片插入失败")
        warnings.append(f"图片插入失败 {path.name}: {exc}")
        return 0
