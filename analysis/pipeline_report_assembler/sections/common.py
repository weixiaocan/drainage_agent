"""Shared helpers for report section renderers."""

from __future__ import annotations

from docx import Document
from docx.text.paragraph import Paragraph

from ..style_writer import set_paragraph_text


def find_paragraph(doc: Document, keyword: str, start: int = 0) -> int:
    for idx in range(start, len(doc.paragraphs)):
        if keyword in doc.paragraphs[idx].text:
            return idx
    return -1


def replace_first_paragraph(doc: Document, keyword: str, text: str) -> bool:
    idx = find_paragraph(doc, keyword)
    if idx < 0:
        return False
    set_paragraph_text(doc.paragraphs[idx], text)
    return True


def clear_paragraph_range(doc: Document, start_idx: int, end_idx: int) -> None:
    for idx in range(max(0, start_idx), min(end_idx, len(doc.paragraphs))):
        set_paragraph_text(doc.paragraphs[idx], "")


def paragraph_after(doc: Document, keyword: str) -> Paragraph | None:
    idx = find_paragraph(doc, keyword)
    if idx < 0 or idx + 1 >= len(doc.paragraphs):
        return None
    return doc.paragraphs[idx + 1]


def delete_body_range_between_paragraphs(doc: Document, start_keyword: str, end_keyword: str) -> bool:
    """Delete all body elements from the start paragraph up to, not including, end."""
    start_idx = find_paragraph(doc, start_keyword)
    end_idx = find_paragraph(doc, end_keyword, start=start_idx + 1 if start_idx >= 0 else 0)
    if start_idx < 0 or end_idx < 0 or end_idx <= start_idx:
        return False

    start_element = doc.paragraphs[start_idx]._p
    end_element = doc.paragraphs[end_idx]._p
    parent = start_element.getparent()
    removing = False
    for child in list(parent):
        if child is start_element:
            removing = True
        if child is end_element:
            break
        if removing:
            parent.remove(child)
    return True
