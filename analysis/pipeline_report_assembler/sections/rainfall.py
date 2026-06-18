"""Rainfall analysis section."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

from ..facts import ReportFacts
from ..report_tables import TABLE_SPECS, render_report_table
from ..style_writer import add_picture_to_paragraph, adjust_table_rows_preserve_style
from ..template_scanner import TemplateMap
from .common import delete_body_range_between_paragraphs, find_paragraph, replace_first_paragraph


def render_rainfall_section(
    doc: Document,
    template_map: TemplateMap,
    context,
    facts: ReportFacts,
    output_dir: Path,
    warnings: list[str],
) -> dict[str, int]:
    stats = {"tables_filled": 0, "text_replaced": 0, "images_inserted": 0}

    if not context.has_rainfall_data:
        if not delete_body_range_between_paragraphs(doc, "降雨分析", "旱天排污规律统计分析"):
            _clear_rainfall_tables(template_map)
            _replace_no_rain_text(doc)
            warnings.append("未能按章节删除降雨分析，已降级为清空降雨内容")
        return stats

    for role in ("rainfall_daily", "rainfall_events"):
        table = template_map.get(role)
        if table is None:
            warnings.append(f"降雨分析缺少表格: {role}")
            continue
        warnings.extend(render_report_table(table, TABLE_SPECS[role], context))
        stats["tables_filled"] += 1

    stats["text_replaced"] += _replace_rainfall_text(doc, facts)
    inserted, image_warnings = _insert_rainfall_images(doc, context, output_dir)
    stats["images_inserted"] += inserted
    warnings.extend(image_warnings)
    return stats


def _replace_rainfall_text(doc: Document, facts: ReportFacts) -> int:
    count = 0
    if replace_first_paragraph(
        doc,
        "雨量计持续监测降雨数据",
        (
            f"监测期间，雨量计持续监测降雨数据。统计监测期内降雨情况，以天（24h）为统计单位，"
            f"降雨日天数为{facts.rainy_days}天，总降雨量{facts.total_rain_mm} mm，"
            f"日最大降雨量为{facts.max_daily_rain_mm} mm，发生在{facts.max_daily_rain_date}。"
            "降雨日各天的日降雨量统计如下表所示："
        ),
    ):
        count += 1
    rainy_pct = round(facts.rainy_days / facts.total_days * 100) if facts.total_days else 0
    non_rainy_pct = 100 - rainy_pct if facts.total_days else 0
    if replace_first_paragraph(
        doc,
        "非降雨日为",
        (
            f"监测期内共{facts.total_days}个自然日，其中降雨日为{facts.rainy_days}天，"
            f"非降雨日为{facts.non_rainy_days}天。降雨日为整个监测期内自然日的{rainy_pct}%，"
            f"非降雨日为{non_rainy_pct}%，如下图所示。"
        ),
    ):
        count += 1
    if replace_first_paragraph(
        doc,
        "累计降雨量",
        (
            f"考虑降雨发生后会引起雨水的径流现象，以下雨发生到结束后12小时为一个降雨场次。"
            f"监测期内共发生有效降雨场次{facts.rainfall_event_count}场，"
            f"累计降雨量{facts.event_total_rain_mm} mm，最大场次降雨量为{facts.max_event_rain_mm} mm。"
            "各降雨场次统计如下表所示："
        ),
    ):
        count += 1
    return count


def _insert_rainfall_images(doc: Document, context, output_dir: Path) -> tuple[int, list[str]]:
    chart_dir = output_dir / "降雨分析图"
    daily_path = chart_dir / "日降雨量时间序列图.png"
    ratio_path = chart_dir / "降雨日占比饼图.png"
    warnings: list[str] = []
    if not daily_path.exists() or not ratio_path.exists():
        warnings.extend(_generate_fallback_charts(context.df("rainfall_daily"), chart_dir))

    inserted = 0
    inserted += _insert_image_before_caption(doc, "图 15", daily_path, warnings)
    inserted += _insert_image_before_caption(doc, "图 16", ratio_path, warnings)
    return inserted, warnings


def _insert_image_before_caption(doc: Document, caption_keyword: str, image_path: Path, warnings: list[str]) -> int:
    idx = find_paragraph(doc, caption_keyword)
    if idx <= 0:
        warnings.append(f"未找到降雨图标题: {caption_keyword}")
        return 0
    if not image_path.exists():
        warnings.append(f"降雨图不存在: {image_path}")
        return 0
    add_picture_to_paragraph(doc.paragraphs[idx - 1], str(image_path), width_inches=5.6)
    return 1


def _generate_fallback_charts(df: pd.DataFrame, chart_dir: Path) -> list[str]:
    warnings: list[str] = []
    if df.empty:
        warnings.append("缺少降雨日数据，无法兜底生成降雨图")
        return warnings
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception as exc:
        warnings.append(f"matplotlib 不可用，无法兜底生成降雨图: {exc}")
        return warnings

    chart_dir.mkdir(parents=True, exist_ok=True)
    plot_df = df.copy()
    plot_df["date"] = pd.to_datetime(plot_df["date"], errors="coerce")
    plot_df["daily_rain_mm"] = pd.to_numeric(plot_df["daily_rain_mm"], errors="coerce").fillna(0)
    plt.rcParams["font.family"] = ["Times New Roman", "SimSun"]
    plt.rcParams["font.sans-serif"] = ["SimSun", "宋体", "Microsoft YaHei", "DejaVu Sans"]
    plt.rcParams["font.serif"] = ["Times New Roman", "SimSun"]
    plt.rcParams["axes.unicode_minus"] = False

    fig, ax = plt.subplots(figsize=(9.2, 4.8), dpi=180)
    labels = [d.strftime("%Y-%m-%d") if not pd.isna(d) else "" for d in plot_df["date"]]
    x = range(len(labels))
    ax.bar(x, plot_df["daily_rain_mm"], color="#5B9BD5", edgecolor="#2F5597", linewidth=0.6)
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels)
    ax.set_title("日降雨量时间序列")
    ax.set_ylabel("降雨量(mm)")
    ax.set_xlabel("日期")
    ax.tick_params(axis="x", rotation=45, labelsize=7.5)
    ax.tick_params(axis="y", labelsize=9)
    ax.grid(False)
    for spine in ax.spines.values():
        spine.set_visible(True)
        spine.set_color("#000000")
        spine.set_linewidth(0.8)
    fig.tight_layout()
    fig.savefig(chart_dir / "日降雨量时间序列图.png", bbox_inches="tight")
    plt.close(fig)

    rainy_days = int((plot_df["daily_rain_mm"] > 0).sum())
    non_rainy_days = len(plot_df) - rainy_days
    total_days = max(1, rainy_days + non_rainy_days)
    fig, ax = plt.subplots(figsize=(4.8, 4.8), dpi=180)
    label_iter = iter(["降雨日", "非降雨日"])
    ax.pie(
        [rainy_days, non_rainy_days],
        labels=["", ""],
        autopct=lambda pct: _pie_autopct(pct, total_days, next(label_iter)),
        pctdistance=0.58,
        startangle=90,
        colors=["#5B9BD5", "#ED7D31"],
        wedgeprops={"edgecolor": "white", "linewidth": 1.0},
        textprops={"fontsize": 10, "color": "black", "ha": "center"},
    )
    ax.axis("equal")
    fig.tight_layout()
    fig.savefig(chart_dir / "降雨日占比饼图.png", bbox_inches="tight")
    plt.close(fig)
    warnings.append("降雨 PNG 缺失，已由报告组装兜底生成")
    return warnings


def _pie_autopct(pct: float, total: int, label: str) -> str:
    count = int(round(pct * total / 100.0))
    return f"{label}\n{count}天\n{pct:.0f}%"


def _clear_rainfall_tables(template_map: TemplateMap) -> None:
    for role in ("rainfall_daily", "rainfall_events", "rainy_overflow_risk"):
        table = template_map.get(role)
        if table is not None:
            adjust_table_rows_preserve_style(table, 0, TABLE_SPECS[role].template_row_idx)


def _replace_no_rain_text(doc: Document) -> None:
    replace_first_paragraph(doc, "雨量计持续监测降雨数据", "监测期间未识别到有效降雨数据，降雨分析相关图表和场次统计不纳入本次报告。")
    replace_first_paragraph(doc, "非降雨日为", "监测期间未识别到有效降雨日。")
    replace_first_paragraph(doc, "累计降雨量", "监测期间未识别到有效降雨场次。")
