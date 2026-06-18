"""Dry-weather discharge pattern section."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..facts import PatternPointFact, ReportFacts
from ..llm_section_writer import LLMSectionWriter
from ..style_writer import clear_cell, set_cell_text, set_paragraph_text
from ..template_scanner import TemplateMap
from .common import find_paragraph


CLASS_TITLES = {
    1: "（1）监测点位流量特征曲线符合生活用水规律",
    2: "（2）监测点位流量特征曲线有波峰或波谷但不符合生活用水规律",
    3: "（3）监测点位流量特征曲线无明显波峰或波谷",
}

CLASS_NAMES = {
    1: "符合生活用水规律",
    2: "有波峰或波谷但不符合典型生活用水规律",
    3: "无明显波峰或波谷",
}


def render_pattern_section(
    doc: Document,
    template_map: TemplateMap,
    facts: ReportFacts,
    image_dir: Path,
    llm_writer: LLMSectionWriter,
    warnings: list[str],
) -> dict[str, int]:
    stats = {"text_replaced": 0, "images_inserted": 0, "llm_generated": 0}
    start_idx = find_paragraph(doc, "第一轮监测点排污规律统计")
    risk_idx = find_paragraph(doc, "污水系统运行风险", start=start_idx + 1 if start_idx >= 0 else 0)
    if start_idx < 0 or risk_idx < 0:
        warnings.append("未找到排污规律章节边界")
        return stats

    templates = _collect_templates(doc, start_idx, risk_idx)
    if (
        templates["normal"] is None
        or templates["subsection"] is None
        or templates["caption"] is None
        or templates["curve_table"] is None
    ):
        warnings.append("排污规律章节模板元素不完整，无法按模板样式重建")
        return stats

    grouped = _group_pattern_points(facts)
    start_element = doc.paragraphs[start_idx]._p
    end_element = doc.paragraphs[risk_idx]._p
    _delete_between(start_element, end_element)

    anchor = start_element
    anchor = _insert_paragraph_after(doc, anchor, templates["normal"], _build_classification_summary(grouped, facts))
    stats["text_replaced"] += 1

    figure_no = 17
    for class_id in (1, 2, 3):
        points = grouped.get(class_id, [])
        if not points:
            continue
        anchor = _insert_paragraph_after(doc, anchor, templates["subsection"], CLASS_TITLES[class_id])
        anchor = _insert_paragraph_after(doc, anchor, templates["normal"], _build_class_intro(class_id, points))
        stats["text_replaced"] += 2

        for point in points:
            anchor = _insert_paragraph_after(doc, anchor, templates["normal"], _point_description(point))
            anchor, inserted = _insert_curve_table_after(doc, anchor, templates["curve_table"], image_dir, point.point_id, warnings)
            stats["images_inserted"] += inserted
            anchor = _insert_paragraph_after(
                doc,
                anchor,
                templates["caption"],
                f"图 {figure_no} {point.point_id}排污规律特征曲线图",
            )
            figure_no += 1
            stats["text_replaced"] += 2

    summary_text, used_llm = llm_writer.generate("排污规律统计本章小结", facts, _build_chapter_summary)
    summary_text = _validate_summary(summary_text, facts) or _build_chapter_summary(facts)
    title_template = templates["summary_title"] if templates["summary_title"] is not None else templates["subsection"]
    anchor = _insert_paragraph_after(doc, anchor, title_template, "本章小结")
    _insert_paragraph_after(doc, anchor, templates["normal"], summary_text)
    stats["text_replaced"] += 2
    stats["llm_generated"] += int(used_llm and summary_text != _build_chapter_summary(facts))
    return stats


def _collect_templates(doc: Document, start_idx: int, end_idx: int) -> dict[str, object]:
    start_element = doc.paragraphs[start_idx]._p
    end_element = doc.paragraphs[end_idx]._p
    parent = start_element.getparent()
    result: dict[str, object] = {
        "normal": None,
        "subsection": None,
        "caption": None,
        "summary_title": None,
        "curve_table": None,
    }

    in_range = False
    for child in list(parent):
        if child is start_element:
            in_range = True
            continue
        if child is end_element:
            break
        if not in_range:
            continue

        tag = _local_name(child)
        text = _element_text(child).strip()
        if tag == "p":
            if text == "本章小结" and result["summary_title"] is None:
                result["summary_title"] = deepcopy(child)
            elif text.startswith("（1）") and result["subsection"] is None:
                result["subsection"] = deepcopy(child)
            elif "排污规律特征曲线图" in text and result["caption"] is None:
                result["caption"] = deepcopy(child)
            elif text and result["normal"] is None:
                result["normal"] = deepcopy(child)
        elif tag == "tbl" and result["curve_table"] is None and _is_curve_table_text(text):
            result["curve_table"] = deepcopy(child)

    return result


def _delete_between(start_element, end_element) -> None:
    parent = start_element.getparent()
    removing = False
    for child in list(parent):
        if child is start_element:
            removing = True
            continue
        if child is end_element:
            break
        if removing:
            parent.remove(child)


def _insert_paragraph_after(doc: Document, anchor, template_element, text: str):
    new_element = deepcopy(template_element)
    anchor.addnext(new_element)
    paragraph = Paragraph(new_element, doc._body)
    set_paragraph_text(paragraph, text)
    return new_element


def _insert_curve_table_after(
    doc: Document,
    anchor,
    template_element,
    image_dir: Path,
    point_id: str,
    warnings: list[str],
) -> tuple[object, int]:
    new_element = deepcopy(template_element)
    anchor.addnext(new_element)
    table = Table(new_element, doc._body)
    inserted = _fill_curve_table(table, image_dir, point_id, warnings)
    return new_element, inserted


def _fill_curve_table(table: Table, image_dir: Path, point_id: str, warnings: list[str]) -> int:
    if not table.rows or len(table.rows[0].cells) < 2:
        warnings.append(f"排污规律曲线表格结构异常: {point_id}")
        return 0
    row = table.rows[0]
    _remove_table_borders(table)
    flow_path = image_dir / f"{point_id}_流量特征曲线.png"
    level_path = image_dir / f"{point_id}_液位特征曲线.png"
    combined_path = image_dir / f"{point_id}_特征曲线.png"
    inserted = 0
    inserted += _insert_curve_image(row.cells[0], flow_path, combined_path, f"{point_id}流量特征曲线", "（a）流量特征曲线图", warnings)
    inserted += _insert_curve_image(row.cells[1], level_path, combined_path, f"{point_id}液位特征曲线", "（b）液位特征曲线图", warnings)
    return inserted


def _insert_curve_image(cell, primary: Path, fallback: Path, label: str, subcaption: str, warnings: list[str]) -> int:
    path = primary if primary.exists() else fallback
    if not path.exists():
        set_cell_text(cell, f"缺少{label}图片\n{subcaption}")
        warnings.append(f"缺少图片: {primary.name}")
        return 0
    try:
        clear_cell(cell)
        image_para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_para.add_run().add_picture(str(path), width=Inches(2.8))
        caption_para = cell.paragraphs[1] if len(cell.paragraphs) > 1 else cell.add_paragraph()
        set_paragraph_text(caption_para, subcaption)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return 1
    except Exception as exc:
        set_cell_text(cell, f"{label}图片插入失败\n{subcaption}")
        warnings.append(f"图片插入失败 {path.name}: {exc}")
        return 0


def _group_pattern_points(facts: ReportFacts) -> dict[int, list[PatternPointFact]]:
    grouped = {1: [], 2: [], 3: []}
    details = facts.pattern_details or []
    if details:
        by_point = {item.point_id: item for item in details}
        for point_id in facts.point_ids:
            item = by_point.get(point_id)
            if item and item.category in grouped:
                grouped[item.category].append(item)
        for item in details:
            if item.category in grouped and item not in grouped[item.category]:
                grouped[item.category].append(item)
        return grouped

    for class_id, points in (facts.pattern_groups or {}).items():
        if class_id not in grouped:
            continue
        for point_id in points:
            grouped[class_id].append(PatternPointFact(point_id=point_id, category=class_id))
    return grouped


def _build_classification_summary(grouped: dict[int, list[PatternPointFact]], facts: ReportFacts) -> str:
    parts = [f"本轮监测的{facts.point_count}个点位根据旱天流量特征曲线形态可分为三类"]
    for class_id in (1, 2, 3):
        points = grouped.get(class_id, [])
        parts.append(f"第{class_id}类{CLASS_NAMES[class_id]}的点位有{len(points)}处，为{_join_points(p.point_id for p in points)}")
    return "；".join(parts) + "。"


def _build_class_intro(class_id: int, points: list[PatternPointFact]) -> str:
    return (
        f"该类共包括{len(points)}个点位：{_join_points(p.point_id for p in points)}。"
        "各点位排污规律具体情况如下。"
    )


def _point_description(point: PatternPointFact) -> str:
    description = str(point.description or "").strip()
    if description:
        if point.point_id not in description:
            description = f"{point.point_id}点位：{description}"
        return _ensure_sentence(description)

    parts = [f"{point.point_id}点位属于{point.category_name or f'第{point.category}类'}"]
    if point.peak_periods:
        parts.append(f"按小时阈值识别的相对高流量区间为{point.peak_periods}")
    if point.valley_periods:
        parts.append(f"相对低流量区间为{point.valley_periods}")
    if point.diagnosis_reason and not str(point.diagnosis_reason).strip().startswith("Kz="):
        parts.append(str(point.diagnosis_reason))
    if point.category == 2:
        parts.append("曲线存在波动但与典型生活用水规律不完全一致，需结合汇水范围、工业排放、商业活动或泵站调度等因素进一步判断")
    elif point.category == 3:
        parts.append("曲线整体波动不明显，需关注持续入渗或恒定排放影响")
    return _ensure_sentence("，".join(parts))


def _build_chapter_summary(facts: ReportFacts) -> str:
    grouped = _group_pattern_points(facts)
    class1 = grouped.get(1, [])
    class2 = grouped.get(2, [])
    class3 = grouped.get(3, [])
    focus = class2 + class3
    focus_points = _join_points(p.point_id for p in focus)
    return (
        f"本章对{facts.point_count}个监测点位的旱天排污规律进行了统计分析。"
        f"其中，第1类点位{len(class1)}处，曲线整体符合生活用水规律；"
        f"第2类点位{len(class2)}处，存在波峰或波谷但与典型生活用水规律不完全一致；"
        f"第3类点位{len(class3)}处，曲线较为平坦或波动特征不明显。"
        f"后续运行诊断中应重点关注{focus_points}等第2类和第3类点位，结合汇水范围、泵站调度和现场管网条件进一步核查异常成因。"
    )


def _validate_summary(text: str, facts: ReportFacts) -> str:
    value = str(text or "").strip()
    if not value:
        return ""
    if "两轮" in value or "32个点位" in value or "1-" in value:
        return ""
    if facts.point_count and str(facts.point_count) not in value:
        return ""
    return value


def _join_points(points: Iterable[str]) -> str:
    values = [str(point).strip() for point in points if str(point).strip()]
    return "、".join(values) if values else "无"


def _ensure_sentence(text: str) -> str:
    value = str(text or "").strip()
    if not value:
        return ""
    if value[-1] not in "。！？；;":
        return value + "。"
    return value


def _local_name(element) -> str:
    return element.tag.rsplit("}", 1)[-1]


def _element_text(element) -> str:
    return "".join(node.text or "" for node in element.iter() if node.tag.endswith("}t"))


def _is_curve_table_text(text: str) -> bool:
    return "流量特征曲线" in text and "液位特征曲线" in text


def _remove_table_borders(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    _set_borders_nil(borders, ("top", "left", "bottom", "right", "insideH", "insideV"))

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            cell_borders = tc_pr.first_child_found_in("w:tcBorders")
            if cell_borders is None:
                cell_borders = OxmlElement("w:tcBorders")
                tc_pr.append(cell_borders)
            _set_borders_nil(cell_borders, ("top", "left", "bottom", "right"))


def _set_borders_nil(parent, edges: tuple[str, ...]) -> None:
    for edge in edges:
        tag = f"w:{edge}"
        element = parent.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            parent.append(element)
        element.set(qn("w:val"), "nil")
        element.set(qn("w:sz"), "0")
