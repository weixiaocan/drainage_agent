"""Risk analysis section."""

from __future__ import annotations

import pandas as pd
from docx import Document
from docx.shared import Pt

from ..facts import ReportFacts
from ..llm_section_writer import LLMSectionWriter
from ..report_tables import TABLE_SPECS, render_report_table
from ..style_writer import set_paragraph_text
from ..template_scanner import TemplateMap
from .common import delete_body_range_between_paragraphs, find_paragraph, replace_first_paragraph


def render_risk_section(
    doc: Document,
    template_map: TemplateMap,
    context,
    facts: ReportFacts,
    llm_writer: LLMSectionWriter,
    warnings: list[str],
    include_dry: bool = True,
    include_rainy: bool = True,
) -> dict[str, int]:
    stats = {"tables_filled": 0, "text_replaced": 0, "llm_generated": 0}

    if not include_dry:
        delete_body_range_between_paragraphs(doc, "旱天运行风险分析", "雨天运行风险分析")
    if not include_rainy or not context.has_rainfall_data:
        delete_body_range_between_paragraphs(doc, "雨天运行风险分析", "本章小结")

    for role in ("dry_risk", "rainy_overflow_risk"):
        if role == "dry_risk" and not include_dry:
            continue
        if role == "rainy_overflow_risk" and (not include_rainy or not context.has_rainfall_data):
            continue
        table = template_map.get(role)
        if table is None:
            warnings.append(f"风险分析缺少表格: {role}")
            continue
        warnings.extend(render_report_table(table, TABLE_SPECS[role], context))
        stats["tables_filled"] += 1

    stats["text_replaced"] += _replace_risk_text(
        doc, context, facts, llm_writer, stats, include_dry=include_dry, include_rainy=include_rainy
    )
    return stats


def _replace_risk_text(
    doc: Document,
    context,
    facts: ReportFacts,
    llm_writer: LLMSectionWriter,
    stats: dict,
    include_dry: bool,
    include_rainy: bool,
) -> int:
    replaced = 0
    dry = context.df("dry_risk")
    rainy = context.df("rainy_overflow_risk")

    if replace_first_paragraph(doc, "最大充满度W5代表运行中风险", "最大充满度1-2代表运行中风险；"):
        replaced += 1
    if replace_first_paragraph(doc, "监测时段内", _risk_intro(facts, include_dry, include_rainy)):
        replaced += 1

    if include_dry:
        replaced += _replace_block(doc, "旱天最大充满度情况如下", _fullness_text(dry, facts), clear_next=4)
        replaced += _replace_block(doc, "旱天溢流风险值情况如下", _overflow_text(dry, facts), clear_next=3)
        replaced += _replace_block(doc, "淤积风险情况如下", _silting_text(dry, facts), clear_next=4)
    if include_rainy and context.has_rainfall_data:
        replaced += _replace_block(doc, "处监测点位在", _rainy_text(rainy, facts), clear_next=3)

    fallback = lambda f: _summary_text(dry, rainy, f, include_dry, include_rainy)
    summary_text, used_llm = llm_writer.generate("风险分析本章小结", facts, fallback)
    if _replace_block(doc, "本章从最大充满度", summary_text, clear_next=5):
        replaced += 1
        stats["llm_generated"] += int(used_llm)
    return replaced


def _risk_intro(facts: ReportFacts, include_dry: bool, include_rainy: bool) -> str:
    scopes = []
    if include_dry:
        scopes.append("旱天运行风险")
    if include_rainy:
        scopes.append(f"{facts.rainfall_event_count}场有效降雨下的雨天溢流风险")
    scope_text = "和".join(scopes) or "运行风险"
    return (
        f"本章对{facts.point_count}个监测点位的{scope_text}进行分析，"
        "以最大充满度、溢流风险值和旱天流速等实际监测指标评估排水系统运行状态。"
    )


def _replace_block(doc: Document, keyword: str, text: str | list[str], clear_next: int) -> int:
    idx = find_paragraph(doc, keyword)
    if idx < 0:
        return 0
    paragraphs = text if isinstance(text, list) else [text]
    for offset, paragraph_text in enumerate(paragraphs[: clear_next + 1]):
        target_idx = idx + offset
        if target_idx >= len(doc.paragraphs):
            break
        set_paragraph_text(doc.paragraphs[target_idx], paragraph_text)
        _format_risk_item_paragraph(doc.paragraphs[target_idx], paragraph_text)
    clear_start = idx + len(paragraphs)
    _clear_range_preserve_table_captions(doc, clear_start, idx + 1 + clear_next)
    return 1


def _clear_range_preserve_table_captions(doc: Document, start_idx: int, end_idx: int) -> None:
    for para_idx in range(max(0, start_idx), min(end_idx, len(doc.paragraphs))):
        text = doc.paragraphs[para_idx].text.strip()
        if text.startswith("表 "):
            continue
        set_paragraph_text(doc.paragraphs[para_idx], "")


def _format_risk_item_paragraph(paragraph, text: str) -> None:
    if not str(text).startswith(tuple("①②③④⑤⑥⑦⑧⑨")):
        return
    paragraph.paragraph_format.space_after = Pt(4)


def _fullness_text(df: pd.DataFrame, facts: ReportFacts) -> list[str]:
    if df.empty or "max_fullness" not in df.columns:
        return [f"监测期间，{facts.point_count}处监测点的旱天最大充满度数据暂不完整。"]
    values = pd.to_numeric(df["max_fullness"], errors="coerce").fillna(0)
    point_col = df["point_id"] if "point_id" in df.columns else pd.Series([""] * len(df))
    low = point_col[values < 0.75].tolist()
    mid_low = point_col[(values >= 0.75) & (values < 1.0)].tolist()
    mid = point_col[(values >= 1.0) & (values <= 2.0)].tolist()
    high = point_col[values > 2.0].tolist()
    return [f"监测期间，{facts.point_count}处监测点的旱天最大充满度情况如下："] + _numbered_items([
        (low, f"{len(low)}处监测点最大充满度小于0.75，运行良好，点位为{_join(low)}。"),
        (mid_low, f"{len(mid_low)}处监测点最大充满度为0.75-1.0，存在运行低风险，点位为{_join(mid_low)}。"),
        (mid, f"{len(mid)}处监测点最大充满度为1.0-2.0，存在运行中风险，点位为{_join(mid)}。"),
        (high, f"{len(high)}处监测点最大充满度大于2.0，存在运行高风险，点位为{_join(high)}。"),
    ])


def _overflow_text(df: pd.DataFrame, facts: ReportFacts) -> list[str]:
    if df.empty or "overflow_value" not in df.columns:
        return [f"监测期间，{facts.point_count}处监测点的旱天溢流风险值数据暂不完整。"]
    values = pd.to_numeric(df["overflow_value"], errors="coerce").fillna(0)
    point_col = df["point_id"] if "point_id" in df.columns else pd.Series([""] * len(df))
    low = point_col[values < 0.7].tolist()
    mid = point_col[(values >= 0.7) & (values < 0.9)].tolist()
    high = point_col[values >= 0.9].tolist()
    return [f"监测期间，{facts.point_count}处监测点的旱天溢流风险值情况如下："] + _numbered_items([
        (low, f"{len(low)}处监测点的溢流风险值小于0.7，溢流风险低，点位为{_join(low)}。"),
        (mid, f"{len(mid)}处监测点的溢流风险值为0.7-0.9，为溢流中风险，点位为{_join(mid)}。"),
        (high, f"{len(high)}处监测点的溢流风险值大于等于0.9，为溢流高风险及以上，点位为{_join(high)}。"),
    ])


def _silting_text(df: pd.DataFrame, facts: ReportFacts) -> list[str]:
    if df.empty or "dry_velocity_mps" not in df.columns:
        return [f"监测期间，{facts.point_count}处监测点的淤积风险数据暂不完整。"]
    values = pd.to_numeric(df["dry_velocity_mps"], errors="coerce").fillna(0)
    point_col = df["point_id"] if "point_id" in df.columns else pd.Series([""] * len(df))
    high = point_col[values <= 0.3].tolist()
    mid = point_col[(values > 0.3) & (values < 0.6)].tolist()
    low = point_col[values >= 0.6].tolist()
    return [f"监测期间，{facts.point_count}处监测点的淤积风险情况如下："] + _numbered_items([
        (high, f"{len(high)}处监测点的平均流速小于等于0.3m/s，为高淤积风险，点位为{_join(high)}。"),
        (mid, f"{len(mid)}处监测点的平均流速为0.3-0.6m/s，为中淤积风险，点位为{_join(mid)}。"),
        (low, f"{len(low)}处监测点的平均流速大于等于0.6m/s，为低淤积风险，点位为{_join(low)}。"),
    ])


def _rainy_text(df: pd.DataFrame, facts: ReportFacts) -> list[str]:
    if df.empty or "overflow_value" not in df.columns:
        return ["监测期间雨天溢流风险数据暂不完整。"]
    values = pd.to_numeric(df["overflow_value"], errors="coerce").fillna(0)
    point_col = df["point_id"] if "point_id" in df.columns else pd.Series([""] * len(df))
    low = point_col[values < 0.7].tolist()
    mid = point_col[(values >= 0.7) & (values < 0.9)].tolist()
    high = point_col[values >= 0.9].tolist()
    return [f"监测期间，{len(df)}处监测点位在选定降雨事件下的溢流风险情况如下所示："] + _numbered_items([
        (low, f"{len(low)}处监测点的溢流风险值小于0.7，溢流风险低，点位为{_join(low)}。"),
        (mid, f"{len(mid)}处监测点的溢流风险值为0.7-0.9，为溢流中风险，点位为{_join(mid)}。"),
        (high, f"{len(high)}处监测点的溢流风险值大于等于0.9，为溢流高风险及以上，点位为{_join(high)}。"),
    ])


def _summary_text(
    dry: pd.DataFrame,
    rainy: pd.DataFrame,
    facts: ReportFacts,
    include_dry: bool = True,
    include_rainy: bool = True,
) -> str:
    parts = [f"本章对{facts.point_count}个监测点位的污水系统运行风险进行了评估。"]
    if include_dry:
        parts.append("旱天结果重点反映最大充满度、溢流风险和低流速淤积风险。")
    if include_rainy:
        parts.append("雨天结果反映选定降雨事件下的液位抬升和溢流安全裕度。")
    parts.append("建议结合管径、井深、上下游关系和运维记录，对中高风险点位优先复核和处置。")
    return "".join(parts)


def _join(points: list) -> str:
    clean = [str(p) for p in points if str(p)]
    return "、".join(clean) if clean else "无"


def _numbered_items(items: list[tuple[list, str]]) -> list[str]:
    numbers = "①②③④⑤⑥⑦⑧⑨"
    active = [text for points, text in items if points]
    result: list[str] = []
    for idx, text in enumerate(active):
        body = str(text).strip().rstrip("。；;")
        punctuation = "。" if idx == len(active) - 1 else "；"
        result.append(f"{numbers[idx]}　{body}{punctuation}")
    return result
