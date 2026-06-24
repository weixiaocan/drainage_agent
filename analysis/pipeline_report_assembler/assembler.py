"""Chapter-oriented report assembly orchestration."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Optional

import pandas as pd
from docx import Document

from .data_context import build_report_context
from .facts import build_report_facts
from .llm_section_writer import LLMSectionWriter
from .sections.pattern import render_pattern_section
from .sections.rainfall import render_rainfall_section
from .sections.risk import render_risk_section
from .sections.site_overview import render_site_overview
from .template_scanner import scan_template
from .validator import validate_report


@dataclass
class ReportConfig:
    """Report assembly configuration."""

    monitoring_start: str = ""
    monitoring_end: str = ""
    monitoring_round: str = "第一轮"
    rainfall_threshold_mm: float = 2.0
    baseinfo_path: str = ""


def run_report_assembler(
    template_file: Path,
    analysis_results: Dict[str, pd.DataFrame],
    site_info_file: Path,
    output_file: Path,
    dry_curve_data: Dict[str, pd.DataFrame] | None = None,
    filter_result_path: Path | None = None,
    config: Dict[str, Any] | None = None,
    has_rainfall_data: bool = True,
    llm_client=None,
    sections: list[str] | None = None,
    point_ids: list[str] | None = None,
) -> Dict[str, Any]:
    """Assemble the Word report by template sections."""
    cfg = _build_config(config)
    template_file = Path(template_file)
    site_info_file = Path(site_info_file)
    output_file = Path(output_file)
    output_file.parent.mkdir(parents=True, exist_ok=True)

    print(f"读取报告模板: {template_file}")
    doc = Document(template_file)
    selected = _selected_section_keys(sections)
    _prune_unselected_sections(doc, selected)
    context = build_report_context(
        analysis_results=analysis_results,
        site_info_file=site_info_file,
        dry_curve_data=dry_curve_data,
        has_rainfall_data=has_rainfall_data,
        point_ids=point_ids,
    )
    template_map = scan_template(doc)
    baseinfo_path = Path(cfg.baseinfo_path) if cfg.baseinfo_path else _default_baseinfo_path(template_file)
    facts = build_report_facts(context, baseinfo_path=baseinfo_path)
    if cfg.monitoring_start or cfg.monitoring_end:
        facts.monitoring_period_text = _scope_period_text(cfg.monitoring_start, cfg.monitoring_end)
        facts.operation_period_text = facts.monitoring_period_text
    llm_writer = LLMSectionWriter(llm_client)

    warnings: list[str] = []
    warnings.extend(context.warnings)
    warnings.extend(template_map.warnings)
    stats = {
        "tables_filled": 0,
        "images_inserted": 0,
        "points_processed": facts.point_count,
        "text_replaced": 0,
        "llm_generated": 0,
        "warnings": 0,
    }

    print(f"报告包含 {len(doc.tables)} 个表格")
    print(f"识别点位: {facts.point_ids}")

    renderers = {
        "monitoring_overview": lambda: render_site_overview(doc, template_map, context, facts, warnings),
        "rainfall_analysis": lambda: render_rainfall_section(
            doc, template_map, context, facts, output_file.parent, warnings
        ),
        "dry_pattern_analysis": lambda: render_pattern_section(
            doc, template_map, facts, output_file.parent / "特征曲线图", llm_writer, warnings
        ),
        "operation_risk_analysis": lambda: render_risk_section(
            doc, template_map, context, facts, llm_writer, warnings
        ),
    }
    for key in selected:
        section_stats = renderers[key]()
        _merge_stats(stats, section_stats)

    validation = validate_report(doc, facts, selected_sections=selected)
    warnings.extend(validation.warnings)
    if validation.critical:
        warnings.extend(validation.critical)
        raise ValueError("报告校验失败: " + "；".join(validation.critical))

    doc.save(output_file)
    stats["warnings"] = len(warnings)
    print(f"保存报告: {output_file}")
    _print_warnings(warnings)
    return {"output_file": output_file, "stats": stats, "warnings": warnings}


SECTION_ALIASES = {
    "monitoring_overview": {"监测概况", "数据体检", "数据质量"},
    "rainfall_analysis": {"降雨分析", "降雨统计", "雨天事件统计", "事件响应", "RDII"},
    "dry_pattern_analysis": {"旱天排污规律统计分析", "排污规律", "排污规律分析", "旱天分析"},
    "operation_risk_analysis": {"污水系统运行风险分析", "污水系统运行风险", "风险评估", "旱天风险", "雨天风险", "溢流风险"},
}


def _selected_section_keys(sections: list[str] | None) -> list[str]:
    if not sections:
        return list(SECTION_ALIASES)
    selected = [key for key, aliases in SECTION_ALIASES.items() if any(section in aliases for section in sections)]
    return selected or list(SECTION_ALIASES)


def _prune_unselected_sections(doc: Document, selected: list[str]) -> None:
    starts: list[tuple[str, object]] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for key, aliases in SECTION_ALIASES.items():
            if text in aliases:
                starts.append((key, paragraph._p))
                break
    if not starts:
        return
    body = doc._element.body
    children = list(body)
    positions = [(key, children.index(element)) for key, element in starts if element in children]
    for idx in range(len(positions) - 1, -1, -1):
        key, start = positions[idx]
        if key in selected:
            continue
        end = positions[idx + 1][1] if idx + 1 < len(positions) else len(children) - 1
        for child in children[start:end]:
            if child.getparent() is body:
                body.remove(child)


def _build_config(config: Optional[Dict[str, Any]]) -> ReportConfig:
    cfg = ReportConfig()
    if config:
        for key, value in config.items():
            if hasattr(cfg, key):
                setattr(cfg, key, value)
    return cfg


def _scope_period_text(start: str, end: str) -> str:
    start_ts = pd.to_datetime(start, errors="coerce") if start else None
    end_ts = pd.to_datetime(end, errors="coerce") if end else None
    start_text = start_ts.strftime("%Y/%m/%d") if start_ts is not None and not pd.isna(start_ts) else "不限"
    end_text = end_ts.strftime("%Y/%m/%d") if end_ts is not None and not pd.isna(end_ts) else "不限"
    return f"{start_text}日-{end_text}日"


def _default_baseinfo_path(template_file: Path) -> Path:
    project_root = template_file.parent.parent
    return project_root / "data" / "baseinfo.xlsx"


def _merge_stats(target: dict[str, int], source: dict[str, int]) -> None:
    for key, value in source.items():
        target[key] = target.get(key, 0) + int(value)


def _print_warnings(warnings: list[str]) -> None:
    if not warnings:
        return
    print("报告组装 warnings:")
    for warning in warnings[:20]:
        print(f"  - {warning}")
    if len(warnings) > 20:
        print(f"  - ... 另有 {len(warnings) - 20} 条")
