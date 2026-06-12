from __future__ import annotations

from pathlib import Path

from docx import Document


def build_report(output_file: Path, title: str, summaries: list[str]) -> dict[str, object]:
    document = Document()
    document.add_heading(title, level=0)
    for summary in summaries:
        document.add_paragraph(summary)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_file)
    return {"output_file": str(output_file), "stats": {"paragraphs": len(summaries)}}
