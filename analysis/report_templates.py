from __future__ import annotations

import hashlib
import json
import sqlite3
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

import pandas as pd
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows

from analysis.baselines import FilterBaselineService
from analysis.modules.patterns import analyze_patterns
from analysis.pattern_charts import save_pattern_curve_pngs
from analysis.reporting import build_report
from analysis.schema import to_display_columns


REQUIRED_PLACEHOLDERS = frozenset(
    {
        "{{PROJECT_NAME}}",
        "{{BATCH_NAME}}",
        "{{ANALYSIS_SUMMARY}}",
        "{{MANUAL_TOPOLOGY_SECTION}}",
    }
)


class InvalidReportTemplate(ValueError):
    """Raised when a DOCX does not satisfy the public placeholder contract."""


@dataclass(frozen=True)
class ReportTemplate:
    template_id: str
    project_id: str
    name: str
    artifact: str
    content_sha256: str
    created_at: str


@dataclass(frozen=True)
class ReportDraft:
    report_id: str
    project_id: str
    batch_id: str
    template_id: str
    version: int
    docx: str
    workbook: str
    created_at: str


class ReportTemplateService:
    """Validate templates and create immutable report-draft versions."""

    required_placeholders = REQUIRED_PLACEHOLDERS

    def __init__(
        self,
        database: Path,
        files_root: Path,
        builtin_template: Path,
    ) -> None:
        self.database = Path(database)
        self.files_root = Path(files_root).resolve()
        self.builtin_template = Path(builtin_template)
        try:
            Document(self.builtin_template)
        except Exception as exc:
            raise InvalidReportTemplate("内置报告模板不是有效的 DOCX 文件") from exc
        with self._connect() as connection:
            connection.execute(
                """
                CREATE TABLE IF NOT EXISTS report_templates (
                    template_id TEXT PRIMARY KEY,
                    project_id TEXT NOT NULL,
                    name TEXT NOT NULL,
                    artifact TEXT NOT NULL,
                    content_sha256 TEXT NOT NULL,
                    created_at TEXT NOT NULL
                )
                """
            )
            connection.execute(
                """
                CREATE TABLE IF NOT EXISTS report_drafts (
                    report_id TEXT PRIMARY KEY,
                    project_id TEXT NOT NULL,
                    batch_id TEXT NOT NULL,
                    template_id TEXT NOT NULL,
                    version INTEGER NOT NULL,
                    docx TEXT NOT NULL,
                    workbook TEXT NOT NULL,
                    created_at TEXT NOT NULL,
                    UNIQUE(project_id, batch_id, version)
                )
                """
            )

    def validate_path(self, path: Path) -> None:
        try:
            document = Document(path)
        except Exception as exc:
            raise InvalidReportTemplate("报告模板不是有效的 DOCX 文件") from exc
        text = self._document_text(document)
        missing = sorted(REQUIRED_PLACEHOLDERS - set(self._tokens(text)))
        if missing:
            raise InvalidReportTemplate(
                "报告模板缺少必需占位符: " + ", ".join(missing)
            )

    def upload(
        self,
        project_id: str,
        name: str,
        filename: str,
        content: bytes,
    ) -> ReportTemplate:
        if Path(filename).suffix.lower() != ".docx":
            raise InvalidReportTemplate("自定义报告模板必须为 .docx 文件")
        try:
            document = Document(BytesIO(content))
        except Exception as exc:
            raise InvalidReportTemplate("报告模板不是有效的 DOCX 文件") from exc
        text = self._document_text(document)
        missing = sorted(REQUIRED_PLACEHOLDERS - set(self._tokens(text)))
        if missing:
            raise InvalidReportTemplate(
                "报告模板缺少必需占位符: " + ", ".join(missing)
            )
        self._require_project(project_id)
        template_id = uuid.uuid4().hex
        artifact = f"templates/{template_id}/template.docx"
        target = self._project_root(project_id) / artifact
        target.parent.mkdir(parents=True, exist_ok=False)
        target.write_bytes(content)
        template = ReportTemplate(
            template_id=template_id,
            project_id=project_id,
            name=name.strip() or filename,
            artifact=artifact,
            content_sha256=hashlib.sha256(content).hexdigest(),
            created_at=_now(),
        )
        with self._connect() as connection:
            connection.execute(
                """
                INSERT INTO report_templates (
                    template_id, project_id, name, artifact,
                    content_sha256, created_at
                ) VALUES (?, ?, ?, ?, ?, ?)
                """,
                (
                    template.template_id,
                    template.project_id,
                    template.name,
                    template.artifact,
                    template.content_sha256,
                    template.created_at,
                ),
            )
        return template

    def create_draft(
        self,
        project_id: str,
        batch_id: str,
        template_id: str,
        *,
        points: list[str] | None = None,
        start: str | None = None,
        end: str | None = None,
        sections: list[str] | None = None,
    ) -> ReportDraft:
        project_name, batch_name = self._project_batch(project_id, batch_id)
        template_path = self._template_path(project_id, template_id)
        results = self._current_results(project_id, batch_id)
        if not results:
            raise ValueError("当前分析批次尚无可用于报告的分析结果")
        version = self._next_version(project_id, batch_id)
        report_id = uuid.uuid4().hex
        root = self._batch_root(project_id, batch_id)
        relative_root = f"exports/{version}-{report_id}"
        docx_relative = f"{relative_root}/report_draft.docx"
        workbook_relative = f"{relative_root}/comprehensive_results.xlsx"
        docx_path = root / docx_relative
        workbook_path = root / workbook_relative
        docx_path.parent.mkdir(parents=True, exist_ok=False)

        tables = self._analysis_tables(results)
        summaries = [
            f"{algorithm} 第 {result['version']} 版"
            for algorithm, result in results
        ]
        template_document = Document(template_path)
        if template_id == "builtin" and template_document.tables:
            baseline_flow = FilterBaselineService(
                self.database,
                self.files_root,
            ).load_flow(project_id, batch_id)
            for column in ("flow_lps", "level_m", "velocity_mps"):
                if column in baseline_flow.columns:
                    baseline_flow[column] = pd.to_numeric(
                        baseline_flow[column],
                        errors="coerce",
                    )
            pattern_result = analyze_patterns(baseline_flow)
            curves = pattern_result.get("curves", {})
            pattern_chart_paths = save_pattern_curve_pngs(
                curves if isinstance(curves, dict) else {},
                baseline_flow,
                docx_path.parent,
                "pattern_charts",
            )
            build_report(
                docx_path,
                "排水监测数据分析报告",
                summaries,
                template_file=template_path,
                analysis_tables=tables,
                site_info_file=root / "standard" / "sites.csv",
                outputs_dir=docx_path.parent,
                sections=sections,
                has_rainfall_data=not tables.get("rainfall_daily", pd.DataFrame()).empty,
                point_ids=points,
                start=start,
                end=end,
                dry_curve_data=(
                    curves if isinstance(curves, dict) else {}
                ),
                pattern_chart_paths=pattern_chart_paths,
                artifact_scope=f"报告第{version}版",
            )
        else:
            replacements = {
                "{{PROJECT_NAME}}": project_name,
                "{{BATCH_NAME}}": batch_name,
                "{{ANALYSIS_SUMMARY}}": "；".join(summaries),
                "{{MANUAL_TOPOLOGY_SECTION}}": (
                    "人工补充模块：当前版本不自动生成空间拓扑关系、上下游或管网结构结论。"
                ),
            }
            self._replace(template_document, replacements)
            template_document.save(docx_path)
        self._write_workbook(workbook_path, tables)

        draft = ReportDraft(
            report_id=report_id,
            project_id=project_id,
            batch_id=batch_id,
            template_id=template_id,
            version=version,
            docx=docx_relative,
            workbook=workbook_relative,
            created_at=_now(),
        )
        with self._connect() as connection:
            connection.execute(
                """
                INSERT INTO report_drafts (
                    report_id, project_id, batch_id, template_id, version,
                    docx, workbook, created_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    draft.report_id,
                    draft.project_id,
                    draft.batch_id,
                    draft.template_id,
                    draft.version,
                    draft.docx,
                    draft.workbook,
                    draft.created_at,
                ),
            )
        return draft

    def list_templates(self, project_id: str) -> list[ReportTemplate]:
        with self._connect() as connection:
            rows = connection.execute(
                """
                SELECT template_id, project_id, name, artifact,
                       content_sha256, created_at
                FROM report_templates
                WHERE project_id = ?
                ORDER BY created_at, template_id
                """,
                (project_id,),
            ).fetchall()
        return [ReportTemplate(*row) for row in rows]

    def list_drafts(
        self, project_id: str, batch_id: str
    ) -> list[ReportDraft]:
        with self._connect() as connection:
            rows = connection.execute(
                """
                SELECT report_id, project_id, batch_id, template_id, version,
                       docx, workbook, created_at
                FROM report_drafts
                WHERE project_id = ? AND batch_id = ?
                ORDER BY version
                """,
                (project_id, batch_id),
            ).fetchall()
        return [ReportDraft(*row) for row in rows]

    def _current_results(
        self, project_id: str, batch_id: str
    ) -> list[tuple[str, dict[str, object]]]:
        with self._connect() as connection:
            rows = connection.execute(
                """
                SELECT current.algorithm, runs.result_json
                FROM current_analysis_results AS current
                JOIN analysis_runs AS runs ON runs.run_id = current.run_id
                WHERE current.project_id = ? AND current.batch_id = ?
                ORDER BY current.algorithm
                """,
                (project_id, batch_id),
            ).fetchall()
        return [(row[0], json.loads(row[1])) for row in rows]

    @staticmethod
    def _analysis_tables(
        results: list[tuple[str, dict[str, object]]],
    ) -> dict[str, pd.DataFrame]:
        tables: dict[str, pd.DataFrame] = {}
        for algorithm, result in results:
            data = result.get("data", {})
            if not isinstance(data, dict):
                continue
            if algorithm == "data_quality":
                tables["data_collection"] = pd.DataFrame(data.get("table", []))
            elif algorithm == "patterns":
                tables["pattern_analysis"] = pd.DataFrame(data.get("table", []))
            elif algorithm == "rainfall":
                tables["rainfall_daily"] = pd.DataFrame(data.get("daily", []))
                tables["rainfall_events"] = pd.DataFrame(data.get("events", []))
            elif algorithm == "event_response":
                tables["rainy_event_stats"] = pd.DataFrame(data.get("table", []))
            elif algorithm == "rdii":
                tables["rdii_total"] = pd.DataFrame(data.get("table", []))
            elif algorithm == "risk":
                tables["dry_analysis"] = pd.DataFrame(data.get("dry_analysis", []))
                tables["dry_risk"] = pd.DataFrame(data.get("dry_risk", []))
                tables["rainy_overflow_risk"] = pd.DataFrame(data.get("rainy_risk", []))
        return tables

    @staticmethod
    def _write_workbook(
        path: Path, tables: dict[str, pd.DataFrame]
    ) -> None:
        sheet_specs = {
            "data_collection": ("数据收集率统计", "data_check"),
            "rainfall_daily": ("降雨概况", "rainfall_daily"),
            "rainfall_events": ("降雨场次分析", "rainfall_events"),
            "pattern_analysis": ("排污规律分析", "patterns"),
            "rainy_event_stats": ("雨天事件统计", "event_response"),
            "rdii_total": ("RDII总量统计", "rdii"),
            "dry_analysis": ("旱天分析", "dry_stats"),
            "dry_risk": ("旱天风险", "dry_risk"),
            "rainy_overflow_risk": ("雨天溢流风险", "rainy_risk"),
        }
        workbook = Workbook()
        workbook.remove(workbook.active)
        for key, (sheet_name, table_type) in sheet_specs.items():
            table = tables.get(key)
            if table is None or table.empty:
                continue
            display_table = to_display_columns(table, table_type)
            sheet = workbook.create_sheet(sheet_name)
            for row in dataframe_to_rows(
                display_table,
                index=False,
                header=True,
            ):
                sheet.append(row)
            sheet.freeze_panes = "A2"
            sheet.sheet_view.showGridLines = False
            thin_side = Side(style="thin", color="D9D9D9")
            table_border = Border(
                left=thin_side,
                right=thin_side,
                top=thin_side,
                bottom=thin_side,
            )
            for row in sheet.iter_rows(
                min_row=1,
                max_row=sheet.max_row,
                min_col=1,
                max_col=sheet.max_column,
            ):
                for cell in row:
                    cell.border = table_border
            for cell in sheet[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="1F4E78")
                cell.alignment = Alignment(
                    horizontal="center",
                    vertical="center",
                    wrap_text=True,
                )
            sheet.row_dimensions[1].height = 30
            for column_index, column_name in enumerate(
                display_table.columns,
                start=1,
            ):
                values = [
                    str(value)
                    for value in display_table[column_name].dropna()
                ]
                max_length = max(
                    [len(str(column_name)), *[len(value) for value in values]],
                    default=len(str(column_name)),
                )
                is_description = any(
                    marker in str(column_name)
                    for marker in ("描述", "原因", "时段")
                )
                width = min(
                    60 if is_description else 24,
                    max(12, max_length + 2),
                )
                sheet.column_dimensions[
                    get_column_letter(column_index)
                ].width = width
                for cell in sheet.iter_cols(
                    min_col=column_index,
                    max_col=column_index,
                    min_row=2,
                ):
                    for item in cell:
                        item.alignment = Alignment(
                            vertical="center",
                            wrap_text=is_description,
                        )
        if not workbook.sheetnames:
            workbook.create_sheet("分析结果")
        workbook.save(path)

    def _template_path(self, project_id: str, template_id: str) -> Path:
        if template_id == "builtin":
            return self.builtin_template
        with self._connect() as connection:
            row = connection.execute(
                """
                SELECT artifact FROM report_templates
                WHERE template_id = ? AND project_id = ?
                """,
                (template_id, project_id),
            ).fetchone()
        if row is None:
            raise LookupError("报告模板不存在或不属于当前监测项目")
        return self._project_root(project_id) / row[0]

    def _project_batch(self, project_id: str, batch_id: str) -> tuple[str, str]:
        with self._connect() as connection:
            row = connection.execute(
                """
                SELECT project.name, batch.name
                FROM projects AS project
                JOIN analysis_batches AS batch ON batch.project_id = project.id
                WHERE project.id = ? AND batch.id = ?
                """,
                (project_id, batch_id),
            ).fetchone()
        if row is None:
            raise LookupError("分析批次不存在或不属于当前监测项目")
        return str(row[0]), str(row[1])

    def _require_project(self, project_id: str) -> None:
        with self._connect() as connection:
            row = connection.execute(
                "SELECT 1 FROM projects WHERE id = ?", (project_id,)
            ).fetchone()
        if row is None:
            raise LookupError("监测项目不存在")

    def _next_version(self, project_id: str, batch_id: str) -> int:
        with self._connect() as connection:
            row = connection.execute(
                """
                SELECT COALESCE(MAX(version), 0) + 1 FROM report_drafts
                WHERE project_id = ? AND batch_id = ?
                """,
                (project_id, batch_id),
            ).fetchone()
        return int(row[0])

    def _project_root(self, project_id: str) -> Path:
        root = (self.files_root / project_id).resolve()
        if not root.is_relative_to(self.files_root):
            raise LookupError("项目标识超出项目目录")
        return root

    def _batch_root(self, project_id: str, batch_id: str) -> Path:
        return self._project_root(project_id) / "batches" / batch_id

    @staticmethod
    def _document_text(document: object) -> str:
        paragraphs = [p.text for p in document.paragraphs]
        cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join([*paragraphs, *cells])

    @staticmethod
    def _tokens(text: str) -> list[str]:
        return [token for token in REQUIRED_PLACEHOLDERS if token in text]

    @staticmethod
    def _replace(document: object, replacements: dict[str, str]) -> None:
        paragraphs = list(document.paragraphs)
        paragraphs.extend(
            paragraph
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
        )
        for paragraph in paragraphs:
            text = paragraph.text
            for token, value in replacements.items():
                text = text.replace(token, value)
            if text != paragraph.text:
                paragraph.text = text

    def _connect(self) -> sqlite3.Connection:
        return sqlite3.connect(self.database)


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()
