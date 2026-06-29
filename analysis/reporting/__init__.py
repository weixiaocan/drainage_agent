from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class ReportModule:
    key: str
    title: str
    aliases: tuple[str, ...]
    keywords: tuple[str, ...]
    sheets: tuple[str, ...]


MODULES: tuple[ReportModule, ...] = (
    ReportModule(
        "monitoring_overview",
        "监测概况",
        ("监测设备安装", "数据收集情况", "数据收集率", "数据体检"),
        ("监测概况", "监测设备安装", "数据收集率", "数据获取情况"),
        ("监测点位安装信息汇总", "数据收集率统计", "数据体检"),
    ),
    ReportModule(
        "rainfall_analysis",
        "降雨分析",
        ("降雨统计", "场次降雨", "雨天事件统计", "雨天溢流风险", "事件响应", "RDII"),
        ("降雨分析", "降雨统计", "场次降雨", "雨天运行风险"),
        ("降雨概况", "降雨场次分析", "日降雨量统计", "场次降雨统计", "雨天事件统计", "RDII总量统计", "雨天溢流风险"),
    ),
    ReportModule(
        "dry_pattern_analysis",
        "旱天排污规律统计分析",
        ("排污规律", "排污规律分析", "旱天流量特征曲线", "旱天分析"),
        ("旱天排污规律统计分析", "排污规律", "旱天流量特征曲线"),
        ("旱天分析", "排污规律分析"),
    ),
    ReportModule(
        "operation_risk_analysis",
        "污水系统运行风险分析",
        ("运行风险分析", "风险评估", "旱天风险", "雨天风险", "溢流风险"),
        ("污水系统运行风险", "运行风险分析", "溢流风险", "淤积风险"),
        ("旱天风险", "雨天溢流风险"),
    ),
)


def _module_by_name(name: str) -> ReportModule | None:
    normalized = name.strip().lower()
    for module in MODULES:
        names = {module.key.lower(), module.title.lower(), *(alias.lower() for alias in module.aliases)}
        if normalized in names:
            return module
    return None


def _selected_modules(sections: Iterable[str] | None) -> list[ReportModule]:
    if not sections:
        return list(MODULES)
    selected: list[ReportModule] = []
    seen: set[str] = set()
    for section in sections:
        module = _module_by_name(section)
        if module and module.key not in seen:
            selected.append(module)
            seen.add(module.key)
    return selected or list(MODULES)


def _find_template_anchor(document: DocumentObject, module: ReportModule) -> Paragraph | None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if any(keyword in text for keyword in module.keywords):
            return paragraph
    return None


def _move_block_after(document: DocumentObject, block, anchor):
    parent = block.getparent()
    if parent is not None:
        parent.remove(block)
    anchor.addnext(block)
    return block


def _add_paragraph_after(document: DocumentObject, anchor, text: str, style: str | None = None):
    paragraph = document.add_paragraph(text)
    if style:
        try:
            paragraph.style = style
        except Exception:
            pass
    return _move_block_after(document, paragraph._p, anchor)


def _add_table_after(document: DocumentObject, anchor, table_df: pd.DataFrame, max_rows: int = 12):
    table_df = table_df.copy()
    if len(table_df) > max_rows:
        table_df = table_df.head(max_rows)
    table = document.add_table(rows=1, cols=max(1, len(table_df.columns)))
    table.style = "Table Grid"
    headers = list(table_df.columns) or ["结果"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = str(header)
    for _, row in table_df.iterrows():
        cells = table.add_row().cells
        for idx, header in enumerate(headers):
            value = row.get(header, "")
            cells[idx].text = "" if pd.isna(value) else str(value)
    return _move_block_after(document, table._tbl, anchor)


def _table_text(table: Table) -> str:
    return "\n".join(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)


def _row_text(table: Table, row_idx: int) -> str:
    if row_idx >= len(table.rows):
        return ""
    return " | ".join(cell.text.strip() for cell in table.rows[row_idx].cells)


def _has_all(text: str, *needles: str) -> bool:
    return all(needle in text for needle in needles)


def _scan_template_tables(document: DocumentObject) -> dict[str, Table]:
    roles: dict[str, Table] = {}
    for table in document.tables:
        text = _table_text(table)
        header0 = _row_text(table, 0)
        header1 = _row_text(table, 1)
        role = ""
        if _has_all(header0, "监测点位", "设备类型", "管径", "井深"):
            role = "site_info"
        elif _has_all(header0, "理论数据条数", "收集率") and ("监测数据条数" in header0 or "记录数" in header0):
            role = "collection_rate"
        elif _has_all(header0, "日期", "日降雨量"):
            role = "rainfall_daily"
        elif _has_all(header0, "开始时间", "结束时间") and "总降雨量" in header0:
            role = "rainfall_events"
        elif _has_all(text, "最大充满度", "淤积风险", "运行风险"):
            role = "dry_risk"
        elif _has_all(text, "最大液位", "溢流风险值", "溢流风险"):
            role = "rainy_overflow_risk"
        if role and role not in roles:
            roles[role] = table
    return roles


def _set_cell_text(cell, text: object) -> None:
    value = "" if text is None or (not isinstance(text, str) and pd.isna(text)) else str(text)
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)
    for extra in cell.paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = str(text)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(str(text))


def _replace_first_paragraph(document: DocumentObject, keyword: str, text: str) -> bool:
    for paragraph in document.paragraphs:
        if keyword in paragraph.text:
            _set_paragraph_text(paragraph, text)
            return True
    return False


def _clone_table_row(table: Table, template_row_idx: int):
    if not table.rows:
        raise ValueError("Cannot clone row from an empty table")
    template_row_idx = min(template_row_idx, len(table.rows) - 1)
    new_row = deepcopy(table.rows[template_row_idx]._tr)
    table._tbl.append(new_row)
    return table.rows[-1]


def _adjust_table_rows(table: Table, target_data_rows: int, template_row_idx: int = 1) -> None:
    current_data_rows = max(0, len(table.rows) - template_row_idx)
    if current_data_rows > target_data_rows:
        for _ in range(current_data_rows - target_data_rows):
            row = table.rows[-1]._element
            row.getparent().remove(row)
    elif current_data_rows < target_data_rows:
        for _ in range(target_data_rows - current_data_rows):
            _clone_table_row(table, template_row_idx)


def _format_number(value: object) -> str:
    if value is None or pd.isna(value):
        return ""
    try:
        return f"{float(value):.2f}".rstrip("0").rstrip(".")
    except (TypeError, ValueError):
        return str(value)


def _format_int(value: object) -> str:
    if value is None or pd.isna(value):
        return ""
    try:
        return str(int(round(float(value))))
    except (TypeError, ValueError):
        return str(value)


def _format_percent(value: object) -> str:
    if value is None or pd.isna(value):
        return ""
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if number <= 1:
        number *= 100
    return f"{number:.1f}%"


def _format_date(value: object) -> str:
    if value is None or pd.isna(value):
        return ""
    parsed = pd.to_datetime(value, errors="coerce")
    if pd.isna(parsed):
        return str(value)[:10]
    return parsed.strftime("%Y-%m-%d")


def _format_datetime(value: object) -> str:
    if value is None or pd.isna(value):
        return ""
    parsed = pd.to_datetime(value, errors="coerce")
    if pd.isna(parsed):
        return str(value)[:16]
    return parsed.strftime("%Y-%m-%d %H:%M")


def _find_source_column(df: pd.DataFrame, aliases: tuple[str, ...]) -> object | None:
    columns = {str(col).strip(): col for col in df.columns}
    for alias in aliases:
        if alias in columns:
            return columns[alias]
    for alias in aliases:
        for text, col in columns.items():
            if alias in text:
                return col
    return None


def _project_table(df: pd.DataFrame, columns: tuple[tuple[str, tuple[str, ...], object], ...]) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame(columns=[output for output, _, _ in columns])
    rows: dict[str, pd.Series] = {}
    for output, aliases, formatter in columns:
        source = _find_source_column(df, aliases)
        series = df[source] if source is not None else pd.Series([""] * len(df), index=df.index)
        rows[output] = series.map(formatter)
    return pd.DataFrame(rows)


def _load_site_info_table(site_info_file: Path | None) -> pd.DataFrame:
    if not site_info_file or not site_info_file.exists():
        return pd.DataFrame()
    raw = pd.read_excel(site_info_file)
    columns = (
        ("监测点位", ("监测点位", "点位编号", "安装点位", "安装监测点位"), str),
        ("设备类型", ("设备类型", "类型"), str),
        ("形状", ("形状", "管道形状", "管型", "绑定管形状"), str),
        ("管径(m)", ("管径(m)", "管径", "管径（m）"), _format_number),
        ("井深(m)", ("井深(m)", "井深", "井深（m）"), _format_number),
        ("设备安装时间", ("设备安装时间", "安装时间"), _format_date),
    )
    return _project_table(raw, columns)


def _template_role_df(role: str, tables: dict[str, pd.DataFrame], site_info_file: Path | None) -> tuple[pd.DataFrame, int]:
    if role == "site_info":
        return _load_site_info_table(site_info_file), 1
    if role == "collection_rate":
        source = tables.get("数据收集率统计", tables.get("数据体检", pd.DataFrame()))
        columns = (
            ("点位编号", ("点位编号", "point_id"), str),
            ("监测数据条数", ("监测数据条数", "记录数", "record_count"), _format_int),
            ("监测天数", ("监测天数", "monitoring_days"), _format_int),
            ("理论数据条数", ("理论数据条数", "theoretical_count"), _format_int),
            ("数据收集率", ("数据收集率", "收集率", "collection_rate"), _format_percent),
        )
        return _project_table(source, columns), 1
    if role == "rainfall_daily":
        source = tables.get("降雨概况", tables.get("日降雨量统计", pd.DataFrame()))
        columns = (
            ("日期", ("日期", "date"), _format_date),
            ("日降雨量(mm)", ("日降雨量(mm)", "日降雨量", "rain_mm"), _format_number),
        )
        result = _project_table(source, columns)
        if "日降雨量(mm)" in result.columns:
            rain = pd.to_numeric(result["日降雨量(mm)"], errors="coerce").fillna(0)
            result = result[rain > 0].copy()
        return result, 1
    if role == "rainfall_events":
        source = tables.get("降雨场次分析", tables.get("场次降雨统计", pd.DataFrame()))
        columns = (
            ("场次编号", ("场次编号", "event_id"), _format_int),
            ("开始时间", ("开始时间", "start_time"), _format_datetime),
            ("结束时间", ("结束时间", "end_time"), _format_datetime),
            ("总降雨量(mm)", ("总降雨量(mm)", "total_rain_mm"), _format_number),
            ("降雨历时(h)", ("降雨历时(h)", "duration_h"), _format_number),
            ("平均强度(mm/h)", ("平均强度(mm/h)", "avg_intensity_mmh"), _format_number),
            ("降雨等级", ("降雨等级", "rain_level"), str),
        )
        return _project_table(source, columns), 1
    if role == "dry_risk":
        source = tables.get("旱天风险", pd.DataFrame())
        columns = (
            ("序号", ("序号", "serial_no"), _format_int),
            ("点位编号", ("点位编号", "point_id"), str),
            ("管径(m)", ("管径(m)", "diameter_m"), _format_number),
            ("井深(m)", ("井深(m)", "well_depth_m"), _format_number),
            ("旱天流速(m/s)", ("旱天流速(m/s)", "dry_velocity_mps"), _format_number),
            ("最大液位(m)", ("最大液位(m)", "max_level_m"), _format_number),
            ("最大充满度", ("最大充满度", "max_fullness"), _format_number),
            ("溢流风险值", ("溢流风险值", "overflow_value"), _format_number),
            ("淤积风险", ("淤积风险", "silting_risk"), str),
            ("运行风险", ("运行风险", "running_risk"), str),
            ("溢流风险", ("溢流风险", "overflow_risk"), str),
        )
        result = _project_table(source, columns)
        if "序号" in result.columns and result["序号"].replace("", pd.NA).isna().all():
            result["序号"] = range(1, len(result) + 1)
        return result, 2
    if role == "rainy_overflow_risk":
        source = tables.get("雨天溢流风险", pd.DataFrame())
        columns = (
            ("点位编号", ("点位编号", "point_id"), str),
            ("最大液位(m)", ("最大液位(m)", "max_level_m"), _format_number),
            ("井深(m)", ("井深(m)", "well_depth_m"), _format_number),
            ("溢流风险值", ("溢流风险值", "overflow_value"), _format_number),
            ("溢流风险", ("溢流风险", "overflow_risk"), str),
        )
        return _project_table(source, columns), 1
    return pd.DataFrame(), 1


def _render_template_table(table: Table, df: pd.DataFrame, template_row_idx: int) -> None:
    _adjust_table_rows(table, len(df), template_row_idx=template_row_idx)
    if df.empty:
        return
    for row_offset, (_, row) in enumerate(df.iterrows()):
        word_row_idx = template_row_idx + row_offset
        if word_row_idx >= len(table.rows):
            break
        cells = table.rows[word_row_idx].cells
        for col_idx, column in enumerate(df.columns):
            if col_idx >= len(cells):
                break
            _set_cell_text(cells[col_idx], row.get(column, ""))


def _fill_template_tables(document: DocumentObject, tables: dict[str, pd.DataFrame], site_info_file: Path | None) -> tuple[int, set[str]]:
    roles = _scan_template_tables(document)
    role_sheets = {
        "site_info": {"监测点位安装信息汇总"},
        "collection_rate": {"数据收集率统计", "数据体检"},
        "rainfall_daily": {"降雨概况", "日降雨量统计"},
        "rainfall_events": {"降雨场次分析", "场次降雨统计"},
        "dry_risk": {"旱天风险"},
        "rainy_overflow_risk": {"雨天溢流风险"},
    }
    filled = 0
    filled_sheets: set[str] = set()
    for role, table in roles.items():
        df, template_row_idx = _template_role_df(role, tables, site_info_file)
        if df.empty:
            continue
        _render_template_table(table, df, template_row_idx)
        filled += 1
        filled_sheets.update(role_sheets.get(role, set()))
    return filled, filled_sheets


def _series(df: pd.DataFrame, aliases: tuple[str, ...]) -> pd.Series:
    source = _find_source_column(df, aliases)
    if source is None:
        return pd.Series(dtype=object)
    return df[source]


def _monitoring_period_text(data_check: pd.DataFrame) -> str:
    starts = pd.to_datetime(_series(data_check, ("开始时间", "start_time")), errors="coerce").dropna()
    ends = pd.to_datetime(_series(data_check, ("结束时间", "end_time")), errors="coerce").dropna()
    if starts.empty or ends.empty:
        days = pd.to_numeric(_series(data_check, ("监测天数", "monitoring_days")), errors="coerce").dropna()
        if days.empty:
            return "本轮监测期"
        return f"{int(days.max())}天监测期"
    return f"{starts.min().strftime('%Y-%m-%d')}至{ends.max().strftime('%Y-%m-%d')}"


def _replace_monitoring_text(document: DocumentObject, tables: dict[str, pd.DataFrame]) -> int:
    data_check = tables.get("数据收集率统计", tables.get("数据体检", pd.DataFrame()))
    if data_check.empty:
        return 0
    point_count = len(data_check)
    period = _monitoring_period_text(data_check)
    records = pd.to_numeric(_series(data_check, ("记录数", "监测数据条数", "record_count")), errors="coerce").fillna(0)
    total_records = int(records.sum())
    rates = pd.to_numeric(_series(data_check, ("数据收集率", "收集率", "collection_rate")), errors="coerce").dropna()
    if not rates.empty and rates.max() <= 1:
        rates = rates * 100
    replaced = 0
    replaced += int(_replace_first_paragraph(document, "本轮共布设", f"本轮共布设{point_count}个流量监测点位，时间段选择{period}。"))
    replaced += int(
        _replace_first_paragraph(
            document,
            "期间对设备持续进行运维",
            f"{period}期间对设备持续进行运维，{point_count}个监测点位在监测期间运行状态良好，共收集分钟级监测数据{total_records}条，具体每台设备的点位获取情况如下表所示。",
        )
    )
    if not rates.empty:
        full_count = int((rates >= 99.9).sum())
        replaced += int(
            _replace_first_paragraph(
                document,
                "有效数据收集率",
                f"有效数据收集率范围为{rates.min():.1f}%-{rates.max():.1f}%，其中{full_count}个监测点位的数据收集率达到99.9%及以上，可支撑后续分析工作。",
            )
        )
    return replaced


def _replace_rainfall_text(document: DocumentObject, tables: dict[str, pd.DataFrame]) -> int:
    daily = tables.get("降雨概况", tables.get("日降雨量统计", pd.DataFrame()))
    events = tables.get("降雨场次分析", tables.get("场次降雨统计", pd.DataFrame()))
    if daily.empty and events.empty:
        return 0
    rain = pd.to_numeric(_series(daily, ("日降雨量(mm)", "日降雨量", "rain_mm")), errors="coerce").fillna(0)
    dates = pd.to_datetime(_series(daily, ("日期", "date")), errors="coerce")
    rainy_days = int((rain > 0).sum())
    total_days = int(len(daily))
    total_rain = float(rain.sum())
    max_rain = float(rain.max()) if len(rain) else 0.0
    max_date = ""
    if len(rain) and not dates.empty:
        max_idx = rain.idxmax()
        if max_idx in dates.index and not pd.isna(dates.loc[max_idx]):
            max_date = dates.loc[max_idx].strftime("%Y-%m-%d")
    event_rain = pd.to_numeric(_series(events, ("总降雨量(mm)", "total_rain_mm")), errors="coerce").fillna(0)
    replaced = 0
    replaced += int(
        _replace_first_paragraph(
            document,
            "雨量计持续监测降雨数据",
            f"监测期间，雨量计持续监测降雨数据。以天（24h）为统计单位，降雨日天数为{rainy_days}天，总降雨量{total_rain:.1f} mm，日最大降雨量为{max_rain:.1f} mm{f'，发生在{max_date}' if max_date else ''}。降雨日各天的日降雨量统计如下表所示：",
        )
    )
    if total_days:
        rainy_pct = round(rainy_days / total_days * 100)
        replaced += int(
            _replace_first_paragraph(
                document,
                "非降雨日为",
                f"监测期内共{total_days}个自然日，其中降雨日为{rainy_days}天，非降雨日为{total_days - rainy_days}天。降雨日占整个监测期内自然日的{rainy_pct}%，非降雨日占{100 - rainy_pct}%。",
            )
        )
    if not events.empty:
        replaced += int(
            _replace_first_paragraph(
                document,
                "累计降雨量",
                f"考虑降雨发生后会引起雨水径流现象，以下雨发生到结束后12小时为一个降雨场次。监测期内共发生有效降雨场次{len(events)}场，累计降雨量{float(event_rain.sum()):.1f} mm，最大场次降雨量为{float(event_rain.max()):.1f} mm。各降雨场次统计如下表所示：",
            )
        )
    return replaced


def _replace_risk_text(document: DocumentObject, tables: dict[str, pd.DataFrame]) -> int:
    dry = tables.get("旱天风险", pd.DataFrame())
    if dry.empty:
        return 0
    blocks = _risk_text_blocks(dry, tables.get("雨天溢流风险", pd.DataFrame()))
    replaced = 0
    replaced += _replace_risk_block(
        document,
        ("旱天最大充满度", "最大充满度情况", "最大充满度"),
        blocks["fullness"],
    )
    replaced += _replace_risk_block(
        document,
        ("旱天溢流风险值", "溢流风险值情况", "溢流风险值"),
        blocks["overflow"],
    )
    replaced += _replace_risk_block(
        document,
        ("淤积风险情况", "淤积风险"),
        blocks["silting"],
    )
    if blocks["rainy"]:
        replaced += _replace_risk_block(
            document,
            ("雨天运行风险分析", "雨天溢流风险", "雨天运行风险"),
            blocks["rainy"],
        )
    replaced += _replace_risk_block(
        document,
        ("本章从最大充满度", "本章小结"),
        blocks["summary"],
    )
    if replaced == 0:
        ordered_blocks = [blocks["fullness"], blocks["overflow"], blocks["silting"], blocks["rainy"], blocks["summary"]]
        replaced = _insert_risk_blocks(document, ordered_blocks)
    return replaced


def _risk_text_blocks(dry: pd.DataFrame, rainy: pd.DataFrame) -> dict[str, list[str]]:
    point_count = len(dry)
    fullness = pd.to_numeric(_series(dry, ("最大充满度", "max_fullness")), errors="coerce")
    overflow = pd.to_numeric(_series(dry, ("溢流风险值", "overflow_value")), errors="coerce")
    velocity = pd.to_numeric(_series(dry, ("旱天流速(m/s)", "dry_velocity_mps")), errors="coerce")
    point_ids = _series(dry, ("点位编号", "point_id")).astype(str)

    def points(mask: pd.Series) -> str:
        values = point_ids[mask.fillna(False)].dropna().tolist()
        return "、".join(values) if values else "无"

    blocks: dict[str, list[str]] = {"fullness": [], "overflow": [], "silting": [], "rainy": [], "summary": []}
    if not fullness.dropna().empty:
        low = fullness < 0.75
        mid_low = (fullness >= 0.75) & (fullness < 1.0)
        mid = (fullness >= 1.0) & (fullness <= 2.0)
        high = fullness > 2.0
        blocks["fullness"] = [
            f"监测期间，{point_count}处监测点的旱天最大充满度情况如下：",
            f"① {int(low.sum())}处监测点最大充满度小于0.75，运行良好，点位为{points(low)}；",
            f"② {int(mid_low.sum())}处监测点最大充满度为0.75-1.0，存在运行低风险，点位为{points(mid_low)}；",
            f"③ {int(mid.sum())}处监测点最大充满度为1.0-2.0，存在运行中风险，点位为{points(mid)}；",
            f"④ {int(high.sum())}处监测点最大充满度大于2.0，存在运行高风险，点位为{points(high)}。",
        ]
    if not overflow.dropna().empty:
        low = overflow < 0.7
        mid = (overflow >= 0.7) & (overflow < 0.9)
        high = (overflow >= 0.9) & (overflow <= 1.0)
        overflowed = overflow > 1.0
        blocks["overflow"] = [
            f"监测期间，{point_count}处监测点的旱天溢流风险值情况如下：",
            f"① {int(low.sum())}处监测点溢流风险值小于0.7，溢流风险低，点位为{points(low)}；",
            f"② {int(mid.sum())}处监测点溢流风险值为0.7-0.9，为溢流中风险，点位为{points(mid)}；",
            f"③ {int(high.sum())}处监测点溢流风险值为0.9-1.0，为溢流高风险，点位为{points(high)}；",
            f"④ {int(overflowed.sum())}处监测点溢流风险值大于1.0，已发生溢流，点位为{points(overflowed)}。",
        ]
    if not velocity.dropna().empty:
        high = velocity <= 0.3
        mid = (velocity > 0.3) & (velocity < 0.6)
        low = velocity >= 0.6
        blocks["silting"] = [
            f"监测期间，{point_count}处监测点的淤积风险情况如下：",
            f"① {int(high.sum())}处监测点平均流速小于等于0.3m/s，为高淤积风险，点位为{points(high)}；",
            f"② {int(mid.sum())}处监测点平均流速为0.3-0.6m/s，为中淤积风险，点位为{points(mid)}；",
            f"③ {int(low.sum())}处监测点平均流速大于等于0.6m/s，为低淤积风险，点位为{points(low)}。",
        ]

    if not rainy.empty:
        rainy_point_ids = _series(rainy, ("点位编号", "point_id")).astype(str)
        rainy_overflow = pd.to_numeric(_series(rainy, ("溢流风险值", "overflow_value")), errors="coerce")

        def rainy_points(mask: pd.Series) -> str:
            values = rainy_point_ids[mask.fillna(False)].dropna().tolist()
            return "、".join(values) if values else "无"

        low = rainy_overflow < 0.7
        mid = (rainy_overflow >= 0.7) & (rainy_overflow < 0.9)
        high = (rainy_overflow >= 0.9) & (rainy_overflow <= 1.0)
        overflowed = rainy_overflow > 1.0
        blocks["rainy"] = [
            f"监测期间，{len(rainy)}处监测点位在选定降雨事件下的溢流风险情况如下：",
            f"① {int(low.sum())}处监测点溢流风险值小于0.7，溢流风险低，点位为{rainy_points(low)}；",
            f"② {int(mid.sum())}处监测点溢流风险值为0.7-0.9，为溢流中风险，点位为{rainy_points(mid)}；",
            f"③ {int(high.sum())}处监测点溢流风险值为0.9-1.0，为溢流高风险，点位为{rainy_points(high)}；",
            f"④ {int(overflowed.sum())}处监测点溢流风险值大于1.0，已发生溢流，点位为{rainy_points(overflowed)}。",
        ]

    concern = []
    if not fullness.dropna().empty:
        concern.extend(point_ids[(fullness >= 1.0).fillna(False)].tolist())
    if not overflow.dropna().empty:
        concern.extend(point_ids[(overflow >= 0.7).fillna(False)].tolist())
    if not velocity.dropna().empty:
        concern.extend(point_ids[(velocity < 0.6).fillna(False)].tolist())
    concern_points = "、".join(dict.fromkeys(str(point) for point in concern if str(point))) or "无"
    blocks["summary"] = [
        f"本章小结：本章从最大充满度、溢流风险和淤积风险三个维度，对{point_count}个监测点位的污水系统运行风险进行了评估。",
        f"监测结果显示，需重点关注的点位包括{concern_points}。建议结合管径、井深、上下游关系和运维记录进一步复核，并优先对中高风险点位开展清淤、排水能力复核和雨天调度优化。",
    ]
    return blocks


def _replace_risk_block(document: DocumentObject, keywords: tuple[str, ...], lines: list[str]) -> int:
    if not lines:
        return 0
    for idx, paragraph in enumerate(document.paragraphs):
        if any(keyword in paragraph.text for keyword in keywords):
            _set_paragraph_text(paragraph, lines[0])
            anchor = paragraph._p
            for line in lines[1:]:
                anchor = _new_paragraph_after(document, anchor, line)
            _clear_following_risk_placeholders(document, idx + len(lines), max_count=4)
            return len(lines)
    return 0


def _insert_risk_blocks(document: DocumentObject, blocks: list[list[str]]) -> int:
    start_idx = _find_paragraph_index(document, "污水系统运行风险")
    if start_idx < 0:
        start_idx = _find_paragraph_index(document, "运行风险")
    if start_idx < 0:
        return 0
    anchor = document.paragraphs[start_idx]._p
    count = 0
    for lines in blocks:
        for line in lines:
            anchor = _new_paragraph_after(document, anchor, line)
            count += 1
    return count


def _clear_following_risk_placeholders(document: DocumentObject, start_idx: int, max_count: int) -> None:
    cleared = 0
    for paragraph in document.paragraphs[start_idx:]:
        text = paragraph.text.strip()
        if not text:
            continue
        if text.startswith("表") or text.startswith("图") or "本章小结" in text:
            break
        if any(keyword in text for keyword in ("风险", "点位", "充满度", "流速", "溢流", "淤积", "2024", "19处")):
            _set_paragraph_text(paragraph, "")
            cleared += 1
        if cleared >= max_count:
            break


def _replace_template_text(document: DocumentObject, tables: dict[str, pd.DataFrame]) -> int:
    return (
        _replace_monitoring_text(document, tables)
        + _replace_rainfall_text(document, tables)
        + _replace_risk_text(document, tables)
    )


def _find_paragraph_index(document: DocumentObject, keyword: str, start: int = 0) -> int:
    for idx, paragraph in enumerate(document.paragraphs[start:], start=start):
        if keyword in paragraph.text:
            return idx
    return -1


def _delete_blocks_between(start_element, end_element) -> None:
    parent = start_element.getparent()
    deleting = False
    for child in list(parent):
        if child is start_element:
            deleting = True
            continue
        if child is end_element:
            break
        if deleting:
            parent.remove(child)


def _new_paragraph_after(document: DocumentObject, anchor, text: str = "", style: str | None = None):
    paragraph = document.add_paragraph(text)
    if style:
        try:
            paragraph.style = style
        except Exception:
            pass
    return _move_block_after(document, paragraph._p, anchor)


def _pattern_category_name(value: object) -> str:
    text = str(value or "")
    if "1" in text or "一" in text or "符合" in text:
        return "第1类，符合生活用水规律"
    if "2" in text or "二" in text or "不符合典型" in text or "有波峰" in text:
        return "第2类，有波峰或波谷但不符合典型生活用水规律"
    return "第3类，曲线无明显波峰或波谷"


def _pattern_class_id(row: pd.Series) -> int:
    category_col = _find_source_column(pd.DataFrame([row]), ("分类", "category"))
    if category_col is not None:
        try:
            value = int(row.get(category_col))
            if value in {1, 2, 3}:
                return value
        except (TypeError, ValueError):
            pass
    text = str(row.get(_find_source_column(pd.DataFrame([row]), ("分类名称", "规律类别", "category_name")) or "", ""))
    if "1" in text or "一" in text or "符合" in text:
        return 1
    if "2" in text or "二" in text or "不符合典型" in text or "有波峰" in text:
        return 2
    return 3


def _pattern_point_id(row: pd.Series) -> str:
    col = _find_source_column(pd.DataFrame([row]), ("点位编号", "point_id"))
    return str(row.get(col, "")).strip() if col is not None else ""


def _pattern_description(row: pd.Series) -> str:
    col = _find_source_column(pd.DataFrame([row]), ("排污规律描述", "description"))
    value = str(row.get(col, "")).strip() if col is not None else ""
    if value:
        return value
    point_id = _pattern_point_id(row)
    kz_col = _find_source_column(pd.DataFrame([row]), ("Kz值", "kz"))
    peak_col = _find_source_column(pd.DataFrame([row]), ("波峰时段", "peak_periods"))
    reason_col = _find_source_column(pd.DataFrame([row]), ("诊断理由", "diagnosis_reason"))
    parts = [f"{point_id}点位属于{_pattern_category_name(row.to_dict())}"]
    if kz_col is not None and str(row.get(kz_col, "")).strip():
        parts.append(f"Kz值为{row.get(kz_col)}")
    if peak_col is not None and str(row.get(peak_col, "")).strip():
        parts.append(f"相对高流量时段为{row.get(peak_col)}")
    if reason_col is not None and str(row.get(reason_col, "")).strip():
        parts.append(str(row.get(reason_col)))
    return "，".join(parts) + "。"


def _pattern_image_path(image_paths: dict[str, list[str]] | None, point_id: str) -> Path | None:
    for path_text in (image_paths or {}).get(point_id, []):
        path = Path(path_text)
        if path.exists() and "流量" in path.name:
            return path
    for path_text in (image_paths or {}).get("selected", []):
        path = Path(path_text)
        if path.exists():
            return path
    return None


def _render_pattern_section(
    document: DocumentObject,
    tables: dict[str, pd.DataFrame],
    pattern_chart_paths: dict[str, list[str]] | None,
) -> tuple[int, int]:
    patterns = tables.get("排污规律分析", pd.DataFrame())
    if patterns.empty:
        return 0, 0
    start_idx = _find_paragraph_index(document, "旱天排污规律统计分析")
    if start_idx < 0:
        start_idx = _find_paragraph_index(document, "排污规律")
    end_idx = _find_paragraph_index(document, "污水系统运行风险", start=start_idx + 1 if start_idx >= 0 else 0)
    if start_idx < 0 or end_idx < 0:
        return 0, 0

    start_element = document.paragraphs[start_idx]._p
    end_element = document.paragraphs[end_idx]._p
    _delete_blocks_between(start_element, end_element)

    rows_by_class: dict[int, list[pd.Series]] = {1: [], 2: [], 3: []}
    for _, row in patterns.iterrows():
        rows_by_class[_pattern_class_id(row)].append(row)

    anchor = start_element
    total_points = sum(len(rows) for rows in rows_by_class.values())
    summary_parts = [
        f"第{class_id}类{_pattern_category_name(class_id)}的点位有{len(rows)}处"
        for class_id, rows in rows_by_class.items()
        if rows
    ]
    anchor = _new_paragraph_after(document, anchor, f"本轮监测的{total_points}个点位根据旱天流量特征曲线形态可分为三类，{'；'.join(summary_parts)}。")
    inserted_images = 0
    rainfall_caption_images = 0
    text_blocks = 1

    figure_no = 1
    for class_id in (1, 2, 3):
        rows = rows_by_class[class_id]
        if not rows:
            continue
        anchor = _new_paragraph_after(document, anchor, f"（{class_id}）{_pattern_category_name(class_id)}")
        anchor = _new_paragraph_after(document, anchor, f"该类共包括{len(rows)}个点位。各点位排污规律具体情况如下。")
        text_blocks += 2
        for row in rows:
            point_id = _pattern_point_id(row)
            anchor = _new_paragraph_after(document, anchor, _pattern_description(row))
            text_blocks += 1
            image_path = _pattern_image_path(pattern_chart_paths, point_id)
            if image_path:
                anchor = _add_picture_after(document, anchor, image_path, width_inches=5.6)
                inserted_images += 1
                anchor = _new_paragraph_after(document, anchor, f"图 {figure_no} {point_id}流量特征曲线图")
                figure_no += 1
                text_blocks += 1

    focus = rows_by_class[2] + rows_by_class[3]
    focus_points = "、".join(_pattern_point_id(row) for row in focus if _pattern_point_id(row)) or "无"
    anchor = _new_paragraph_after(
        document,
        anchor,
        f"本章小结：本章对{total_points}个监测点位的旱天排污规律进行了统计分析。后续运行诊断中应重点关注第2类和第3类点位，重点点位包括{focus_points}。",
    )
    text_blocks += 1
    return text_blocks, inserted_images


def _add_picture_after(document: DocumentObject, anchor, image_path: Path, width_inches: float = 5.8):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    return _move_block_after(document, paragraph._p, anchor)


def _table_for(module: ReportModule, tables: dict[str, pd.DataFrame]) -> list[tuple[str, pd.DataFrame]]:
    result: list[tuple[str, pd.DataFrame]] = []
    for sheet in module.sheets:
        table = tables.get(sheet)
        if table is not None and not table.empty:
            result.append((sheet, table))
    return result


def _numeric_sum(df: pd.DataFrame, candidates: Iterable[str]) -> float | None:
    for name in candidates:
        if name in df.columns:
            values = pd.to_numeric(df[name], errors="coerce").dropna()
            if not values.empty:
                return float(values.sum())
    return None


def _numeric_mean(df: pd.DataFrame, candidates: Iterable[str]) -> float | None:
    for name in candidates:
        if name in df.columns:
            values = pd.to_numeric(df[name], errors="coerce").dropna()
            if not values.empty:
                return float(values.mean())
    return None


def _module_summary(module: ReportModule, module_tables: list[tuple[str, pd.DataFrame]], fallback_summaries: list[str]) -> str:
    if not module_tables:
        return f"本节尚未在综合结果表中找到对应模块数据，报告仅保留模板结构。"

    tables = {name: table for name, table in module_tables}
    if module.key == "monitoring_overview":
        table = tables.get("数据收集率统计", tables.get("数据体检", module_tables[0][1]))
        avg_rate = _numeric_mean(table, ("数据收集率", "collection_rate"))
        if avg_rate is not None:
            return f"本轮监测共统计 {len(table)} 个监测点位的数据获取情况，平均数据收集率为 {avg_rate:.1%}。"
        return f"本轮监测共统计 {len(table)} 个监测点位的数据获取情况。"

    if module.key == "rainfall_analysis":
        daily = tables.get("降雨概况", tables.get("日降雨量统计", pd.DataFrame()))
        events = tables.get("降雨场次分析", tables.get("场次降雨统计", pd.DataFrame()))
        total_rain = _numeric_sum(daily, ("日降雨量(mm)", "rain_mm"))
        rainy_days = 0
        if not daily.empty:
            rain_col = "日降雨量(mm)" if "日降雨量(mm)" in daily.columns else "rain_mm"
            if rain_col in daily.columns:
                rainy_days = int(pd.to_numeric(daily[rain_col], errors="coerce").fillna(0).gt(0).sum())
        if total_rain is not None:
            return f"降雨分析共统计 {len(daily)} 天，其中降雨日 {rainy_days} 天，总降雨量 {total_rain:.1f} mm，识别降雨场次 {len(events)} 场。"
        return f"降雨分析共输出 {len(daily)} 条日统计和 {len(events)} 条场次统计。"

    if module.key == "dry_pattern_analysis":
        table = tables.get("排污规律分析", module_tables[0][1])
        category_col = "规律类别" if "规律类别" in table.columns else "category_name"
        if category_col in table.columns:
            counts = table[category_col].astype(str).value_counts().to_dict()
            count_text = "，".join(f"{name} {count} 个" for name, count in counts.items())
            return f"旱天排污规律统计分析覆盖 {len(table)} 个点位，分类结果为：{count_text}。"
        return f"旱天排污规律统计分析覆盖 {len(table)} 个点位。"

    if module.key == "operation_risk_analysis":
        total_rows = sum(len(table) for _, table in module_tables)
        risk_cols = ["运行风险", "溢流风险", "淤积风险", "running_risk", "overflow_risk", "silting_risk"]
        fragments: list[str] = []
        for _, table in module_tables:
            for col in risk_cols:
                if col in table.columns:
                    counts = table[col].astype(str).value_counts().head(4).to_dict()
                    fragments.extend(f"{name} {count} 项" for name, count in counts.items())
                    break
        if fragments:
            return f"污水系统运行风险分析共输出 {total_rows} 条记录，主要风险分布为：{'，'.join(fragments[:6])}。"
        return f"污水系统运行风险分析共输出 {total_rows} 条记录。"

    summary = next((item for item in fallback_summaries if module.title in item), "")
    if summary:
        return summary
    return f"{module.title}已根据综合结果表生成。"


def _rainfall_chart_paths(chart_paths: dict[str, str] | None) -> list[Path]:
    return [Path(value) for value in (chart_paths or {}).values() if value and Path(value).is_file()]


def _rainfall_chart_map(chart_paths: dict[str, str] | None) -> dict[str, Path]:
    values = chart_paths or {}
    paths = {
        "daily": Path(values["daily_bar"]) if values.get("daily_bar") else None,
        "ratio": Path(values["rainy_ratio"]) if values.get("rainy_ratio") else None,
    }
    return {key: path for key, path in paths.items() if path is not None and path.is_file()}


def _add_picture_before_paragraph(document: DocumentObject, paragraph: Paragraph, image_path: Path, width_inches: float = 5.6):
    image_paragraph = document.add_paragraph()
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    block = image_paragraph._p
    parent = block.getparent()
    if parent is not None:
        parent.remove(block)
    paragraph._p.addprevious(block)


def _insert_rainfall_images_at_captions(document: DocumentObject, rainfall_chart_paths: dict[str, str] | None) -> int:
    chart_map = _rainfall_chart_map(rainfall_chart_paths)
    if not chart_map:
        return 0
    targets = (
        ("daily", ("日降雨量时间序列", "日降雨量统计图", "图 15", "图15")),
        ("ratio", ("降雨日占比", "降雨日与非降雨日占比", "非降雨日", "图 16", "图16")),
    )
    inserted = 0
    used: set[str] = set()
    for key, keywords in targets:
        path = chart_map.get(key)
        if not path:
            continue
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text and any(keyword in text for keyword in keywords):
                _add_picture_before_paragraph(document, paragraph, path)
                inserted += 1
                used.add(key)
                break
    return inserted


def build_report(
    output_file: Path,
    title: str,
    summaries: list[str] | None = None,
    template_file: Path | None = None,
    analysis_tables: dict[str, pd.DataFrame] | None = None,
    site_info_file: Path | None = None,
    outputs_dir: Path | None = None,
    sections: list[str] | None = None,
    dry_curve_data: dict[str, pd.DataFrame] | None = None,
    has_rainfall_data: bool = True,
    point_ids: list[str] | None = None,
    start: str | None = None,
    end: str | None = None,
    rainfall_chart_paths: dict[str, str] | None = None,
    pattern_chart_paths: dict[str, list[str]] | None = None,
    artifact_scope: str = "全网_全时段",
) -> dict[str, object]:
    if Path(output_file).suffix.lower() != ".docx":
        raise ValueError(f"报告输出必须是 .docx 文件: {Path(output_file).name}")
    summaries = summaries or []
    tables = analysis_tables or {}
    template_used = bool(template_file and template_file.exists())
    if template_used:
        from .pipeline_report_assembler.assembler import run_report_assembler

        result = run_report_assembler(
            template_file=template_file,
            analysis_results=tables,
            site_info_file=site_info_file or Path(),
            output_file=output_file,
            dry_curve_data=dry_curve_data,
            filter_result_path=None,
            config={"monitoring_start": start or "", "monitoring_end": end or ""},
            has_rainfall_data=has_rainfall_data,
            llm_client=None,
            sections=sections,
            point_ids=point_ids,
            rainfall_chart_paths=rainfall_chart_paths,
            pattern_chart_paths=pattern_chart_paths,
            artifact_scope=artifact_scope,
        )
        stats = result.get("stats", {})
        return {
            "output_file": str(result.get("output_file", output_file)),
            "template_used": True,
            "template_file": str(template_file),
            "templated_sections": sections or ["监测概况", "降雨分析", "旱天排污规律统计分析", "污水系统运行风险分析"],
            "generated_sections": [],
            "missing_sheets": [],
            "warnings": [str(item) for item in result.get("warnings", [])],
            "stats": {
                "paragraphs": stats.get("paragraphs", 0),
                "tables": stats.get("tables_filled", 0),
                "inserted_tables": stats.get("tables_filled", 0),
                "filled_template_tables": stats.get("tables_filled", 0),
                "inserted_images": stats.get("images_inserted", 0),
                "text_replaced": stats.get("text_replaced", 0),
                "points_processed": stats.get("points_processed", 0),
                "llm_generated": stats.get("llm_generated", 0),
            },
        }

    document = Document(str(template_file)) if template_used else Document()
    if not template_used:
        document.add_heading(title, level=0)

    templated_sections: list[str] = []
    generated_sections: list[str] = []
    missing_sheets: list[str] = []
    inserted_tables = 0
    inserted_images = 0
    filled_template_sheets: set[str] = set()
    filled_template_table_count = 0
    text_replaced = 0
    rainfall_caption_images = 0
    if template_used:
        filled_template_table_count, filled_template_sheets = _fill_template_tables(document, tables, site_info_file)
        inserted_tables += filled_template_table_count
        text_replaced = _replace_template_text(document, tables)
        rainfall_caption_images = _insert_rainfall_images_at_captions(document, rainfall_chart_paths)
        inserted_images += rainfall_caption_images
        pattern_text, pattern_images = _render_pattern_section(document, tables, pattern_chart_paths)
        if pattern_text:
            text_replaced += pattern_text
            inserted_images += pattern_images
            filled_template_sheets.update({"排污规律分析", "旱天分析"})

    for module in _selected_modules(sections):
        module_tables = _table_for(module, tables)
        missing_sheets.extend(
            sheet
            for sheet in module.sheets
            if sheet not in tables and sheet not in filled_template_sheets
        )
        anchor_paragraph = _find_template_anchor(document, module) if template_used else None
        if anchor_paragraph is None:
            if document.paragraphs:
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            anchor_paragraph = document.add_heading(module.title, level=1)
            generated_sections.append(module.title)
        else:
            templated_sections.append(module.title)

        anchor = anchor_paragraph._p
        anchor = _add_paragraph_after(document, anchor, _module_summary(module, module_tables, summaries))
        for sheet_name, table_df in module_tables:
            if sheet_name in filled_template_sheets:
                continue
            anchor = _add_paragraph_after(document, anchor, f"表：{sheet_name}")
            anchor = _add_table_after(document, anchor, table_df)
            inserted_tables += 1

        if module.key == "rainfall_analysis":
            if rainfall_caption_images == 0:
                for image_path in _rainfall_chart_paths(rainfall_chart_paths):
                    anchor = _add_picture_after(document, anchor, image_path)
                    inserted_images += 1

    output_file.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_file)
    return {
        "output_file": str(output_file),
        "template_used": template_used,
        "template_file": str(template_file) if template_used else None,
        "templated_sections": templated_sections,
        "generated_sections": generated_sections,
        "missing_sheets": sorted(set(missing_sheets)),
        "stats": {
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "inserted_tables": inserted_tables,
            "filled_template_tables": filled_template_table_count,
            "inserted_images": inserted_images,
            "text_replaced": text_replaced,
        },
    }
