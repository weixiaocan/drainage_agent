"""Chapter-oriented report assembly orchestration."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Optional
from zipfile import is_zipfile

import pandas as pd
from docx import Document

from .data_context import build_report_context
from .facts import build_report_facts
from .llm_section_writer import LLMSectionWriter
from .sections.pattern import render_pattern_section
from .sections.rainfall import render_rainfall_section
from .sections.risk import render_risk_section
from .sections.site_overview import render_site_overview
from .template_scanner import scan_template
from .validator import validate_report


@dataclass
class ReportConfig:
    """Report assembly configuration."""

    monitoring_start: str = ""
    monitoring_end: str = ""
    monitoring_round: str = "第一轮"
    rainfall_threshold_mm: float = 2.0
    baseinfo_path: str = ""


def run_report_assembler(
    template_file: Path,
    analysis_results: Dict[str, pd.DataFrame],
    site_info_file: Path,
    output_file: Path,
    dry_curve_data: Dict[str, pd.DataFrame] | None = None,
    filter_result_path: Path | None = None,
    config: Dict[str, Any] | None = None,
    has_rainfall_data: bool = True,
    llm_client=None,
    sections: list[str] | None = None,
    point_ids: list[str] | None = None,
    rainfall_chart_paths: dict[str, str] | None = None,
    pattern_chart_paths: dict[str, list[str]] | None = None,
    artifact_scope: str = "全网_全时段",
) -> Dict[str, Any]:
    """Assemble the Word report by template sections."""
    cfg = _build_config(config)
    template_file = Path(template_file)
    site_info_file = Path(site_info_file)
    output_file = Path(output_file)
    if output_file.suffix.lower() != ".docx":
        raise ValueError(f"报告输出必须是 .docx 文件: {output_file.name}")
    output_file.parent.mkdir(parents=True, exist_ok=True)

    print(f"读取报告模板: {template_file}")
    doc = Document(template_file)
    selected = _selected_section_keys(sections)
    dry_only = _is_dry_only_report(sections, selected)
    context = build_report_context(
        analysis_results=analysis_results,
        site_info_file=site_info_file,
        dry_curve_data=dry_curve_data,
        has_rainfall_data=has_rainfall_data,
        point_ids=point_ids,
        rainfall_chart_paths=rainfall_chart_paths,
        pattern_chart_paths=pattern_chart_paths,
        artifact_scope=artifact_scope,
    )
    _resolve_pattern_chart_paths(context, output_file.parent)
    template_map = scan_template(doc)
    baseinfo_path = Path(cfg.baseinfo_path) if cfg.baseinfo_path else _default_baseinfo_path(template_file)
    facts = build_report_facts(context, baseinfo_path=baseinfo_path)
    if cfg.monitoring_start or cfg.monitoring_end:
        facts.monitoring_period_text = _scope_period_text(cfg.monitoring_start, cfg.monitoring_end)
        facts.operation_period_text = facts.monitoring_period_text
    llm_writer = LLMSectionWriter(llm_client)

    warnings: list[str] = []
    warnings.extend(context.warnings)
    warnings.extend(template_map.warnings)
    stats = {
        "tables_filled": 0,
        "images_inserted": 0,
        "points_processed": facts.point_count,
        "text_replaced": 0,
        "llm_generated": 0,
        "warnings": 0,
    }

    print(f"报告包含 {len(doc.tables)} 个表格")
    print(f"识别点位: {facts.point_ids}")

    include_dry_risk, include_rainy_risk = _selected_risk_modes(sections)
    renderers = {
        "monitoring_overview": lambda: render_site_overview(doc, template_map, context, facts, warnings),
        "rainfall_analysis": lambda: render_rainfall_section(
            doc, template_map, context, facts, output_file.parent, warnings
        ),
        "dry_pattern_analysis": lambda: render_pattern_section(
            doc, template_map, context, facts, llm_writer, warnings
        ),
        "operation_risk_analysis": lambda: render_risk_section(
            doc,
            template_map,
            context,
            facts,
            llm_writer,
            warnings,
            include_dry=include_dry_risk,
            include_rainy=include_rainy_risk,
        ),
    }
    for key in selected:
        section_stats = renderers[key]()
        if key == "dry_pattern_analysis":
            expected_images = 2 * len(facts.pattern_details)
            if section_stats.get("images_inserted", 0) != expected_images:
                raise ValueError(
                    f"排污规律图片插入不完整: 应插入 {expected_images} 张，"
                    f"实际 {section_stats.get('images_inserted', 0)} 张"
                )
        _merge_stats(stats, section_stats)

    _prune_unselected_sections(doc, selected)
    if dry_only:
        _prune_heading_blocks(doc, {"雨天运行风险分析", "雨天风险"})

    validation = validate_report(
        doc,
        facts,
        selected_sections=selected,
        include_dry_risk=include_dry_risk,
        include_rainy_risk=include_rainy_risk,
    )
    warnings.extend(validation.warnings)
    if validation.critical:
        warnings.extend(validation.critical)
        raise ValueError("报告校验失败: " + "；".join(validation.critical))

    _apply_explicit_heading_numbers(doc, selected)
    doc.save(output_file)
    if not is_zipfile(output_file):
        raise ValueError(f"生成的报告不是有效 Word 文档: {output_file}")
    stats["warnings"] = len(warnings)
    print(f"保存报告: {output_file}")
    _print_warnings(warnings)
    return {"output_file": output_file, "stats": stats, "warnings": warnings}


SECTION_ALIASES = {
    "monitoring_overview": {"监测概况", "数据概况", "概述与数据质量", "数据体检", "数据质量"},
    "rainfall_analysis": {"降雨分析", "降雨统计", "雨天事件统计", "事件响应", "RDII"},
    "dry_pattern_analysis": {
        "旱天排污规律统计分析", "旱天排污规律", "旱天排污规律分析", "点位特征对比分析",
        "排污规律", "排污规律分析", "旱天分析",
    },
    "operation_risk_analysis": {
        "污水系统运行风险分析", "污水系统运行风险", "风险评估", "旱天风险",
        "旱天运行风险评估", "结论与建议", "雨天风险", "溢流风险",
    },
}

SECTION_CHAPTER_TITLES = {
    "monitoring_overview": {"监测概况", "概述与数据质量"},
    "rainfall_analysis": {"降雨分析"},
    "dry_pattern_analysis": {"旱天排污规律统计分析", "旱天排污规律分析"},
    "operation_risk_analysis": {"污水系统运行风险分析", "污水系统运行风险"},
}


def _selected_section_keys(sections: list[str] | None) -> list[str]:
    if not sections:
        return list(SECTION_ALIASES)
    selected = [
        key
        for key, aliases in SECTION_ALIASES.items()
        if any(
            section == alias or section.startswith(f"{alias}（") or section.startswith(f"{alias}(")
            for section in sections
            for alias in aliases
        )
    ]
    return selected or list(SECTION_ALIASES)


def _selected_risk_modes(sections: list[str] | None) -> tuple[bool, bool]:
    if not sections:
        return True, True
    requested = set(sections)
    full = bool(requested.intersection({"污水系统运行风险分析", "污水系统运行风险", "运行风险分析", "风险评估"}))
    include_dry = full or bool(requested.intersection({"旱天风险", "旱天运行风险评估", "结论与建议"}))
    include_rainy = full or bool(requested.intersection({"雨天风险", "雨天溢流风险", "溢流风险"}))
    return include_dry, include_rainy


def _is_dry_only_report(sections: list[str] | None, selected: list[str]) -> bool:
    if not sections:
        return False
    include_dry, include_rainy = _selected_risk_modes(sections)
    return (
        "rainfall_analysis" not in selected
        and not include_rainy
        and ("dry_pattern_analysis" in selected or include_dry)
    )


def _resolve_pattern_chart_paths(context, output_dir: Path) -> None:
    """Fill missing chart mappings from the exact scoped output directory."""
    scoped_dir = output_dir / "特征曲线图" / context.artifact_scope
    resolved: dict[str, list[str]] = {}
    for point_id in context.point_ids:
        mapped = [Path(value) for value in context.pattern_chart_paths.get(point_id, [])]
        candidates = [path for path in mapped if path.is_file()]
        if scoped_dir.is_dir():
            for path in sorted(scoped_dir.glob(f"{point_id}_*.png")):
                if path not in candidates:
                    candidates.append(path)
        resolved[point_id] = [str(path) for path in candidates]
    if "selected" in context.pattern_chart_paths:
        resolved["selected"] = list(context.pattern_chart_paths["selected"])
    context.pattern_chart_paths = resolved


def _prune_unselected_sections(doc: Document, selected: list[str]) -> None:
    starts: list[tuple[str, object]] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for key, aliases in SECTION_ALIASES.items():
            if text in aliases:
                starts.append((key, paragraph._p))
                break
    if not starts:
        return
    body = doc._element.body
    children = list(body)
    positions = [(key, children.index(element)) for key, element in starts if element in children]
    for idx in range(len(positions) - 1, -1, -1):
        key, start = positions[idx]
        if key in selected:
            continue
        end = positions[idx + 1][1] if idx + 1 < len(positions) else len(children) - 1
        for child in children[start:end]:
            if child.getparent() is body:
                body.remove(child)


def _prune_heading_blocks(doc: Document, titles: set[str]) -> None:
    starts: list[tuple[int, int, object]] = []
    body = doc._element.body
    children = list(body)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text not in titles or paragraph._p not in children:
            continue
        level = _heading_level(paragraph.style.name)
        starts.append((children.index(paragraph._p), level, paragraph._p))
    for start, level, _element in sorted(starts, reverse=True):
        end = len(children) - 1
        for idx in range(start + 1, len(children)):
            child = children[idx]
            paragraph = next((p for p in doc.paragraphs if p._p is child), None)
            if paragraph is None:
                continue
            next_level = _heading_level(paragraph.style.name)
            if next_level and (not level or next_level <= level):
                end = idx
                break
        for child in children[start:end]:
            if child.getparent() is body:
                body.remove(child)


def _heading_level(style_name: str) -> int | None:
    if style_name.startswith("标题"):
        suffix = style_name.removeprefix("标题")
        return int(suffix) if suffix.isdigit() else None
    if style_name.startswith("Heading "):
        suffix = style_name.removeprefix("Heading ")
        return int(suffix) if suffix.isdigit() else None
    return None


def _apply_explicit_heading_numbers(doc: Document, selected: list[str]) -> None:
    """Renumber retained chapters after template pruning.

    The source template uses Word list numbering. Once preceding chapters are
    removed, a retained chapter can start at Heading 2 and Word may display no
    number at all. Explicit prefixes make the partial report deterministic in
    Word, browser previews, and converted PDFs.
    """
    chapter_titles = {
        title
        for key in selected
        for title in SECTION_CHAPTER_TITLES[key]
    }
    heading_styles = [style for style in doc.styles if _heading_level(style.name)]
    for style in heading_styles:
        p_pr = style.element.pPr
        if p_pr is not None and p_pr.numPr is not None:
            p_pr.remove(p_pr.numPr)

    chapter = 0
    sublevels = [0, 0]
    chapter_source_level: int | None = None
    for paragraph in doc.paragraphs:
        source_level = _heading_level(paragraph.style.name)
        if source_level is None or source_level == 0:
            continue
        plain_text = re.sub(r"^\d+(?:\.\d+)*\s+", "", paragraph.text.strip())
        if plain_text in chapter_titles:
            chapter += 1
            sublevels = [0, 0]
            chapter_source_level = source_level
            target_level = 1
            prefix = f"{chapter} "
        elif chapter and chapter_source_level is not None:
            target_level = max(2, min(3, source_level - chapter_source_level + 1))
            index = target_level - 2
            sublevels[index] += 1
            if index == 0:
                sublevels[1] = 0
                prefix = f"{chapter}.{sublevels[0]} "
            else:
                prefix = f"{chapter}.{sublevels[0]}.{sublevels[1]} "
        else:
            continue

        style_name = next(
            (name for name in (f"标题{target_level}", f"Heading {target_level}") if name in doc.styles),
            None,
        )
        if style_name:
            paragraph.style = doc.styles[style_name]
        p_pr = paragraph._p.pPr
        if p_pr is not None and p_pr.numPr is not None:
            p_pr.remove(p_pr.numPr)
        if paragraph.runs:
            paragraph.runs[0].text = prefix + plain_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(prefix + plain_text)


def _build_config(config: Optional[Dict[str, Any]]) -> ReportConfig:
    cfg = ReportConfig()
    if config:
        for key, value in config.items():
            if hasattr(cfg, key):
                setattr(cfg, key, value)
    return cfg


def _scope_period_text(start: str, end: str) -> str:
    start_ts = pd.to_datetime(start, errors="coerce") if start else None
    end_ts = pd.to_datetime(end, errors="coerce") if end else None
    start_text = start_ts.strftime("%Y/%m/%d") if start_ts is not None and not pd.isna(start_ts) else None
    end_text = end_ts.strftime("%Y/%m/%d") if end_ts is not None and not pd.isna(end_ts) else None
    if start_text and not end_text:
        return f"{start_text}日之后"
    if end_text and not start_text:
        return f"{end_text}日之前"
    if start_text and end_text:
        return f"{start_text}日-{end_text}日"
    return "全时段"


def _default_baseinfo_path(template_file: Path) -> Path:
    project_root = template_file.parent.parent
    return project_root / "data" / "baseinfo.xlsx"


def _merge_stats(target: dict[str, int], source: dict[str, int]) -> None:
    for key, value in source.items():
        target[key] = target.get(key, 0) + int(value)


def _print_warnings(warnings: list[str]) -> None:
    if not warnings:
        return
    print("报告组装 warnings:")
    for warning in warnings[:20]:
        print(f"  - {warning}")
    if len(warnings) > 20:
        print(f"  - ... 另有 {len(warnings) - 20} 条")
