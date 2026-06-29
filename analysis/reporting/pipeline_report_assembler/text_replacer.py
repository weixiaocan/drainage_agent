"""规则文字替换模块

根据模板中的占位文字，用实际数据进行规则替换。
按模块组织：监测概况部分和降雨分析部分。
"""

import re
from datetime import datetime
from typing import Any, Dict, Optional

import pandas as pd
from docx import Document


class TextReplacer:
    """文字段落规则替换器"""

    def __init__(self, context: Dict[str, Any]):
        """
        Args:
            context: 替换上下文数据
        """
        self.context = context

    def replace_site_info_paragraphs(self, doc: Document) -> int:
        """
        替换监测概况部分的文字（属于点位信息模块）

        Returns:
            替换的段落数量
        """
        replaced = 0

        for para in doc.paragraphs:
            original_text = para.text
            new_text = self._replace_site_info_text(original_text)

            if new_text != original_text:
                self._update_paragraph(para, new_text)
                replaced += 1

        return replaced

    def replace_rainfall_paragraphs(self, doc: Document) -> int:
        """
        替换降雨分析部分的文字（属于降雨分析模块）

        Returns:
            替换的段落数量
        """
        replaced = 0

        for para in doc.paragraphs:
            original_text = para.text
            new_text = self._replace_rainfall_text(original_text)

            if new_text != original_text:
                self._update_paragraph(para, new_text)
                replaced += 1

        return replaced

    def _replace_site_info_text(self, text: str) -> str:
        """替换监测概况相关文字"""
        result = text

        # 点位数
        if self.context.get("point_count"):
            result = re.sub(
                r'共布设\d+个流量监测点位',
                f'共布设{self.context["point_count"]}个流量监测点位',
                result
            )

        # 监测时段
        if self.context.get("start_date") and self.context.get("end_date"):
            start = self._format_date(self.context["start_date"])
            end = self._format_date(self.context["end_date"])
            result = re.sub(
                r'时间段选择[\d/]+日?-[\d/]+日?',
                f'时间段选择{start}-{end}',
                result
            )

        # 运维时段
        if self.context.get("operation_start") and self.context.get("operation_end"):
            op_start = self._format_date_cn(self.context["operation_start"])
            op_end = self._format_date_cn(self.context["operation_end"])
            result = re.sub(
                r'\d{4}年\d{1,2}月\d{1,2}日至\d{1,2}月\d{1,2}日期间',
                f'{op_start}至{op_end}期间',
                result
            )

        # 点位运行状态
        if self.context.get("point_count"):
            result = re.sub(
                r'\d+个监测点位在监测期间运行状态良好',
                f'{self.context["point_count"]}个监测点位在监测期间运行状态良好',
                result
            )

        # 数据条数
        if self.context.get("data_count_wan"):
            result = re.sub(
                r'共收集分钟级监测数据超\d+万条',
                f'共收集分钟级监测数据超{self.context["data_count_wan"]}万条',
                result
            )

        # 数据收集率描述
        if self.context.get("collection_rate_desc"):
            result = re.sub(
                r'\d+个点位的有效数据收集率[^。，]*',
                self.context["collection_rate_desc"],
                result
            )

        return result

    def _replace_rainfall_text(self, text: str) -> str:
        """替换降雨分析相关文字"""
        result = text

        # 降雨天数
        if self.context.get("rainy_days"):
            result = re.sub(
                r'降雨日天数为\d+天',
                f'降雨日天数为{self.context["rainy_days"]}天',
                result
            )

        # 总降雨量
        if self.context.get("total_rainfall"):
            result = re.sub(
                r'总降雨量[\d.]+\s*mm',
                f'总降雨量{self.context["total_rainfall"]}mm',
                result
            )

        # 最大日降雨量
        if self.context.get("max_daily_rainfall"):
            result = re.sub(
                r'日最大降雨量为[\d.]+\s*mm',
                f'日最大降雨量为{self.context["max_daily_rainfall"]}mm',
                result
            )

        # 最大降雨日期
        if self.context.get("max_rainfall_date"):
            result = re.sub(
                r'发生在[\d-]+',
                f'发生在{self.context["max_rainfall_date"]}',
                result
            )

        # 监测期总天数
        if self.context.get("total_days"):
            result = re.sub(
                r'监测期内共\d+个自然日',
                f'监测期内共{self.context["total_days"]}个自然日',
                result
            )

        # 降雨场次
        if self.context.get("rainfall_events"):
            result = re.sub(
                r'有效降雨场次\d+场',
                f'有效降雨场次{self.context["rainfall_events"]}场',
                result
            )

        if self.context.get("event_total_rainfall"):
            result = re.sub(
                r'累计降雨量[\d.]+\s*mm',
                f'累计降雨量{self.context["event_total_rainfall"]}mm',
                result
            )

        if self.context.get("max_event_rainfall"):
            result = re.sub(
                r'最大场次降雨量为[\d.]+\s*mm',
                f'最大场次降雨量为{self.context["max_event_rainfall"]}mm',
                result
            )

        return result

    def _format_date(self, date: Any) -> str:
        """格式化日期为 YYYY/M/D 格式"""
        if isinstance(date, datetime):
            return f"{date.year}/{date.month}/{date.day}"
        if isinstance(date, str):
            for fmt in ["%Y-%m-%d", "%Y/%m/%d", "%Y年%m月%d日"]:
                try:
                    dt = datetime.strptime(date, fmt)
                    return f"{dt.year}/{dt.month}/{dt.day}"
                except ValueError:
                    continue
            return date
        return str(date)

    def _format_date_cn(self, date: Any) -> str:
        """格式化日期为 YYYY年M月D日 格式"""
        if isinstance(date, datetime):
            return f"{date.year}年{date.month}月{date.day}日"
        if isinstance(date, str):
            for fmt in ["%Y-%m-%d", "%Y/%m/%d", "%Y年%m月%d日"]:
                try:
                    dt = datetime.strptime(date, fmt)
                    return f"{dt.year}年{dt.month}月{dt.day}日"
                except ValueError:
                    continue
            return date
        return str(date)

    def _update_paragraph(self, para, new_text: str) -> None:
        """更新段落文字，保留格式"""
        first_run_font = None
        if para.runs:
            first_run = para.runs[0]
            first_run_font = {
                "name": first_run.font.name,
                "size": first_run.font.size,
                "bold": first_run.font.bold,
            }

        for run in para.runs:
            run.text = ""

        if para.runs:
            para.runs[0].text = new_text
        else:
            run = para.add_run(new_text)
            if first_run_font:
                run.font.name = first_run_font["name"]
                run.font.size = first_run_font["size"]
                run.font.bold = first_run_font["bold"]

    # 保留旧方法以兼容
    def replace_in_document(self, doc: Document) -> int:
        """在文档中执行所有规则替换（兼容旧调用）"""
        replaced = self.replace_site_info_paragraphs(doc)
        replaced += self.replace_rainfall_paragraphs(doc)
        return replaced


def build_site_info_context(
    collection_df: Optional[pd.DataFrame] = None,
    baseinfo: Optional[Dict] = None,
) -> Dict[str, Any]:
    """
    构建监测概况部分的替换上下文。

    Args:
        collection_df: 数据收集率统计 DataFrame
        baseinfo: 项目基本信息字典

    Returns:
        替换上下文字典
    """
    context: Dict[str, Any] = {}

    # 从数据收集率统计提取
    if collection_df is not None and not collection_df.empty:
        context["point_count"] = len(collection_df)
        context["data_count"] = collection_df["监测数据条数"].sum()
        context["data_count_wan"] = int(context["data_count"] // 10000)
        context["monitoring_days"] = collection_df["监测天数"].max()

        # 生成收集率描述
        rates = collection_df["数据收集率(%)"]
        high_rate_count = (rates >= 99).sum()
        if high_rate_count == len(collection_df):
            context["collection_rate_desc"] = f"{len(collection_df)}个点位的有效数据收集率均超过99%"
        else:
            context["collection_rate_desc"] = f"{high_rate_count}个点位的有效数据收集率超过99%，其余点位收集率良好"

    # 从基本信息提取
    if baseinfo:
        if "监测开始时间" in baseinfo:
            context["start_date"] = baseinfo["监测开始时间"]
        if "监测结束时间" in baseinfo:
            context["end_date"] = baseinfo["监测结束时间"]
        if "监测轮次" in baseinfo:
            context["monitoring_round"] = baseinfo["监测轮次"]

    return context


def build_rainfall_context(
    rainfall_daily_df: Optional[pd.DataFrame] = None,
    rainfall_event_df: Optional[pd.DataFrame] = None,
) -> Dict[str, Any]:
    """
    构建降雨分析部分的替换上下文。

    Args:
        rainfall_daily_df: 日降雨量统计 DataFrame
        rainfall_event_df: 场次降雨统计 DataFrame

    Returns:
        替换上下文字典
    """
    context: Dict[str, Any] = {}

    # 从日降雨量统计提取
    if rainfall_daily_df is not None and not rainfall_daily_df.empty:
        # 兼容不同列名格式
        rain_col = None
        for col in rainfall_daily_df.columns:
            if col == "daily_rain_mm" or "日降雨量" in str(col):
                rain_col = col
                break

        if rain_col:
            rainy_rows = rainfall_daily_df[rainfall_daily_df[rain_col] > 0]
            context["rainy_days"] = len(rainy_rows)
            context["total_rainfall"] = round(rainfall_daily_df[rain_col].sum(), 1)
            context["max_daily_rainfall"] = round(rainy_rows[rain_col].max(), 1) if len(rainy_rows) > 0 else 0

            # 总监测天数
            context["total_days"] = len(rainfall_daily_df)

            # 最大降雨日期
            if len(rainy_rows) > 0:
                max_idx = rainy_rows[rain_col].idxmax()
                date_col = "date" if "date" in rainy_rows.columns else "日期"
                max_date = rainy_rows.loc[max_idx, date_col]
                if isinstance(max_date, datetime):
                    context["max_rainfall_date"] = f"{max_date.year}年{max_date.month}月{max_date.day}日"
                else:
                    context["max_rainfall_date"] = str(max_date)[:10]

    # 从场次降雨统计提取
    if rainfall_event_df is not None and not rainfall_event_df.empty:
        context["rainfall_events"] = len(rainfall_event_df)

        # 兼容不同列名格式
        total_col = None
        for col in rainfall_event_df.columns:
            if col == "total_rain_mm" or "总降雨量" in str(col):
                total_col = col
                break

        if total_col:
            context["event_total_rainfall"] = round(rainfall_event_df[total_col].sum(), 1)
            context["max_event_rainfall"] = round(rainfall_event_df[total_col].max(), 1)

    return context


def build_context_from_data(
    collection_df: Optional[pd.DataFrame] = None,
    rainfall_daily_df: Optional[pd.DataFrame] = None,
    rainfall_event_df: Optional[pd.DataFrame] = None,
    site_info_df: Optional[pd.DataFrame] = None,
    baseinfo: Optional[Dict] = None,
) -> Dict[str, Any]:
    """
    从数据源构建替换上下文（兼容旧调用）。

    合并监测概况和降雨分析两部分上下文。
    """
    context = build_site_info_context(collection_df, baseinfo)
    context.update(build_rainfall_context(rainfall_daily_df, rainfall_event_df))
    return context
