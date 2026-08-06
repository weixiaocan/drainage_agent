"""Report validation for template residue and fact consistency."""

from __future__ import annotations

import re
from dataclasses import dataclass, field

from docx import Document

from .facts import ReportFacts


@dataclass
class ValidationResult:
    warnings: list[str] = field(default_factory=list)
    critical: list[str] = field(default_factory=list)

    @property
    def ok(self) -> bool:
        return not self.critical


def validate_report(
    doc: Document,
    facts: ReportFacts,
    selected_sections: list[str] | None = None,
    include_dry_risk: bool = True,
    include_rainy_risk: bool = True,
) -> ValidationResult:
    result = ValidationResult()
    text = _document_text(doc)

    legacy_points = sorted(set(re.findall(r"(?<!\d)1-[\d-]+#", text)))
    if legacy_points:
        result.critical.append(f"报告仍包含旧模板点位编号: {', '.join(legacy_points[:10])}")

    legacy_times = ["2024/9/18", "2024/11/26", "2026年2月1日至2月11日"]
    found_times = [item for item in legacy_times if item in text]
    if found_times:
        result.critical.append(f"报告仍包含模板旧时间: {', '.join(found_times)}")

    if facts.device_count != 13 and "13台流量监测设备" in text:
        result.critical.append("报告仍包含模板旧设备数量: 13台流量监测设备")

    residue_fragments = ["44个流量监测点位", "32个点位", "0.W", "最大充满度W", "____/__/__"]
    found_residue = [item for item in residue_fragments if item in text]
    if found_residue:
        result.critical.append(f"报告仍包含模板占位内容: {', '.join(found_residue)}")

    allowed_points = {str(point).upper() for point in facts.point_ids}
    mentioned_points = set(re.findall(r"(?<![A-Za-z0-9])W\d+(?![A-Za-z0-9])", text, flags=re.IGNORECASE))
    unexpected_points = sorted(point for point in mentioned_points if point.upper() not in allowed_points)
    if unexpected_points:
        result.critical.append(f"报告包含本次数据中不存在的点位: {', '.join(unexpected_points[:20])}")

    if facts.point_count and f"共布设{facts.point_count}个流量监测点位" not in text:
        result.warnings.append("未找到与事实一致的点位总数描述")

    selected = set(selected_sections or ["monitoring_overview", "rainfall_analysis", "dry_pattern_analysis", "operation_risk_analysis"])
    headings = {
        "monitoring_overview": "监测概况",
        "rainfall_analysis": "降雨分析",
        "dry_pattern_analysis": "旱天排污规律统计分析",
        "operation_risk_analysis": "污水系统运行风险",
    }
    for key, heading in headings.items():
        if key not in selected and heading in text:
            result.critical.append(f"报告仍包含未选择章节: {heading}")
    if "operation_risk_analysis" in selected and include_dry_risk and "表 12 旱天运行状态统计表" not in text:
        result.critical.append("报告缺少表题: 表 12 旱天运行状态统计表")
    if "operation_risk_analysis" in selected and include_rainy_risk and "雨天运行风险分析" in text and "表 13 第二轮监测雨天运行状态统计表" not in text:
        result.critical.append("报告缺少表题: 表 13 第二轮监测雨天运行状态统计表")

    pattern_start = text.find("旱天排污规律统计分析")
    risk_start = text.find("污水系统运行风险")
    if pattern_start >= 0 and risk_start > pattern_start:
        pattern_text = text[pattern_start:risk_start]
        if "两轮" in pattern_text or "32个点位" in pattern_text:
            result.critical.append("排污规律章节仍包含模板旧统计口径")
        if "本章小结" not in pattern_text:
            result.critical.append("排污规律章节缺少本章小结")
        else:
            summary_tail = pattern_text.split("本章小结", 1)[1].strip()
            if len(summary_tail) < 20:
                result.critical.append("排污规律章节本章小结为空或过短")
        legacy_pattern_captions = re.findall(r"图\s*\d+\s+1-[\d-]+#排污规律特征曲线图", pattern_text)
        if legacy_pattern_captions:
            result.critical.append("排污规律章节仍包含模板旧图题")

    return result


def _document_text(doc: Document) -> str:
    values = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        values.extend(cell.text for row in table.rows for cell in row.cells)
    for section in doc.sections:
        values.extend(paragraph.text for paragraph in section.header.paragraphs)
        values.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(values)
