"""表格行数动态管理模块

提供表格行数的动态调整功能，支持：
- 增加行（复制模板行格式）
- 删除多余行
- 调整特征曲线图表格数量
"""

from copy import deepcopy
from typing import List

from docx.document import Document
from docx.shared import Pt, Twips
from docx.table import Table, _Row
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def adjust_table_rows(table: Table, target_rows: int, template_row_idx: int = 1) -> None:
    """
    调整表格数据行数至目标数量。

    Args:
        table: docx 表格对象
        target_rows: 目标数据行数（不含表头）
        template_row_idx: 模板行索引（用于复制格式），默认为1（第一行数据）

    行为：
        - 当前行数 > 目标：删除末尾多余行
        - 当前行数 < 目标：复制模板行添加新行
    """
    current_data_rows = len(table.rows) - template_row_idx

    if target_rows == current_data_rows:
        return

    if target_rows < current_data_rows:
        # 删除多余行
        _remove_rows(table, current_data_rows - target_rows)
    else:
        # 添加新行
        _add_rows_from_template(table, target_rows - current_data_rows, template_row_idx)


def _remove_rows(table: Table, count: int) -> None:
    """从表格末尾删除指定数量的行"""
    for _ in range(count):
        if len(table.rows) > 1:  # 至少保留一行（表头）
            row = table.rows[-1]._element
            row.getparent().remove(row)


def _add_rows_from_template(table: Table, count: int, template_idx: int) -> None:
    """
    从模板行复制格式并添加新行。

    复制内容：
        - 单元格宽度
        - 字体样式
        - 边框样式
        - 对齐方式
    """
    if template_idx >= len(table.rows):
        template_idx = len(table.rows) - 1

    for _ in range(count):
        new_tr = deepcopy(table.rows[template_idx]._tr)
        table._tbl.append(new_tr)


def _copy_row_format(source_row: _Row, target_row: _Row) -> None:
    """复制行格式到目标行"""
    for src_cell, tgt_cell in zip(source_row.cells, target_row.cells):
        # 复制单元格宽度（跳过 None）
        if src_cell.width is not None:
            tgt_cell.width = src_cell.width

        # 复制段落格式
        for src_para, tgt_para in zip(src_cell.paragraphs, tgt_cell.paragraphs):
            tgt_para.alignment = src_para.alignment
            if src_para.paragraph_format.space_before is not None:
                tgt_para.paragraph_format.space_before = src_para.paragraph_format.space_before
            if src_para.paragraph_format.space_after is not None:
                tgt_para.paragraph_format.space_after = src_para.paragraph_format.space_after

            # 复制字体样式
            if src_para.runs:
                src_run = src_para.runs[0]
                if tgt_para.runs:
                    tgt_run = tgt_para.runs[0]
                else:
                    tgt_run = tgt_para.add_run()

                if src_run.font.name is not None:
                    tgt_run.font.name = src_run.font.name
                if src_run.font.size is not None:
                    tgt_run.font.size = src_run.font.size
                if src_run.font.bold is not None:
                    tgt_run.font.bold = src_run.font.bold
                if src_run.font.color.rgb:
                    tgt_run.font.color.rgb = src_run.font.color.rgb

        # 复制边框和底纹
        _copy_cell_shading(src_cell, tgt_cell)


def _copy_cell_shading(source_cell, target_cell) -> None:
    """复制单元格底纹"""
    src_tc = source_cell._tc
    tgt_tc = target_cell._tc

    src_tcPr = src_tc.tcPr
    if src_tcPr is not None:
        tgt_tcPr = tgt_tc.get_or_add_tcPr()
        # 复制底纹元素
        for child in src_tcPr:
            if 'shd' in child.tag.lower():
                tgt_tcPr.append(deepcopy(child))


def adjust_curve_image_tables(doc: Document, point_names: List[str], start_table_idx: int = 4) -> None:
    """
    调整特征曲线图表格数量。

    模板中表格4-16用于放置各点位的特征曲线图。
    如果实际点位数与模板不同，需要增删表格。

    Args:
        doc: Word 文档对象
        point_names: 点位名称列表
        start_table_idx: 特征曲线图表格起始索引（默认为4）

    行为：
        - 点位数 < 模板表格数：删除多余表格
        - 点位数 > 模板表格数：复制表格添加
    """
    target_count = len(point_names)
    current_count = _count_curve_tables(doc, start_table_idx)

    if target_count == current_count:
        return

    if target_count < current_count:
        # 删除多余表格
        _remove_curve_tables(doc, start_table_idx, current_count - target_count)
    else:
        # 复制表格添加
        _add_curve_tables(doc, start_table_idx, target_count - current_count)


def _count_curve_tables(doc: Document, start_idx: int) -> int:
    """统计特征曲线图表格数量"""
    tables = doc.tables
    count = 0

    for i in range(start_idx, len(tables)):
        table = tables[i]
        # 特征曲线图表格特征：1行 x 2列
        if len(table.rows) == 1 and len(table.rows[0].cells) == 2:
            # 检查是否包含"特征曲线图"字样
            cell_text = table.rows[0].cells[0].text + table.rows[0].cells[1].text
            if "特征曲线图" in cell_text or "流量特征曲线" in cell_text or "液位特征曲线" in cell_text:
                count += 1
            else:
                break  # 遇到非特征曲线图表格，停止计数
        else:
            break  # 格式不同，停止计数

    return count


def _remove_curve_tables(doc: Document, start_idx: int, count: int) -> None:
    """删除多余的特征曲线图表格"""
    for _ in range(count):
        if start_idx < len(doc.tables):
            table = doc.tables[start_idx]
            table._element.getparent().remove(table._element)


def _add_curve_tables(doc: Document, start_idx: int, count: int) -> None:
    """
    复制特征曲线图表格并添加。

    由于 docx 库不支持直接复制表格，这里采用重新创建的方式。
    """
    if start_idx >= len(doc.tables):
        return

    template_table = doc.tables[start_idx]

    for _ in range(count):
        # 在模板表格后插入新表格
        new_table = _create_curve_table(doc, template_table)
        # 将新表格移动到正确位置
        template_table._element.addprevious(new_table._element)


def _create_curve_table(doc: Document, template_table: Table) -> Table:
    """创建特征曲线图表格"""
    # 创建 1行 x 2列 的表格
    new_table = doc.add_table(rows=1, cols=2)

    # 复制格式
    if len(template_table.rows) > 0:
        src_row = template_table.rows[0]
        tgt_row = new_table.rows[0]

        for i, (src_cell, tgt_cell) in enumerate(zip(src_row.cells, tgt_row.cells)):
            if src_cell.width is not None:
                tgt_cell.width = src_cell.width

            # 设置默认文字
            if i == 0:
                tgt_cell.text = "（a）流量特征曲线图"
            else:
                tgt_cell.text = "（b）液位特征曲线图"

            # 复制段落格式
            if src_cell.paragraphs:
                for src_para, tgt_para in zip(src_cell.paragraphs, tgt_cell.paragraphs):
                    tgt_para.alignment = src_para.alignment

    return new_table


def get_table_row_count(table: Table, header_rows: int = 1) -> int:
    """获取表格数据行数（不含表头）"""
    return max(0, len(table.rows) - header_rows)


def clear_table_data(table: Table, start_row: int = 1) -> None:
    """清空表格数据（保留表头）"""
    for row in table.rows[start_row:]:
        for cell in row.cells:
            cell.text = ""
