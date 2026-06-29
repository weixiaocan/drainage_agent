"""Style-preserving Word write helpers."""

from __future__ import annotations

from copy import deepcopy
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """Replace paragraph text while preserving the first run style."""
    if paragraph.runs:
        first = paragraph.runs[0]
        first.text = str(text)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(str(text))


def set_cell_text(cell: _Cell, text: Any) -> None:
    """Replace cell text while preserving paragraph/run style."""
    value = "" if text is None else str(text)
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    set_paragraph_text(paragraph, value)
    for extra in cell.paragraphs[1:]:
        set_paragraph_text(extra, "")


def clear_cell(cell: _Cell) -> None:
    """Clear text and drawings in a cell without replacing the cell object."""
    for paragraph in cell.paragraphs:
        clear_paragraph(paragraph)


def clear_paragraph(paragraph: Paragraph) -> None:
    """Clear all runs in a paragraph."""
    for run in paragraph.runs:
        run.text = ""
        for drawing in run._element.xpath(".//w:drawing"):
            drawing.getparent().remove(drawing)


def add_picture_to_cell(cell: _Cell, image_path: str, width_inches: float = 2.8) -> None:
    """Insert a picture into a cell while preserving paragraph alignment."""
    clear_cell(cell)
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(image_path, width=Inches(width_inches))


def add_picture_to_paragraph(paragraph: Paragraph, image_path: str, width_inches: float = 5.5) -> None:
    """Insert a picture into an existing paragraph."""
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(image_path, width=Inches(width_inches))


def clone_table_row(table: Table, template_row_idx: int):
    """Append an exact XML clone of a template row and return the new row."""
    if not table.rows:
        raise ValueError("Cannot clone row from an empty table")
    if template_row_idx >= len(table.rows):
        template_row_idx = len(table.rows) - 1
    new_tr = deepcopy(table.rows[template_row_idx]._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def adjust_table_rows_preserve_style(table: Table, target_data_rows: int, template_row_idx: int = 1) -> None:
    """Resize data rows by deleting extras or XML-cloning the template row."""
    current_data_rows = max(0, len(table.rows) - template_row_idx)
    if current_data_rows > target_data_rows:
        for _ in range(current_data_rows - target_data_rows):
            row = table.rows[-1]._element
            row.getparent().remove(row)
    elif current_data_rows < target_data_rows:
        for _ in range(target_data_rows - current_data_rows):
            clone_table_row(table, template_row_idx)


def remove_paragraph_drawings(paragraph: Paragraph) -> int:
    """Remove drawing elements from a paragraph."""
    removed = 0
    for run in paragraph.runs:
        drawings = list(run._element.xpath(".//w:drawing"))
        for drawing in drawings:
            drawing.getparent().remove(drawing)
            removed += 1
    return removed
