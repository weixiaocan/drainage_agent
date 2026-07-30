from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

import pytest

from analysis.report_templates import (
    InvalidReportTemplate,
    ReportTemplateService,
)
from analysis.runs import AnalysisRequest, AnalysisRunner
from analysis.io.standard import STANDARD_FLOW_COLUMNS, STANDARD_FLOW_UNITS
from web.projects import ProjectRepository


def _write_builtin_template(path: Path) -> None:
    document = Document()
    document.add_heading("{{PROJECT_NAME}} / {{BATCH_NAME}}", level=1)
    document.add_paragraph("{{ANALYSIS_SUMMARY}}")
    document.add_heading("空间拓扑关系", level=2)
    document.add_paragraph("{{MANUAL_TOPOLOGY_SECTION}}")
    document.save(path)


def _document_bytes(*paragraphs: str) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    content = BytesIO()
    document.save(content)
    return content.getvalue()


def _analysis_result(tmp_path: Path):
    database = tmp_path / "state" / "drainage.sqlite3"
    files_root = tmp_path / "projects"
    projects = ProjectRepository(database, files_root)
    project = projects.create("报告项目")
    batch = projects.create_batch(project.id, "报告批次")
    standard = projects.batch_workspace(project.id, batch.id) / "standard"
    standard.mkdir(parents=True, exist_ok=True)
    (standard / "flow.csv").write_text(
        ",".join(STANDARD_FLOW_COLUMNS)
        + "\n2026-03-07T00:00:00,D1,W1,2.5,1.2,0.4\n",
        encoding="utf-8",
    )
    (standard / "manifest.json").write_text(
        json.dumps(
            {
                "contract_version": 1,
                "kind": "standard_flow",
                "columns": STANDARD_FLOW_COLUMNS,
                "units": STANDARD_FLOW_UNITS,
                "file": "flow.csv",
            }
        ),
        encoding="utf-8",
    )
    AnalysisRunner(database, files_root).run(
        AnalysisRequest(project.id, batch.id, "data_quality")
    )
    return database, files_root, projects, project, batch


def test_builtin_contract_generates_versioned_drafts_and_comprehensive_tables(
    tmp_path: Path,
) -> None:
    database, files_root, projects, project, batch = _analysis_result(tmp_path)
    builtin = tmp_path / "builtin.docx"
    _write_builtin_template(builtin)
    reports = ReportTemplateService(database, files_root, builtin)

    first = reports.create_draft(project.id, batch.id, "builtin")
    second = reports.create_draft(project.id, batch.id, "builtin")

    assert first.version == 1
    assert second.version == 2
    assert first.report_id != second.report_id
    first_docx = projects.batch_workspace(project.id, batch.id) / first.docx
    second_docx = projects.batch_workspace(project.id, batch.id) / second.docx
    assert first_docx.is_file() and second_docx.is_file()
    text = "\n".join(p.text for p in Document(first_docx).paragraphs)
    assert "报告项目 / 报告批次" in text
    assert "人工补充模块" in text
    workbook = load_workbook(
        projects.batch_workspace(project.id, batch.id) / first.workbook
    )
    assert "数据收集率统计" in workbook.sheetnames
    assert workbook["数据收集率统计"]["A1"].value == "点位编号"
    assert workbook["数据收集率统计"]["A2"].value == "W1"
    assert reports.list_drafts(project.id, batch.id) == [first, second]


def test_custom_template_validation_is_actionable_and_project_scoped(
    tmp_path: Path,
) -> None:
    database, files_root, projects, project, batch = _analysis_result(tmp_path)
    other = projects.create("其他项目")
    builtin = tmp_path / "builtin.docx"
    _write_builtin_template(builtin)
    reports = ReportTemplateService(database, files_root, builtin)

    with pytest.raises(InvalidReportTemplate, match="ANALYSIS_SUMMARY"):
        reports.upload(
            project.id,
            "缺占位符",
            "invalid.docx",
            _document_bytes(
                "{{PROJECT_NAME}}",
                "{{BATCH_NAME}}",
                "{{MANUAL_TOPOLOGY_SECTION}}",
            ),
        )
    custom = reports.upload(
        project.id,
        "单位模板",
        "custom.docx",
        _document_bytes(*sorted(reports.required_placeholders)),
    )

    assert custom.project_id == project.id
    with pytest.raises(LookupError, match="不属于"):
        reports.create_draft(other.id, batch.id, custom.template_id)


def test_web_exposes_contract_templates_and_versioned_report_downloads(
    tmp_path: Path,
) -> None:
    from fastapi.testclient import TestClient

    from quality.tests.test_web_app import FakeAgent, make_deps
    from web.app import create_app

    app = create_app(
        tmp_path,
        deps_factory=make_deps,
        agent_factory=lambda _deps: FakeAgent(),
    )
    with TestClient(app) as client:
        project = client.post(
            "/api/projects", json={"name": "Web 报告项目"}
        ).json()
        batch = client.post(
            f"/api/projects/{project['id']}/batches",
            json={"name": "Web 报告批次"},
        ).json()
        standard = app.state.projects.batch_workspace(
            project["id"], batch["id"]
        ) / "standard"
        standard.mkdir(parents=True, exist_ok=True)
        (standard / "flow.csv").write_text(
            ",".join(STANDARD_FLOW_COLUMNS)
            + "\n2026-03-07T00:00:00,D1,W1,2.5,1.2,0.4\n",
            encoding="utf-8",
        )
        (standard / "manifest.json").write_text(
            json.dumps(
                {
                    "contract_version": 1,
                    "kind": "standard_flow",
                    "columns": STANDARD_FLOW_COLUMNS,
                    "units": STANDARD_FLOW_UNITS,
                    "file": "flow.csv",
                }
            ),
            encoding="utf-8",
        )
        app.state.analysis_runner.run(
            AnalysisRequest(project["id"], batch["id"], "data_quality")
        )

        templates = client.get(
            f"/api/projects/{project['id']}/report-templates"
        )
        created = client.post(
            f"/api/projects/{project['id']}/batches/{batch['id']}/reports",
            json={"template_id": "builtin"},
        )

        assert templates.json()[0]["template_id"] == "builtin"
        assert created.status_code == 201
        assert client.get(created.json()["docx_url"]).status_code == 200
        assert client.get(created.json()["workbook_url"]).status_code == 200
