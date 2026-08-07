from __future__ import annotations

import json
import base64
from pathlib import Path
from types import SimpleNamespace

from docx import Document
import pytest

from quality.eval.eval_stage2.run_eval import (
    apply_after_seed_mutation,
    canonical_case_id,
    compact_pending_results,
    completed_case_ids,
    fresh_root,
    normalize_case,
    preserve_artifacts,
    select_cases,
    trace_evidence,
    tree_snapshot,
    tool_seq,
    validate_cases,
)
from quality.eval.eval_stage2.view import load_checks, load_results, render_report
from quality.eval.check import (
    CheckContext,
    check_case_execution_completed,
    check_coverage_guard_no_analysis_without_data,
    check_expected_tool_contract,
    check_report_has_independent_curve_images,
    check_single_analysis_no_unrelated_tools,
    load_cases,
)


def test_fresh_root_copies_prompt_without_copying_env(tmp_path, monkeypatch) -> None:
    project = tmp_path / "project"
    (project / "resources" / "data").mkdir(parents=True)
    (project / "resources" / "templates").mkdir()
    (project / "agent" / "prompts").mkdir(parents=True)
    (project / "agent" / "prompts" / "system.md").write_text("system prompt", encoding="utf-8")
    (project / ".env").write_text("SECRET=value", encoding="utf-8")
    root = tmp_path / "isolated"
    root.mkdir()
    monkeypatch.setattr("quality.eval.eval_stage2.run_eval.PROJECT", project)

    fresh_root(root)

    assert (root / "agent" / "prompts" / "system.md").read_text(encoding="utf-8") == "system prompt"
    assert not (root / ".env").exists()
    assert all((root / "var" / name).is_dir() for name in ("outputs", "workspace", "logs"))


def test_fresh_root_can_remove_optional_inputs_from_isolated_copy(tmp_path, monkeypatch) -> None:
    project = tmp_path / "project"
    (project / "resources" / "data").mkdir(parents=True)
    (project / "resources" / "templates").mkdir()
    (project / "agent" / "prompts").mkdir(parents=True)
    (project / "resources" / "data" / "降雨数据.csv").write_text("rain", encoding="utf-8")
    (project / "resources" / "data" / "点位信息.xlsx").write_bytes(b"site")
    (project / "resources" / "templates" / "监测数据分析报告模板-更新.docx").write_bytes(b"docx")
    root = tmp_path / "isolated"
    root.mkdir()
    monkeypatch.setattr("quality.eval.eval_stage2.run_eval.PROJECT", project)

    fresh_root(root, {"fixture": "default", "remove_inputs": ["rainfall", "site_info"]})

    assert not (root / "resources" / "data" / "降雨数据.csv").exists()
    assert not (root / "resources" / "data" / "点位信息.xlsx").exists()
    assert (project / "resources" / "data" / "降雨数据.csv").exists()
    assert (project / "resources" / "data" / "点位信息.xlsx").exists()


def test_fresh_root_can_overlay_versioned_fixture_without_mutating_source(tmp_path, monkeypatch) -> None:
    project = tmp_path / "project"
    (project / "resources" / "data").mkdir(parents=True)
    (project / "resources" / "templates").mkdir()
    (project / "agent" / "prompts").mkdir(parents=True)
    (project / "resources" / "data" / "点位信息.xlsx").write_bytes(b"original")
    fixture = project / "quality" / "eval" / "fixtures" / "bad.xlsx"
    fixture.parent.mkdir(parents=True)
    fixture.write_bytes(b"adversarial")
    root = tmp_path / "isolated"
    root.mkdir()
    monkeypatch.setattr("quality.eval.eval_stage2.run_eval.PROJECT", project)

    fresh_root(root, {
        "fixture": "default",
        "overlay_inputs": {"site_info": "quality/eval/fixtures/bad.xlsx"},
    })

    assert (root / "resources" / "data" / "点位信息.xlsx").read_bytes() == b"adversarial"
    assert (project / "resources" / "data" / "点位信息.xlsx").read_bytes() == b"original"
    assert fixture.read_bytes() == b"adversarial"


def test_after_seed_mutation_only_changes_isolated_flow_copy(tmp_path: Path) -> None:
    project_flow = tmp_path / "project" / "resources" / "data" / "flow" / "W1.csv"
    isolated_flow = tmp_path / "isolated" / "resources" / "data" / "flow" / "W1.csv"
    project_flow.parent.mkdir(parents=True)
    isolated_flow.parent.mkdir(parents=True)
    project_flow.write_text("header\nvalue", encoding="utf-8")
    isolated_flow.write_text(project_flow.read_text(encoding="utf-8"), encoding="utf-8")

    apply_after_seed_mutation(tmp_path / "isolated", {"after_seed_mutation": "append_flow_newline"})

    assert project_flow.read_text(encoding="utf-8") == "header\nvalue"
    assert isolated_flow.read_text(encoding="utf-8") == "header\nvalue\n"


def test_tool_seq_reads_only_supplied_messages() -> None:
    call = SimpleNamespace(part_kind="tool-call", tool_name="check_data", args={"points": ["W1"]})
    text = SimpleNamespace(part_kind="text", content="done")
    messages = [SimpleNamespace(parts=[call, text])]

    assert tool_seq(messages) == [{"tool": "check_data", "args": {"points": ["W1"]}}]


def test_preserve_artifacts_replaces_stale_case_directory(tmp_path, monkeypatch) -> None:
    project = tmp_path / "project"
    root = tmp_path / "isolated"
    for name in ("outputs", "workspace", "logs"):
        (root / "var" / name).mkdir(parents=True)
    (root / "var" / "logs" / "trace.jsonl").write_text("trace", encoding="utf-8")
    stale = project / "quality" / "eval" / "eval_stage2" / "artifacts" / "E001"
    stale.mkdir(parents=True)
    (stale / "stale.txt").write_text("stale", encoding="utf-8")
    monkeypatch.setattr("quality.eval.eval_stage2.run_eval.STAGE_DIR", project / "quality" / "eval" / "eval_stage2")

    destination = preserve_artifacts(root, "E001")

    assert not (destination / "stale.txt").exists()
    assert (destination / "logs" / "trace.jsonl").read_text(encoding="utf-8") == "trace"


def test_preserve_artifacts_keeps_generated_report_charts(tmp_path, monkeypatch) -> None:
    project = tmp_path / "project"
    root = tmp_path / "isolated"
    for name in ("outputs", "workspace", "logs"):
        (root / "var" / name).mkdir(parents=True)
    charts = root / "results" / "generated" / "特征曲线图" / "W1_全时段"
    charts.mkdir(parents=True)
    (charts / "W1_流量特征曲线.png").write_bytes(b"flow-png")
    (charts / "W1_液位特征曲线.png").write_bytes(b"level-png")
    monkeypatch.setattr(
        "quality.eval.eval_stage2.run_eval.STAGE_DIR",
        project / "quality" / "eval" / "eval_stage2",
    )

    destination = preserve_artifacts(root, "CI003")

    preserved = destination / "results" / "generated" / "特征曲线图" / "W1_全时段"
    assert (preserved / "W1_流量特征曲线.png").read_bytes() == b"flow-png"
    assert (preserved / "W1_液位特征曲线.png").read_bytes() == b"level-png"


def test_preserve_artifacts_collapses_derived_case_directories(tmp_path, monkeypatch) -> None:
    project = tmp_path / "project"
    artifacts = project / "quality" / "eval" / "eval_stage2" / "artifacts"
    root = tmp_path / "isolated"
    for name in ("outputs", "workspace", "logs"):
        (root / "var" / name).mkdir(parents=True)
    (root / "var" / "outputs" / "latest.txt").write_text("latest", encoding="utf-8")
    stale = artifacts / "M003A_SCOPE_GUARD"
    stale.mkdir(parents=True)
    (stale / "stale.txt").write_text("stale", encoding="utf-8")
    monkeypatch.setattr("quality.eval.eval_stage2.run_eval.STAGE_DIR", project / "quality" / "eval" / "eval_stage2")

    destination = preserve_artifacts(root, "M003A_SCOPE_RULE2")

    assert canonical_case_id("M003A_SCOPE_RULE2") == "M003A"
    assert destination == artifacts / "M003A"
    assert not stale.exists()
    assert (destination / "outputs" / "latest.txt").read_text(encoding="utf-8") == "latest"


def test_normalize_multiturn_case_preserves_key_turns() -> None:
    case = normalize_case({
        "id": "M001",
        "category": "指代",
        "conversation_goal": "继承任务并替换点位",
        "state_under_test": ["task", "points"],
        "turns": [
            {"prompt": "先看 W1", "expect": "调用工具"},
            {
                "prompt": "W6 呢",
                "key": True,
                "expected": {
                    "response": "继承上下文",
                    "inherited": ["task"],
                    "replaced": ["points"],
                },
            },
        ],
    })

    assert [turn["key"] for turn in case["turns"]] == [False, True]
    assert case["turns"][1]["expect"] == "继承上下文"
    assert case["turns"][1]["expected"]["replaced"] == ["points"]
    assert case["conversation_goal"] == "继承任务并替换点位"
    assert case["state_under_test"] == ["task", "points"]


def test_normalize_structured_single_case_preserves_eval_contract() -> None:
    case = normalize_case({
        "id": "E021",
        "scenario": "缺少降雨资料时请求 RDII",
        "dimensions": {"task": "RDII", "data_state": "missing_rainfall"},
        "setup": {"fixture": "default", "remove_inputs": ["rainfall"]},
        "prompt": "分析 W1 第 6 场降雨的 RDII。",
        "expected": {
            "response": "说明缺少降雨资料",
            "forbidden": ["编造 RDII"],
        },
    })

    assert case["scenario"] == "缺少降雨资料时请求 RDII"
    assert case["dimensions"]["data_state"] == "missing_rainfall"
    assert case["setup"]["remove_inputs"] == ["rainfall"]
    assert case["turns"][0]["expect"] == "说明缺少降雨资料"
    assert case["expected"]["forbidden"] == ["编造 RDII"]


def test_validate_cases_rejects_duplicate_ids() -> None:
    duplicate = {"id": "E001", "prompt": "检查数据"}

    try:
        validate_cases([duplicate, duplicate])
    except ValueError as exc:
        assert "重复" in str(exc)
    else:
        raise AssertionError("duplicate Eval ids should fail validation")


def test_select_cases_keeps_requested_order_and_rejects_unknown_ids() -> None:
    cases = [{"id": "E001"}, {"id": "E002"}, {"id": "E003"}]

    assert [case["id"] for case in select_cases(cases, ["E003", "E001"])] == [
        "E003",
        "E001",
    ]
    with pytest.raises(ValueError, match="E999"):
        select_cases(cases, ["E999"])


def test_tree_snapshot_and_trace_evidence_capture_objective_state(tmp_path: Path) -> None:
    root = tmp_path / "root"
    root.mkdir()
    (root / "result.txt").write_text("done", encoding="utf-8")
    trace = SimpleNamespace(path=root / "trace.jsonl")
    trace.path.write_text(
        "\n".join([
            '{"run_id":"run-1","event":"tool_call","tool_name":"check_data"}',
            '{"run_id":"run-2","event":"tool_result","status":"ok"}',
        ]),
        encoding="utf-8",
    )

    snapshot = tree_snapshot(root)
    evidence = trace_evidence(trace, "run-1")

    assert {item["path"] for item in snapshot} == {"result.txt", "trace.jsonl"}
    assert len(evidence) == 1
    assert evidence[0]["tool_name"] == "check_data"


def test_multiturn_view_skips_meta_and_renders_turns(tmp_path: Path) -> None:
    source = tmp_path / "results.jsonl"
    destination = tmp_path / "report.html"
    source.write_text(
        "\n".join([
            '{"_meta":{"model":"test-model","case_count":1}}',
            '{"id":"M001","category":"指代","turns":['
            '{"n":1,"run_id":"run-1","prompt":"先看 W1","expect":"调用工具","key":true,'
            '"output":"完成","tool_calls":[{"tool":"check_data","args":"{\\"points\\":[\\"W1\\"]}"}]}],'
            '"trace":"trace.jsonl","error":null}',
        ]),
        encoding="utf-8",
    )

    meta, rows = load_results(source)
    count = render_report(source, destination)
    html = destination.read_text(encoding="utf-8")

    assert meta["model"] == "test-model"
    assert rows[0]["turns"][0]["tool_calls"][0]["args"] == {"points": ["W1"]}
    assert count == 1
    assert "逐轮人工判定" not in html
    assert "M001" in html and "先看 W1" in html and "调用工具" in html


def test_eval_view_loads_objective_check_sidecar(tmp_path: Path) -> None:
    source = tmp_path / "results_single.jsonl"
    destination = tmp_path / "report.html"
    row = {
        "id": "E001",
        "category": "analysis",
        "scenario": "Inspect W1",
        "dimensions": {"clarity": "clear"},
        "turns": [{
            "n": 1,
            "prompt": "Inspect W1",
            "expected": {"response": "Summarize status"},
            "output": "W1 is normal",
            "tool_calls": [],
            "trace_events": [],
        }],
    }
    source.write_text(json.dumps(row) + "\n", encoding="utf-8")
    checks = {"checks": [{
        "case_id": "E001",
        "check": "tool_contract",
        "basis": "trace",
        "status": "fail",
        "reason": "missing call",
        "turn": 1,
    }]}
    source.with_name("results_single_checks.json").write_text(
        json.dumps(checks), encoding="utf-8"
    )

    assert load_checks(source)[0]["status"] == "fail"
    assert render_report(source, destination) == 1
    html = destination.read_text(encoding="utf-8")
    assert "Drainage Agent Eval 评审" in html
    assert "自动失败" in html
    assert 'data-filter="uncertain">人工不确定' in html
    assert 'if(filter==="uncertain")return g.verdict==="uncertain"' in html
    assert "Inspect W1" in html
    assert "Summarize status" in html
    assert "W1 is normal" in html
    assert "整例人工判定" in html


def test_report_curve_check_reads_preserved_generated_charts(tmp_path: Path) -> None:
    root = tmp_path / "artifacts" / "CI003"
    outputs = root / "outputs"
    outputs.mkdir(parents=True)
    png = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
    )
    charts = root / "results" / "generated" / "特征曲线图" / "W1_全时段"
    charts.mkdir(parents=True)
    flow = charts / "W1_流量特征曲线.png"
    level = charts / "W1_液位特征曲线.png"
    flow.write_bytes(png)
    level.write_bytes(png)
    document = Document()
    document.add_picture(str(flow))
    document.add_picture(str(level))
    document.save(outputs / "W1_全时段_分析报告.docx")
    manifest = {
        "results": {"generate_report": {"params": {"points": ["W1"]}}}
    }
    (outputs / "manifest.json").write_text(json.dumps(manifest), encoding="utf-8")
    results = tmp_path / "results.jsonl"
    results.write_text(
        json.dumps({"id": "CI003", "root": str(root), "turns": []}) + "\n",
        encoding="utf-8",
    )

    case = load_cases(results)[0]
    checked = check_report_has_independent_curve_images(
        case, CheckContext(tmp_path, {"W1"}, None, None, None, None)
    )

    assert checked[0].status == "pass"


def test_completed_case_ids_ignores_meta_and_partial_line(tmp_path: Path) -> None:
    pending = tmp_path / "results.jsonl.tmp"
    pending.write_text(
        '{"_meta":{"case_count":2}}\n{"id":"M001","turns":[]}\n{"id":',
        encoding="utf-8",
    )

    assert completed_case_ids(pending) == {"M001"}


def test_compact_pending_results_keeps_last_complete_case(tmp_path: Path) -> None:
    pending = tmp_path / "results.jsonl.tmp"
    pending.write_text(
        "\n".join([
            '{"_meta":{"case_count":1}}',
            '{"id":"M001","turns":[{"output":"old"}]}',
            '{"id":"M001","turns":[{"output":"new"}]}',
            '{"id":',
        ]),
        encoding="utf-8",
    )

    compact_pending_results(pending)
    rows = [json.loads(line) for line in pending.read_text(encoding="utf-8").splitlines()]

    assert len(rows) == 2
    assert rows[1]["turns"][0]["output"] == "new"


def test_structured_expected_tools_are_loaded_and_checked(tmp_path: Path) -> None:
    results = tmp_path / "results.jsonl"
    results.write_text(json.dumps({
        "id": "E001",
        "expected": {"tools": {"must_call": ["check_data"], "must_not_call": ["generate_report"]}},
        "turns": [{
            "n": 1,
            "prompt": "检查数据",
            "tool_calls": [{"tool": "check_data", "args": {}}],
        }],
        "root": str(tmp_path / "artifacts"),
    }, ensure_ascii=False), encoding="utf-8")
    case = load_cases(results)[0]
    ctx = CheckContext(tmp_path, set(), None, None, None, None)

    checked = check_expected_tool_contract(case, ctx)

    assert case.expected["tools"]["must_call"] == ["check_data"]
    assert checked[0].status == "pass"


def test_case_execution_error_fails_even_when_no_turn_was_recorded(tmp_path: Path) -> None:
    results = tmp_path / "results.jsonl"
    results.write_text(json.dumps({
        "id": "M008",
        "turns": [],
        "root": str(tmp_path / "artifacts"),
        "error": "ValueError('day is out of range for month')",
    }, ensure_ascii=False), encoding="utf-8")
    case = load_cases(results)[0]

    checked = check_case_execution_completed(
        case, CheckContext(tmp_path, set(), None, None, None, None)
    )

    assert checked[0].status == "fail"
    assert "day is out of range for month" in checked[0].reason


def test_structured_per_turn_tool_contract_is_loaded_and_checked(tmp_path: Path) -> None:
    results = tmp_path / "results.jsonl"
    results.write_text(json.dumps({
        "id": "M001",
        "turns": [{
            "n": 1,
            "prompt": "W6 呢",
            "expected": {"tools": {"must_call": ["analyze_patterns"]}, "inherited": ["task"]},
            "tool_calls": [{"tool": "analyze_patterns", "args": {"points": ["W6"]}}],
        }],
        "root": str(tmp_path / "artifacts"),
    }, ensure_ascii=False), encoding="utf-8")
    case = load_cases(results)[0]
    ctx = CheckContext(tmp_path, set(), None, None, None, None)

    checked = check_expected_tool_contract(case, ctx)

    assert case.turns[0].expected["inherited"] == ["task"]
    assert checked[0].status == "pass"
    assert checked[0].turn == 1


def test_single_turn_tool_contract_is_not_checked_twice(tmp_path: Path) -> None:
    contract = {"must_call": ["check_data"]}
    results = tmp_path / "results.jsonl"
    results.write_text(json.dumps({
        "id": "E001",
        "expected": {"tools": contract},
        "turns": [{
            "n": 1,
            "prompt": "检查数据",
            "expected": {"tools": contract},
            "tool_calls": [{"tool": "check_data", "args": {}}],
        }],
        "root": str(tmp_path / "artifacts"),
    }, ensure_ascii=False), encoding="utf-8")
    case = load_cases(results)[0]

    checked = check_expected_tool_contract(
        case, CheckContext(tmp_path, set(), None, None, None, None)
    )

    assert len(checked) == 1
    assert checked[0].status == "pass"
    assert checked[0].turn == 1


def test_coverage_guard_allows_analysis_for_mixed_valid_and_invalid_points(tmp_path: Path) -> None:
    root = tmp_path / "artifacts"
    outputs = root / "outputs"
    outputs.mkdir(parents=True)
    manifest = {"results": {"analyze_patterns": {"artifacts": []}}}
    (outputs / "manifest.json").write_text(json.dumps(manifest), encoding="utf-8")
    results = tmp_path / "results.jsonl"
    results.write_text(json.dumps({
        "id": "CI005",
        "root": str(root),
        "turns": [{
            "n": 1,
            "prompt": "比较 W1、W6、W999 的旱天流量。",
            "expect": "指出 W999 无效并继续处理 W1/W6",
            "tool_calls": [{"tool": "analyze_patterns", "args": {"points": ["W1", "W6", "W999"]}}],
        }],
    }, ensure_ascii=False) + "\n", encoding="utf-8")

    checked = check_coverage_guard_no_analysis_without_data(
        load_cases(results)[0],
        CheckContext(tmp_path, {"W1", "W6"}, None, None, None, None),
    )

    assert checked[0].status == "skip"


def test_coverage_guard_allows_analysis_tool_to_return_needs_input(tmp_path: Path) -> None:
    root = tmp_path / "artifacts"
    outputs = root / "outputs"
    outputs.mkdir(parents=True)
    (outputs / "manifest.json").write_text(
        json.dumps({"results": {"assess_risk": {"status": "needs_input"}}}),
        encoding="utf-8",
    )
    trace = root / "trace.jsonl"
    trace.write_text(json.dumps({
        "event": "tool_result",
        "run_id": "run-1",
        "tool_name": "assess_risk",
        "status": "needs_input",
    }), encoding="utf-8")
    results = tmp_path / "results.jsonl"
    results.write_text(json.dumps({
        "id": "E009A",
        "root": str(root),
        "trace": str(trace),
        "turns": [{
            "n": 1,
            "run_id": "run-1",
            "prompt": "分析第 4 场降雨风险，该事件没有流量覆盖。",
            "expect": "说明无覆盖",
            "tool_calls": [{"tool": "assess_risk", "args": {"event_ids": [4]}}],
        }],
    }, ensure_ascii=False), encoding="utf-8")

    checked = check_coverage_guard_no_analysis_without_data(
        load_cases(results)[0],
        CheckContext(tmp_path, set(), None, None, None, None),
    )

    assert checked[0].status == "pass"


def test_rainy_risk_allows_event_response_analysis(tmp_path: Path) -> None:
    results = tmp_path / "results.jsonl"
    results.write_text(json.dumps({
        "id": "CI006",
        "root": str(tmp_path / "artifacts"),
        "turns": [{
            "n": 2,
            "prompt": "那就第 6 场。",
            "expect": "承接原任务并完成事件6雨天风险",
            "tool_calls": [
                {"tool": "assess_risk", "args": {"scope": "rainy", "event_ids": [6]}},
                {"tool": "analyze_event_response", "args": {"event_ids": [6]}},
            ],
        }],
    }, ensure_ascii=False) + "\n", encoding="utf-8")

    checked = check_single_analysis_no_unrelated_tools(
        load_cases(results)[0],
        CheckContext(tmp_path, set(), None, None, None, None),
    )

    assert checked[0].status == "pass"


def test_rain_only_report_does_not_require_dry_weather_curve_images(tmp_path: Path) -> None:
    from quality.eval.check import check_report_has_independent_curve_images

    root = tmp_path / "artifacts"
    outputs = root / "outputs"
    outputs.mkdir(parents=True)
    (outputs / "manifest.json").write_text(json.dumps({
        "results": {
            "generate_report": {
                "params": {"sections": ["响应", "RDII", "风险"], "event_ids": [6]}
            }
        }
    }, ensure_ascii=False), encoding="utf-8")
    (outputs / "雨天报告.docx").write_bytes(b"not inspected after scope skip")
    results = tmp_path / "results.jsonl"
    results.write_text(json.dumps({
        "id": "E016B",
        "root": str(root),
        "turns": [{"n": 1, "prompt": "生成含响应、RDII 和风险的报告"}],
    }, ensure_ascii=False), encoding="utf-8")

    checked = check_report_has_independent_curve_images(
        load_cases(results)[0],
        CheckContext(tmp_path, {"W1"}, None, None, None, None),
    )

    assert checked[0].status == "skip"
    assert "dry-weather" in checked[0].reason
