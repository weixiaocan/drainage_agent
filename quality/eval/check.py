from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Iterable, Literal

import pandas as pd
from docx import Document
from openpyxl import load_workbook

PROJECT = Path(__file__).resolve().parents[2]
if str(PROJECT) not in sys.path:
    sys.path.insert(0, str(PROJECT))

from analysis.io import load_flow, load_rain, load_sites


Status = Literal["pass", "fail", "skip"]
Basis = Literal["trace", "artifact"]

ANALYSIS_TOOLS = {
    "analyze_rainfall",
    "analyze_event_response",
    "analyze_patterns",
    "analyze_rdii",
    "assess_risk",
    "generate_report",
}
NON_ANALYSIS_TOOLS = {"list_results", "check_data", "data_filter", "run_python"}
REPORT_EXTENSIONS = {".doc", ".docx"}
POINT_RE = re.compile(r"(?<![A-Za-z0-9])W\d+(?![A-Za-z0-9])", re.IGNORECASE)

SITE_COLUMNS = ("point_id", "点位编号", "监测点编号", "点位", "安装点位", "安装监测点位")
TEMPLATE_BLACKLIST = (
    "1-1#",
    "1-9#",
    "44个流量监测点位",
    "32个点位",
    "13台流量监测设备",
    "____/__/__",
    "{{",
    "}}",
    "${",
    "TODO",
    "示例点位",
)
SECTION_KEYWORDS = {
    "rainfall": ("降雨分析", "降雨日分析", "降雨场次分析", "雨天事件统计", "RDII", "降雨入流"),
    "dry": ("旱天排污规律统计分析", "排污规律统计", "流量特征曲线", "液位特征曲线"),
    "risk": ("污水系统运行风险分析", "运行风险分析", "旱天风险", "雨天溢流风险"),
    "monitoring": ("监测概况", "监测设备安装", "监测数据质量"),
}


@dataclass
class ToolCall:
    tool: str
    args: dict[str, Any] = field(default_factory=dict)
    turn: int | None = None
    run_id: str | None = None


@dataclass
class TurnRecord:
    n: int
    prompt: str = ""
    expect: str = ""
    output: str = ""
    tool_calls: list[ToolCall] = field(default_factory=list)
    run_id: str | None = None


@dataclass
class CaseRecord:
    stage: str
    case_id: str
    root: Path
    trace: Path | None
    error: str | None
    turns: list[TurnRecord]
    manifest: dict[str, Any] = field(default_factory=dict)

    @property
    def outputs(self) -> Path:
        standard = self.root / "var" / "outputs"
        if standard.exists():
            return standard
        return self.root / "outputs"

    @property
    def all_tool_calls(self) -> list[ToolCall]:
        calls: list[ToolCall] = []
        for turn in self.turns:
            calls.extend(turn.tool_calls)
        return calls


@dataclass
class CheckResult:
    case_id: str
    check: str
    basis: Basis
    status: Status
    reason: str
    turn: int | None = None


@dataclass
class CheckContext:
    project: Path
    allowed_sites: set[str]
    flow_start: pd.Timestamp | None
    flow_end: pd.Timestamp | None
    rain_start: pd.Timestamp | None
    rain_end: pd.Timestamp | None


CheckFn = Callable[[CaseRecord, CheckContext], list[CheckResult]]


def _parse_args_value(value: Any) -> dict[str, Any]:
    if isinstance(value, dict):
        return value
    if value in (None, ""):
        return {}
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
            return parsed if isinstance(parsed, dict) else {}
        except json.JSONDecodeError:
            return {}
    return {}


def _tool_call(raw: dict[str, Any], turn: int | None = None, run_id: str | None = None) -> ToolCall:
    return ToolCall(
        tool=str(raw.get("tool") or raw.get("tool_name") or ""),
        args=_parse_args_value(raw.get("args")),
        turn=turn,
        run_id=run_id,
    )


def _stage_from_path(path: Path) -> str:
    parts = {part.lower() for part in path.parts}
    if "eval_stage1" in parts:
        return "stage1"
    if "eval_stage2" in parts:
        return "stage2"
    return path.parent.name


def load_cases(results_path: Path, artifacts_root: Path | None = None) -> list[CaseRecord]:
    cases: list[CaseRecord] = []
    stage = _stage_from_path(results_path)
    for line in results_path.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        record = json.loads(line)
        if "_meta" in record:
            continue
        case_id = str(record.get("id", ""))
        original_root = Path(record.get("root") or "")
        root = (artifacts_root / case_id) if artifacts_root else original_root
        trace_text = record.get("trace") or ""
        trace = None
        if trace_text:
            trace_path = Path(trace_text)
            if artifacts_root and original_root and trace_path.is_relative_to(original_root):
                trace_path = root / trace_path.relative_to(original_root)
            trace = trace_path
        turns: list[TurnRecord] = []
        if "turns" in record:
            for raw_turn in record.get("turns") or []:
                n = int(raw_turn.get("n") or len(turns) + 1)
                run_id = raw_turn.get("run_id")
                calls = [_tool_call(raw, turn=n, run_id=run_id) for raw in raw_turn.get("tool_calls") or []]
                turns.append(
                    TurnRecord(
                        n=n,
                        prompt=str(raw_turn.get("prompt") or ""),
                        expect=str(raw_turn.get("expect") or ""),
                        output=str(raw_turn.get("output") or ""),
                        tool_calls=calls,
                        run_id=run_id,
                    )
                )
        else:
            calls = [_tool_call(raw, turn=1) for raw in record.get("tool_calls") or []]
            turns.append(
                TurnRecord(
                    n=1,
                    prompt=str(record.get("prompt") or ""),
                    expect=str(record.get("pass_when") or ""),
                    output=str(record.get("output") or ""),
                    tool_calls=calls,
                )
            )
        manifest = _load_manifest(_outputs_root(root) / "manifest.json")
        cases.append(
            CaseRecord(
                stage=stage,
                case_id=case_id,
                root=root,
                trace=trace,
                error=record.get("error"),
                turns=turns,
                manifest=manifest,
            )
        )
    return cases


def _outputs_root(root: Path) -> Path:
    standard = root / "var" / "outputs"
    if standard.exists():
        return standard
    return root / "outputs"


def _load_manifest(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {}


def build_context(project: Path) -> CheckContext:
    sites = load_sites(root=project)
    allowed_sites: set[str] = set()
    for column in SITE_COLUMNS:
        if column in sites.columns:
            allowed_sites = {str(value).strip().upper() for value in sites[column].dropna() if str(value).strip()}
            break
    flow = load_flow(root=project)
    rain = load_rain(root=project)
    return CheckContext(
        project=project,
        allowed_sites=allowed_sites,
        flow_start=flow["timestamp"].min() if not flow.empty else None,
        flow_end=flow["timestamp"].max() if not flow.empty else None,
        rain_start=rain["timestamp"].min() if not rain.empty else None,
        rain_end=rain["timestamp"].max() if not rain.empty else None,
    )


def result(
    case: CaseRecord,
    check: str,
    basis: Basis,
    status: Status,
    reason: str,
    turn: int | None = None,
) -> CheckResult:
    return CheckResult(case.case_id, check, basis, status, reason, turn)


def _outputs_files(case: CaseRecord, suffixes: Iterable[str] | None = None) -> list[Path]:
    if not case.outputs.exists():
        return []
    wanted = {suffix.lower() for suffix in suffixes or []}
    files = [path for path in case.outputs.rglob("*") if path.is_file()]
    if wanted:
        files = [path for path in files if path.suffix.lower() in wanted]
    return sorted(files)


def _report_paths(case: CaseRecord) -> list[Path]:
    return [path for path in _outputs_files(case, {".doc", ".docx", ".md"}) if "分析报告" in path.stem or path.suffix.lower() in {".doc", ".docx"}]


def _combined_tables(case: CaseRecord) -> list[Path]:
    return [path for path in _outputs_files(case, {".xlsx"}) if "综合分析结果" in path.stem]


def _document_text(path: Path) -> str:
    doc = Document(path)
    values: list[str] = []
    values.extend(paragraph.text for paragraph in doc.paragraphs)
    for table in doc.tables:
        values.extend(cell.text for row in table.rows for cell in row.cells)
    for section in doc.sections:
        values.extend(paragraph.text for paragraph in section.header.paragraphs)
        values.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(value for value in values if value)


def _inline_shape_count(path: Path) -> int:
    return len(Document(path).inline_shapes)


def _manifest_result(case: CaseRecord, tool_name: str) -> dict[str, Any]:
    results = case.manifest.get("results") if isinstance(case.manifest, dict) else {}
    value = (results or {}).get(tool_name)
    return value if isinstance(value, dict) else {}


def _report_params(case: CaseRecord) -> dict[str, Any]:
    params = _manifest_result(case, "generate_report").get("params")
    return params if isinstance(params, dict) else {}


def _expected_points_for_report(case: CaseRecord, ctx: CheckContext) -> list[str]:
    params = _report_params(case)
    points = params.get("points")
    if isinstance(points, list) and points:
        return sorted({str(point).upper() for point in points})
    mentioned: set[str] = set()
    for path in _combined_tables(case):
        mentioned.update(_point_ids_from_workbook(path))
    if mentioned:
        return sorted(mentioned, key=_point_sort_key)
    return sorted(ctx.allowed_sites, key=_point_sort_key)


def _point_sort_key(point_id: str) -> tuple[str, int, str]:
    match = re.match(r"^([A-Za-z]+)(\d+)$", point_id)
    if match:
        return match.group(1), int(match.group(2)), point_id
    return point_id, -1, point_id


def _point_ids_from_workbook(path: Path) -> set[str]:
    points: set[str] = set()
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except Exception:
        return points
    for sheet in workbook.worksheets:
        rows = sheet.iter_rows(min_row=1, max_row=min(sheet.max_row, 20), values_only=True)
        header = next(rows, None)
        if not header:
            continue
        point_indexes = [
            idx
            for idx, value in enumerate(header)
            if str(value or "").strip() in {"点位编号", "监测点编号", "point_id", "安装监测点位"}
        ]
        for row in sheet.iter_rows(min_row=2, values_only=True):
            for idx in point_indexes:
                if idx < len(row) and row[idx]:
                    text = str(row[idx]).strip().upper()
                    if POINT_RE.fullmatch(text):
                        points.add(text)
    return points


def _has_partial_scope(call: ToolCall) -> bool:
    points = call.args.get("points")
    has_points = isinstance(points, list) and bool(points)
    return bool(has_points or call.args.get("start") or call.args.get("end"))


def _has_generate_report(case: CaseRecord) -> bool:
    return any(call.tool == "generate_report" for call in case.all_tool_calls) or bool(_manifest_result(case, "generate_report"))


def _trace_events(case: CaseRecord) -> list[dict[str, Any]]:
    if not case.trace or not case.trace.exists():
        return []
    events: list[dict[str, Any]] = []
    for line in case.trace.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        try:
            item = json.loads(line)
        except json.JSONDecodeError:
            continue
        if isinstance(item, dict):
            events.append(item)
    return events


def _turn_trace_events(case: CaseRecord, turn: TurnRecord) -> list[dict[str, Any]]:
    events = _trace_events(case)
    if not turn.run_id:
        return []
    return [event for event in events if event.get("run_id") == turn.run_id]


def _has_needs_confirmation(events: list[dict[str, Any]]) -> bool:
    return any(
        event.get("event") == "tool_result"
        and event.get("tool_name") == "data_filter"
        and event.get("status") == "needs_confirmation"
        for event in events
    )


def _tool_calls_after_data_filter_confirmation(events: list[dict[str, Any]]) -> list[str]:
    seen_confirmation = False
    calls: list[str] = []
    for event in events:
        if (
            event.get("event") == "tool_result"
            and event.get("tool_name") == "data_filter"
            and event.get("status") == "needs_confirmation"
        ):
            seen_confirmation = True
            continue
        if seen_confirmation and event.get("event") == "tool_call":
            calls.append(str(event.get("tool_name")))
    return calls


def _looks_like_no_coverage_request(text: str) -> bool:
    if "W999" in text or "2030" in text:
        return True
    if re.search(r"[12]\s*月", text) or re.search(r"一\s*月|二\s*月", text):
        return True
    if re.search(r"第\s*4\s*场", text):
        return True
    return False


def check_coverage_guard_no_analysis_without_data(case: CaseRecord, ctx: CheckContext) -> list[CheckResult]:
    name = "coverage_guard_no_analysis_without_data"
    if case.error:
        return [result(case, name, "trace", "skip", f"case has error: {case.error}")]
    results: list[CheckResult] = []
    forbidden = {"analyze_event_response", "analyze_patterns", "analyze_rdii", "assess_risk"}
    for turn in case.turns:
        text = f"{turn.prompt}\n{turn.expect}"
        if not _looks_like_no_coverage_request(text):
            continue
        if not turn.tool_calls:
            results.append(result(case, name, "trace", "skip", "no tool_calls available for this turn", turn.n))
            continue
        bad = [
            call.tool
            for call in turn.tool_calls
            if call.tool in forbidden and _manifest_result(case, call.tool)
        ]
        status: Status = "fail" if bad else "pass"
        reason = (
            f"no-coverage analysis produced manifest result(s): {bad}"
            if bad
            else "no forbidden analysis result was materialized"
        )
        results.append(result(case, name, "trace", status, reason, turn.n))
    if not results:
        return [result(case, name, "trace", "skip", "no explicit no-coverage request detected")]
    return results


def _primary_intents(text: str) -> set[str]:
    intents: set[str] = set()
    if any(token in text for token in ("排污规律", "特征曲线", "旱天统计")):
        intents.add("analyze_patterns")
    if "RDII" in text or "降雨入流" in text:
        intents.add("analyze_rdii")
    if any(token in text for token in ("风险", "溢流", "满管", "淤积")):
        intents.add("assess_risk")
    if any(token in text for token in ("降雨事件", "降雨分析", "最大雨", "雨量")):
        intents.add("analyze_rainfall")
    if any(token in text for token in ("响应", "事件响应")):
        intents.add("analyze_event_response")
    return intents


def check_single_analysis_no_unrelated_tools(case: CaseRecord, ctx: CheckContext) -> list[CheckResult]:
    name = "single_analysis_no_unrelated_tools"
    if case.error:
        return [result(case, name, "trace", "skip", f"case has error: {case.error}")]
    checks: list[CheckResult] = []
    for turn in case.turns:
        text = f"{turn.prompt}\n{turn.expect}"
        if _looks_like_no_coverage_request(text):
            continue
        if any(token in text for token in ("报告", "完整", "全部章节", "再加", "上述", "导出", "综合")):
            continue
        intents = _primary_intents(text)
        if len(intents) != 1:
            continue
        if not turn.tool_calls:
            checks.append(result(case, name, "trace", "skip", "no tool_calls available for this turn", turn.n))
            continue
        allowed = set(intents) | NON_ANALYSIS_TOOLS
        if "analyze_rdii" in intents:
            allowed.add("analyze_rainfall")
        if "analyze_event_response" in intents:
            allowed.add("analyze_rainfall")
        if "assess_risk" in intents:
            allowed.add("analyze_rainfall")
        bad = [call.tool for call in turn.tool_calls if call.tool in ANALYSIS_TOOLS and call.tool not in allowed]
        status: Status = "fail" if bad else "pass"
        reason = f"unrelated analysis tools called: {bad}; intent={sorted(intents)}" if bad else f"tools match intent={sorted(intents)}"
        checks.append(result(case, name, "trace", status, reason, turn.n))
    if not checks:
        return [result(case, name, "trace", "skip", "no single-analysis turn detected")]
    return checks


def check_partial_scope_no_combined_xlsx(case: CaseRecord, ctx: CheckContext) -> list[CheckResult]:
    name = "partial_scope_no_combined_xlsx"
    if case.error:
        return [result(case, name, "trace", "skip", f"case has error: {case.error}")]
    if not case.all_tool_calls:
        return [result(case, name, "trace", "skip", "no tool_calls available")]
    partial = any(_has_partial_scope(call) for call in case.all_tool_calls if call.tool in ANALYSIS_TOOLS)
    tables = _combined_tables(case)
    if not partial:
        return [result(case, name, "trace", "skip", "no partial point/time analysis detected")]
    if _has_generate_report(case):
        return [result(case, name, "trace", "skip", "partial scope belongs to report generation; covered by artifact checks")]
    if tables:
        return [result(case, name, "trace", "fail", f"partial analysis wrote combined table(s): {[p.name for p in tables]}")]
    return [result(case, name, "trace", "pass", "partial analysis did not write combined table")]


def check_report_extension_doc_or_docx(case: CaseRecord, ctx: CheckContext) -> list[CheckResult]:
    name = "report_extension_doc_or_docx"
    reports = _report_paths(case)
    if not reports:
        return [result(case, name, "artifact", "skip", "no report artifact found")]
    bad = [path.name for path in reports if path.suffix.lower() not in REPORT_EXTENSIONS]
    if bad:
        return [result(case, name, "artifact", "fail", f"report artifact is not doc/docx: {bad}")]
    return [result(case, name, "artifact", "pass", f"report artifact extensions ok: {[p.name for p in reports]}")]


def check_report_has_independent_curve_images(case: CaseRecord, ctx: CheckContext) -> list[CheckResult]:
    name = "report_has_independent_curve_images"
    reports = [path for path in _report_paths(case) if path.suffix.lower() == ".docx"]
    if not reports:
        return [result(case, name, "artifact", "skip", "no docx report artifact found")]
    expected_points = _expected_points_for_report(case, ctx)
    if not expected_points:
        return [result(case, name, "artifact", "skip", "no expected report points resolved")]
    curve_files = _outputs_files(case, {".png"})
    missing: list[str] = []
    for point in expected_points:
        flow = [path for path in curve_files if path.name == f"{point}_流量特征曲线.png"]
        level = [path for path in curve_files if path.name == f"{point}_液位特征曲线.png"]
        if not flow:
            missing.append(f"{point} flow")
        if not level:
            missing.append(f"{point} level")
    shape_failures: list[str] = []
    minimum_curve_images = 2 * len(expected_points)
    for report in reports:
        count = _inline_shape_count(report)
        if count < minimum_curve_images:
            shape_failures.append(f"{report.name}: inline images {count} < expected curve images {minimum_curve_images}")
    if missing or shape_failures:
        pieces = []
        if missing:
            pieces.append(f"missing curve files: {missing[:20]}")
        if shape_failures:
            pieces.append("; ".join(shape_failures))
        return [result(case, name, "artifact", "fail", " | ".join(pieces))]
    return [result(case, name, "artifact", "pass", f"{len(expected_points)} point(s) have independent flow/level curve files and report images")]


def check_report_no_placeholders_or_fake_sites(case: CaseRecord, ctx: CheckContext) -> list[CheckResult]:
    name = "report_no_placeholders_or_fake_sites"
    reports = [path for path in _report_paths(case) if path.suffix.lower() == ".docx"]
    if not reports:
        return [result(case, name, "artifact", "skip", "no docx report artifact found")]
    failures: list[str] = []
    for report in reports:
        text = _document_text(report)
        blacklisted = [item for item in TEMPLATE_BLACKLIST if item in text]
        mentioned = {match.group(0).upper() for match in POINT_RE.finditer(text)}
        fake = sorted(mentioned - ctx.allowed_sites, key=_point_sort_key)
        if blacklisted:
            failures.append(f"{report.name}: template residue {blacklisted}")
        if fake:
            failures.append(f"{report.name}: non-whitelisted site(s) {fake}")
    if failures:
        return [result(case, name, "artifact", "fail", " | ".join(failures))]
    return [result(case, name, "artifact", "pass", "no template blacklist terms or fake site ids found")]


def _selected_section_kinds(case: CaseRecord) -> set[str]:
    params = _report_params(case)
    sections = params.get("sections")
    if not isinstance(sections, list) or not sections:
        return set()
    text = "\n".join(str(section) for section in sections)
    selected: set[str] = set()
    if any(token in text for token in ("降雨", "雨天", "RDII", "事件响应")):
        selected.add("rainfall")
    if any(token in text for token in ("旱天", "排污规律", "特征曲线")):
        selected.add("dry")
    if "风险" in text:
        selected.add("risk")
    if any(token in text for token in ("监测", "概况", "数据质量")):
        selected.add("monitoring")
    return selected


def check_report_excludes_unselected_sections(case: CaseRecord, ctx: CheckContext) -> list[CheckResult]:
    name = "report_excludes_unselected_sections"
    reports = [path for path in _report_paths(case) if path.suffix.lower() == ".docx"]
    if not reports:
        return [result(case, name, "artifact", "skip", "no docx report artifact found")]
    selected = _selected_section_kinds(case)
    if not selected:
        return [result(case, name, "artifact", "skip", "generate_report sections unavailable")]
    failures: list[str] = []
    for report in reports:
        text = _document_text(report)
        for kind, keywords in SECTION_KEYWORDS.items():
            if kind in selected:
                continue
            found = [keyword for keyword in keywords if keyword in text]
            if found:
                failures.append(f"{report.name}: unselected {kind} section keyword(s) {found[:5]}")
    if failures:
        return [result(case, name, "artifact", "fail", " | ".join(failures))]
    return [result(case, name, "artifact", "pass", f"only selected section kinds present: {sorted(selected)}")]


def _date_forms(ts: pd.Timestamp) -> set[str]:
    return {
        ts.strftime("%Y/%m/%d"),
        ts.strftime("%Y-%m-%d"),
        f"{ts.year}年{ts.month}月{ts.day}日",
        f"{ts.year}/{ts.month:02d}/{ts.day:02d}日",
    }


def check_report_period_matches_real_data_bounds(case: CaseRecord, ctx: CheckContext) -> list[CheckResult]:
    name = "report_period_matches_real_data_bounds"
    reports = [path for path in _report_paths(case) if path.suffix.lower() == ".docx"]
    if not reports:
        return [result(case, name, "artifact", "skip", "no docx report artifact found")]
    if ctx.flow_start is None or ctx.flow_end is None:
        return [result(case, name, "artifact", "skip", "real flow data bounds unavailable")]
    params = _report_params(case)
    if params.get("start") or params.get("end"):
        return [result(case, name, "artifact", "skip", "report has explicit start/end; skip full-period guard for baseline compatibility")]
    start_forms = _date_forms(ctx.flow_start)
    end_forms = _date_forms(ctx.flow_end)
    failures: list[str] = []
    for report in reports:
        text = _document_text(report)
        has_start = any(form in text for form in start_forms)
        has_end = any(form in text for form in end_forms)
        if not (has_start and has_end):
            failures.append(
                f"{report.name}: expected real flow bounds {ctx.flow_start.date()} to {ctx.flow_end.date()} not found"
            )
    if failures:
        return [result(case, name, "artifact", "fail", " | ".join(failures))]
    return [result(case, name, "artifact", "pass", f"report contains real flow bounds {ctx.flow_start.date()} to {ctx.flow_end.date()}")]


def check_combined_table_name_matches_report(case: CaseRecord, ctx: CheckContext) -> list[CheckResult]:
    name = "combined_table_name_matches_report"
    reports = [path for path in _report_paths(case) if path.suffix.lower() in REPORT_EXTENSIONS]
    if not reports:
        return [result(case, name, "artifact", "skip", "no report artifact found")]
    tables = _combined_tables(case)
    if not tables:
        return [result(case, name, "artifact", "skip", "no combined table artifact found")]
    failures: list[str] = []
    for report in reports:
        expected = f"{report.stem.replace('分析报告', '综合分析结果')}.xlsx"
        if not any(table.name == expected for table in tables):
            failures.append(f"{report.name}: expected paired table {expected}, found {[table.name for table in tables]}")
    if failures:
        return [result(case, name, "artifact", "fail", " | ".join(failures))]
    return [result(case, name, "artifact", "pass", "combined table name matches report name")]


def _rainfall_event_ids(path: Path) -> list[int]:
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except Exception:
        return []
    ids: list[int] = []
    for sheet in workbook.worksheets:
        if "降雨" not in sheet.title and "雨" not in sheet.title:
            continue
        header = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True), None)
        if not header:
            continue
        indexes = [idx for idx, value in enumerate(header) if str(value or "").strip() in {"场次编号", "降雨场次编号", "event_id"}]
        for row in sheet.iter_rows(min_row=2, values_only=True):
            for idx in indexes:
                if idx < len(row) and row[idx] not in (None, ""):
                    try:
                        ids.append(int(row[idx]))
                    except (TypeError, ValueError):
                        pass
    return sorted(set(ids))


def check_rainfall_event_ids_contiguous_in_window(case: CaseRecord, ctx: CheckContext) -> list[CheckResult]:
    name = "rainfall_event_ids_contiguous_in_window"
    if case.error:
        return [result(case, name, "artifact", "skip", f"case has error: {case.error}")]
    tables = _combined_tables(case)
    if not tables:
        return [result(case, name, "artifact", "skip", "no combined table artifact found")]
    checked = False
    failures: list[str] = []
    for table in tables:
        ids = _rainfall_event_ids(table)
        if not ids:
            continue
        checked = True
        expected = list(range(1, max(ids) + 1))
        if ids != expected:
            failures.append(f"{table.name}: event ids {ids} are not contiguous from 1 ({expected})")
    if failures:
        return [result(case, name, "artifact", "fail", " | ".join(failures))]
    if not checked:
        return [result(case, name, "artifact", "skip", "no rainfall event id table found")]
    return [result(case, name, "artifact", "pass", "rainfall event ids are contiguous from 1")]


def check_hitl_filter_confirmation(case: CaseRecord, ctx: CheckContext) -> list[CheckResult]:
    name = "hitl_filter_confirmation"
    marked_turns = [turn for turn in case.turns if "hitl_" in turn.expect]
    if not marked_turns:
        return [result(case, name, "trace", "skip", "no hitl filter expectation")]
    if case.error:
        return [result(case, name, "trace", "fail", f"case has error: {case.error}")]
    checks: list[CheckResult] = []
    for turn in marked_turns:
        events = _turn_trace_events(case, turn)
        if "hitl_stop_after_filter" in turn.expect:
            bad_after = _tool_calls_after_data_filter_confirmation(events)
            if not _has_needs_confirmation(events):
                checks.append(result(case, name, "trace", "fail", "data_filter did not return needs_confirmation", turn.n))
            elif bad_after:
                checks.append(result(case, name, "trace", "fail", f"tool calls after confirmation stop: {bad_after}", turn.n))
            else:
                checks.append(result(case, name, "trace", "pass", "data_filter hard-stopped with needs_confirmation", turn.n))
        if "hitl_resume_without_refilter" in turn.expect:
            refilter = [call.tool for call in turn.tool_calls if call.tool == "data_filter"]
            if refilter:
                checks.append(result(case, name, "trace", "fail", "confirmation turn reran data_filter", turn.n))
            elif "确认" in turn.output and "筛选结果" in turn.output and not turn.tool_calls:
                checks.append(result(case, name, "trace", "fail", "confirmation turn appears to stop again without analysis", turn.n))
            else:
                checks.append(result(case, name, "trace", "pass", "confirmation resumed without rerunning data_filter", turn.n))
        if "hitl_confirmed_fresh_no_repeat_stop" in turn.expect:
            if _has_needs_confirmation(events) or "请确认或修改后告知继续" in turn.output:
                checks.append(result(case, name, "trace", "fail", "fresh confirmed filter result stopped again", turn.n))
            else:
                checks.append(result(case, name, "trace", "pass", "fresh confirmed filter result did not repeat-stop", turn.n))
        if "hitl_ambiguous_confirmation" in turn.expect:
            if turn.tool_calls:
                checks.append(result(case, name, "trace", "fail", f"ambiguous confirmation called tools: {[c.tool for c in turn.tool_calls]}", turn.n))
            elif "确认用当前筛选结果" not in turn.output:
                checks.append(result(case, name, "trace", "fail", "ambiguous confirmation did not ask clarification", turn.n))
            else:
                checks.append(result(case, name, "trace", "pass", "ambiguous confirmation asked clarification without tools", turn.n))
    return checks


CHECKS: list[CheckFn] = [
    check_hitl_filter_confirmation,
    check_coverage_guard_no_analysis_without_data,
    check_single_analysis_no_unrelated_tools,
    check_partial_scope_no_combined_xlsx,
    check_report_extension_doc_or_docx,
    check_report_has_independent_curve_images,
    check_report_no_placeholders_or_fake_sites,
    check_report_excludes_unselected_sections,
    check_report_period_matches_real_data_bounds,
    check_combined_table_name_matches_report,
    check_rainfall_event_ids_contiguous_in_window,
]


def run_checks(cases: list[CaseRecord], ctx: CheckContext) -> list[CheckResult]:
    results: list[CheckResult] = []
    for case in cases:
        for check in CHECKS:
            try:
                results.extend(check(case, ctx))
            except Exception as exc:
                results.append(
                    result(
                        case,
                        check.__name__.removeprefix("check_"),
                        "artifact",
                        "fail",
                        f"check crashed: {exc!r}",
                    )
                )
    return results


def _status_rank(status: Status) -> int:
    return {"fail": 0, "pass": 1, "skip": 2}[status]


def print_text_report(results: list[CheckResult]) -> None:
    by_case: dict[str, list[CheckResult]] = {}
    for item in results:
        by_case.setdefault(item.case_id, []).append(item)
    for case_id in sorted(by_case):
        print(f"\n[{case_id}]")
        for item in sorted(by_case[case_id], key=lambda value: (_status_rank(value.status), value.check, value.turn or 0)):
            turn = f" turn={item.turn}" if item.turn is not None else ""
            print(f"  {item.status.upper():4} {item.check} ({item.basis}{turn}) - {item.reason}")
    passed = sum(1 for item in results if item.status == "pass")
    failed_items = [item for item in results if item.status == "fail"]
    skipped = sum(1 for item in results if item.status == "skip")
    failed_cases = sorted({item.case_id for item in failed_items})
    print(
        f"\n客观项 {passed} 通过/{len(failed_items)} 失败/{skipped} 跳过"
        + (f"，涉及用例: {', '.join(failed_cases)}" if failed_cases else "")
    )


def print_summary_report(results: list[CheckResult]) -> None:
    passed = sum(1 for item in results if item.status == "pass")
    failed_items = [item for item in results if item.status == "fail"]
    skipped = sum(1 for item in results if item.status == "skip")
    print(f"客观项 {passed} 通过 / {len(failed_items)} 失败 / {skipped} 跳过")
    if not failed_items:
        return
    print("失败项:")
    for item in sorted(failed_items, key=lambda value: (value.case_id, value.check, value.turn or 0)):
        turn = f" turn={item.turn}" if item.turn is not None else ""
        print(f"  - {item.case_id}{turn} | {item.check} | {item.reason}")


def default_results_paths(stage: str) -> list[Path]:
    if stage == "stage1":
        return [PROJECT / "quality" / "eval" / "eval_stage1" / "results.jsonl"]
    if stage == "stage2":
        return [PROJECT / "quality" / "eval" / "eval_stage2" / "results.jsonl"]
    return [
        PROJECT / "quality" / "eval" / "eval_stage1" / "results.jsonl",
        PROJECT / "quality" / "eval" / "eval_stage2" / "results.jsonl",
    ]


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Objective eval artifact checker.")
    parser.add_argument("--stage", choices=["stage1", "stage2", "all"], default="all")
    parser.add_argument("--results", action="append", type=Path, help="results.jsonl path; can be passed multiple times")
    parser.add_argument("--artifacts-root", type=Path, default=None, help="override artifacts root for injection copies")
    parser.add_argument("--format", choices=["text", "json"], default="text")
    args = parser.parse_args(argv)

    results_paths = args.results or default_results_paths(args.stage)
    cases: list[CaseRecord] = []
    for path in results_paths:
        if not path.exists():
            print(f"missing results file: {path}", file=sys.stderr)
            return 2
        cases.extend(load_cases(path, artifacts_root=args.artifacts_root))
    ctx = build_context(PROJECT)
    results = run_checks(cases, ctx)
    if args.format == "json":
        print(json.dumps([item.__dict__ for item in results], ensure_ascii=False, indent=2))
    else:
        print_text_report(results)
    return 1 if any(item.status == "fail" for item in results) else 0


if __name__ == "__main__":
    raise SystemExit(main())
